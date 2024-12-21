import streamlit as st
import pandas as pd
import openai
from langchain.llms import OpenAI
from langchain.chat_models import ChatOpenAI
from langchain.document_loaders import PyPDFLoader
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import FAISS
from langchain.chains import ConversationalRetrievalChain
from langchain.memory import ConversationBufferMemory
from markdown_pdf import MarkdownPdf, Section
from io import BytesIO
import re
from docx import Document


# Configuration de l'API OpenAI
hide_streamlit_style = """
    <style>
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    </style>
    """
st.markdown(hide_streamlit_style, unsafe_allow_html=True)

# Configuration de l'API OpenAI

api_key = st.secrets["API_KEY"]
openai.api_key = api_key


# Initialiser le dictionnaire principal dans session_state
if "data" not in st.session_state:
    st.session_state["data"] = {}
    
# Section 1 : Informations Générales
def page_informations_generales():
    st.title("Informations Générales")
    
    # Accès au dictionnaire principal
    data = st.session_state["data"]
    
    # Collecte des entrées et stockage dans le dictionnaire principal
    data["informations_generales"] = data.get("informations_generales", {})
    info = data["informations_generales"]
    info["prenom_nom"] = st.text_input("Prénom, nom :", value=info.get("prenom_nom", ""))
    info["intitule_projet"] = st.text_input("Intitulé de votre projet :", value=info.get("intitule_projet", ""))
    info["statut_juridique"] = st.selectbox(
        "Votre statut juridique :",
        ["Micro-entreprise", "EURL", "SARL", "SAS", "SASU"],
        index=["Micro-entreprise", "EURL", "SARL", "SAS", "SASU"].index(info.get("statut_juridique", "Micro-entreprise"))
    )
    info["telephone"] = st.text_input("Votre numéro de téléphone :", value=info.get("telephone", ""))
    info["email"] = st.text_input("Votre adresse e-mail :", value=info.get("email", ""))
    info["ville"] = st.text_input("Votre ville ou commune d'activité :", value=info.get("ville", ""))
    info["type_vente"] = st.selectbox(
        "Vente de marchandises ou de services ?",
        ["Marchandises", "Services", "Mixte"],
        index=["Marchandises", "Services", "Mixte"].index(info.get("type_vente", "Marchandises"))
    )
    
    # Mise à jour des données dans le dictionnaire principal
    st.session_state["data"]["informations_generales"] = info

# Section 2 : Besoins de Démarrage
def page_besoins_demarrage():
    st.title("Besoins de Démarrage")
    
    # Accès au dictionnaire principal
    data = st.session_state["data"]
    
    # Liste des besoins
    besoins = [
        "Frais d’établissement", "Frais d’ouverture de compteurs", "Logiciels, formations",
        "Dépôt de marque", "Droits d’entrée", "Achat fonds de commerce ou parts",
        "Droit au bail", "Caution ou dépôt de garantie", "Frais de dossier",
        "Frais de notaire", "Enseigne et éléments de communication", "Véhicule",
        "Matériel professionnel", "Matériel autre", "Matériel de bureau",
        "Stock de matières et produits", "Trésorerie de départ"
    ]
    
    # Initialiser le dictionnaire pour stocker les besoins
    data["besoins_demarrage"] = data.get("besoins_demarrage", {})
    besoins_demarrage = data["besoins_demarrage"]
    
    total_besoins = 0.0
    
    for besoin in besoins:
        montant = st.number_input(
            f"{besoin} (€)",
            min_value=0.0,
            key=f"besoin_{besoin}",
            value=besoins_demarrage.get(besoin, 0.0)
        )
        besoins_demarrage[besoin] = montant
        total_besoins += montant
    
    data["total_besoins"] = total_besoins
    
    st.write("---")
    st.markdown(f"**Total des Besoins de Démarrage :** {total_besoins:.2f} €")
    
    # Durée d'amortissement
    data["duree_amortissement"] = st.number_input(
        "Durée d'amortissement des investissements (en années) :",
        min_value=1,
        key="duree_amortissement",
        value=data.get("duree_amortissement", 3)
    )
    
    # Mise à jour des données dans le dictionnaire principal
    st.session_state["data"] = data
    
import streamlit as st

def calculer_pret_interet_fixe(montant, taux_annuel, duree_mois):
    """
    Calcule les détails d'un prêt avec intérêts fixes par mois.

    Args:
        montant (float): Montant du prêt en euros.
        taux_annuel (float): Taux d'intérêt annuel en pourcentage.
        duree_mois (int): Durée du prêt en mois.

    Returns:
        dict: Détails du prêt incluant mensualité, total à rembourser, principal mensuel,
              intérêts totaux et intérêts par année.
    """
    if duree_mois <= 0:
        return {
            "mensualite": 0.0,
            "total_a_rembourser": 0.0,
            "principal_mensuel": 0.0,
            "interet_mensuel": 0.0,
            "interets_totaux": 0.0,
            "interets_annee1": 0.0,
            "interets_annee2": 0.0,
            "interets_annee3": 0.0
        }

    taux_mensuel = taux_annuel / 100 / 12

    # Calcul de la mensualité en utilisant la formule PMT
    try:
        mensualite = (taux_mensuel * montant) / (1 - (1 + taux_mensuel) ** (-duree_mois))
    except ZeroDivisionError:
        mensualite = 0.0

    # Principal mensuel fixe
    principal_mensuel = montant / duree_mois

    # Intérêt mensuel
    interet_mensuel = mensualite - principal_mensuel

    # Total à rembourser
    total_a_rembourser = mensualite * duree_mois

    # Intérêts totaux
    interets_totaux = interet_mensuel * duree_mois

    # Intérêts par année, limités à 12 mois maximum
    interets_annee1 = interet_mensuel * min(duree_mois, 12)
    interets_annee2 = interet_mensuel * min(max(duree_mois - 12, 0), 12)
    interets_annee3 = interet_mensuel * min(max(duree_mois - 24, 0), 12)

    return {
        "mensualite": round(mensualite, 2),
        "total_a_rembourser": round(total_a_rembourser, 2),
        "principal_mensuel": round(principal_mensuel, 2),
        "interet_mensuel": round(interet_mensuel, 2),
        "interets_totaux": round(interets_totaux, 2),
        "interets_annee1": round(interets_annee1, 2),
        "interets_annee2": round(interets_annee2, 2),
        "interets_annee3": round(interets_annee3, 2)
    }

def page_financement():
    st.title("Financement des Besoins de Démarrage")
    
    data = st.session_state.get("data", {})
    
    # Initialiser la section des financements
    if "financements" not in data:
        data["financements"] = {}
    
    financements_dict = data["financements"]
    
    total_financement = 0.0
    
    st.subheader("Apports")
    
    # Apport personnel ou familial
    apport_personnel = st.number_input(
        "Apport personnel ou familial (€)",
        min_value=0.0,
        key="financement_apport_personnel",
        value=financements_dict.get("Apport personnel ou familial", 4000.00)
    )
    financements_dict["Apport personnel ou familial"] = apport_personnel
    total_financement += apport_personnel
    
    # Apports en nature (en valeur)
    apport_nature = st.number_input(
        "Apports en nature (en valeur) (€)",
        min_value=0.0,
        key="financement_apport_nature",
        value=financements_dict.get("Apports en nature (en valeur)", 1200.00)
    )
    financements_dict["Apports en nature (en valeur)"] = apport_nature
    total_financement += apport_nature
    
    st.subheader("Prêts")
    
    # Nombre de prêts (maximum 3)
    num_prets = 3  # Limité à 3 prêts comme demandé
    
    interets_prets = {
        "annee1": 0.0,
        "annee2": 0.0,
        "annee3": 0.0
    }
    
    for i in range(1, num_prets + 1):
        st.markdown(f"#### Prêt {i}")
        pret_name = st.text_input(
            f"Nom du prêt {i}",
            value=financements_dict.get(f"Prêt {i}", {}).get("nom", f"Prêt {i}"),
            key=f"pret_{i}_nom"
        )
        pret_montant = st.number_input(
            f"Montant du {pret_name} (€)",
            min_value=0.0,
            value=financements_dict.get(f"Prêt {i}", {}).get("montant", 0.0),
            key=f"pret_{i}_montant"
        )
        pret_taux = st.number_input(
            f"Taux du {pret_name} (%)",
            min_value=0.0,
            max_value=100.0,
            value=financements_dict.get(f"Prêt {i}", {}).get("taux", 0.0),
            key=f"pret_{i}_taux"
        )
        pret_duree = st.number_input(
            f"Durée du {pret_name} (en mois)",
            min_value=1,
            value=financements_dict.get(f"Prêt {i}", {}).get("duree", 12),
            key=f"pret_{i}_duree"
        )
        
        # Stocker les détails du prêt
        financements_dict[f"Prêt {i}"] = {
            "nom": pret_name,
            "montant": pret_montant,
            "taux": pret_taux,
            "duree": pret_duree
        }
        total_financement += pret_montant
        
        # Calculer les détails du remboursement du prêt
        if pret_montant > 0 and pret_taux > 0 and pret_duree > 0:
            pret_info = calculer_pret_interet_fixe(pret_montant, pret_taux, pret_duree)
            # Stocker les résultats du calcul
            financements_dict[f"Prêt {i}"].update(pret_info)
            # Ajouter les intérêts par année
            interets_prets["annee1"] += pret_info["interets_annee1"]
            interets_prets["annee2"] += pret_info["interets_annee2"]
            interets_prets["annee3"] += pret_info["interets_annee3"]
            
            # Afficher les détails du prêt pour vérification
            st.write(f"**Détails du {pret_name}:**")
            st.write(f"Mensualité : {pret_info['mensualite']:.2f} €")
            st.write(f"Total à rembourser : {pret_info['total_a_rembourser']:.2f} €")
            st.write(f"Principal mensuel : {pret_info['principal_mensuel']:.2f} €")
            st.write(f"Intérêt mensuel : {pret_info['interet_mensuel']:.2f} €")
            st.write(f"Intérêts totaux : {pret_info['interets_totaux']:.2f} €")
            st.write(f"Intérêts Année 1 : {pret_info['interets_annee1']:.2f} €")
            st.write(f"Intérêts Année 2 : {pret_info['interets_annee2']:.2f} €")
            st.write(f"Intérêts Année 3 : {pret_info['interets_annee3']:.2f} €")
            st.write("---")
    
    st.subheader("Subventions")
    
    # Nombre de subventions (maximum 2)
    num_subventions = 2  # Limité à 2 subventions comme demandé
    
    for i in range(1, num_subventions + 1):
        st.markdown(f"#### Subvention {i}")
        subvention_name = st.text_input(
            f"Nom de la subvention {i}",
            value=financements_dict.get(f"Subvention {i}", {}).get("nom", f"Subvention {i}"),
            key=f"subvention_{i}_nom"
        )
        subvention_montant = st.number_input(
            f"Montant de {subvention_name} (€)",
            min_value=0.0,
            value=financements_dict.get(f"Subvention {i}", {}).get("montant", 0.0),
            key=f"subvention_{i}_montant"
        )
        # Stocker les détails de la subvention
        financements_dict[f"Subvention {i}"] = {
            "nom": subvention_name,
            "montant": subvention_montant
        }
        total_financement += subvention_montant
    
    st.subheader("Autres Financements")
    
    # Autre financement
    autre_financement = st.number_input(
        "Autre financement (€)",
        min_value=0.0,
        key="financement_autre",
        value=financements_dict.get("Autre financement", 1000.00)
    )
    financements_dict["Autre financement"] = autre_financement
    total_financement += autre_financement
    
    st.write("---")
    st.markdown(f"**Total des Financements :** {total_financement:,.2f} €")
    
    # Validation du total des financements
    besoin_total = data.get("besoins", 0.0)  # Assurez-vous que cette clé existe dans vos données
    if besoin_total > 0 and total_financement != besoin_total:
        st.error(f"Le total des financements ({total_financement:,.2f} €) ne correspond pas au besoin total ({besoin_total:,.2f} €). Veuillez ajuster les montants.")
    elif besoin_total > 0:
        st.success(f"Le total des financements correspond au besoin total ({besoin_total:,.2f} €).")
    
    # Stocker les données dans la session
    data["financements"] = financements_dict
    data["total_financement"] = total_financement
    data["interets_prets"] = interets_prets  # Stocker les intérêts des prêts
    
    st.session_state["data"] = data



# Section 4 : Charges Fixes sur 3 Années
def page_charges_fixes():
    st.title("Charges Fixes sur 3 Années")
    
    data = st.session_state["data"]
    
    charges_fixes = [
        "Assurances véhicule et RC pro", "Téléphone, internet", "Autres abonnements",
        "Carburant", "Frais de déplacement / hébergement", "Eau, électricité, gaz",
        "Mutuelle", "Fournitures diverses", "Entretien Moto livraison et matériel",
        "Nettoyage des locaux", "Budget publicité et communication", "Emplacements",
        "Expert comptable, avocats", "Frais bancaires et terminal carte bleue", "Taxes, CFE"
    ]
    
    data["charges_fixes"] = data.get("charges_fixes", {"annee1": {}, "annee2": {}, "annee3": {}})
    charges_fixes_dict = data["charges_fixes"]
    
    total_annee1 = total_annee2 = total_annee3 = 0.0
    
    st.subheader("Charges Fixes par Défaut")
    for charge in charges_fixes:
        col1, col2, col3 = st.columns(3)
        with col1:
            montant1 = st.number_input(
                f"{charge} - Année 1 (€)",
                min_value=0.0,
                key=f"charge_{charge}_annee1",
                value=charges_fixes_dict["annee1"].get(charge, 0.0)
            )
            charges_fixes_dict["annee1"][charge] = montant1
        with col2:
            montant2 = st.number_input(
                f"{charge} - Année 2 (€)",
                min_value=0.0,
                key=f"charge_{charge}_annee2",
                value=charges_fixes_dict["annee2"].get(charge, 0.0)
            )
            charges_fixes_dict["annee2"][charge] = montant2
        with col3:
            montant3 = st.number_input(
                f"{charge} - Année 3 (€)",
                min_value=0.0,
                key=f"charge_{charge}_annee3",
                value=charges_fixes_dict["annee3"].get(charge, 0.0)
            )
            charges_fixes_dict["annee3"][charge] = montant3
        
        total_annee1 += montant1
        total_annee2 += montant2
        total_annee3 += montant3
    
    # Charges supplémentaires
    st.write("---")
    st.subheader("Ajouter des Charges Supplémentaires")
    
    if "charges_supplementaires" not in data:
        data["charges_supplementaires"] = []
    
    nouvelle_charge = st.text_input("Nom de la nouvelle charge :", key="nouvelle_charge")
    
    if st.button("Ajouter la charge"):
        if nouvelle_charge and nouvelle_charge not in data["charges_supplementaires"]:
            data["charges_supplementaires"].append(nouvelle_charge)
            charges_fixes_dict["annee1"][nouvelle_charge] = 0.0
            charges_fixes_dict["annee2"][nouvelle_charge] = 0.0
            charges_fixes_dict["annee3"][nouvelle_charge] = 0.0
            st.session_state["nouvelle_charge"] = ""
    
    for charge in data["charges_supplementaires"]:
        col1, col2, col3 = st.columns(3)
        with col1:
            montant1 = st.number_input(
                f"{charge} - Année 1 (€)",
                min_value=0.0,
                key=f"charge_{charge}_supp_annee1",
                value=charges_fixes_dict["annee1"].get(charge, 0.0)
            )
            charges_fixes_dict["annee1"][charge] = montant1
        with col2:
            montant2 = st.number_input(
                f"{charge} - Année 2 (€)",
                min_value=0.0,
                key=f"charge_{charge}_supp_annee2",
                value=charges_fixes_dict["annee2"].get(charge, 0.0)
            )
            charges_fixes_dict["annee2"][charge] = montant2
        with col3:
            montant3 = st.number_input(
                f"{charge} - Année 3 (€)",
                min_value=0.0,
                key=f"charge_{charge}_supp_annee3",
                value=charges_fixes_dict["annee3"].get(charge, 0.0)
            )
            charges_fixes_dict["annee3"][charge] = montant3
        
        total_annee1 += montant1
        total_annee2 += montant2
        total_annee3 += montant3
    
    data["total_charges_fixes_annee1"] = total_annee1
    data["total_charges_fixes_annee2"] = total_annee2
    data["total_charges_fixes_annee3"] = total_annee3
    
    st.write("---")
    st.markdown(f"**Total Charges Fixes Année 1 :** {total_annee1:.2f} €")
    st.markdown(f"**Total Charges Fixes Année 2 :** {total_annee2:.2f} €")
    st.markdown(f"**Total Charges Fixes Année 3 :** {total_annee3:.2f} €")
    
    st.session_state["data"] = data

# Section 5 : Chiffre d'Affaires Prévisionnel
def page_chiffre_affaires():
    st.title("Chiffre d'Affaires Prévisionnel")
    
    data = st.session_state["data"]
    type_vente = data["informations_generales"].get("type_vente", "Marchandises")
    
    data["chiffre_affaires"] = data.get("chiffre_affaires", {})
    chiffre_affaires_dict = data["chiffre_affaires"]
    
    def calcul_chiffre_affaires(nom_vente):
        mois = [f"Mois {i}" for i in range(1, 13)]
        data_ca = []
        
        st.subheader(f"Année 1 - {nom_vente}")
        for mois_nom in mois:
            col1, col2, col3 = st.columns(3)
            key_jours = f"{nom_vente}_{mois_nom}_jours"
            key_ca_moyen = f"{nom_vente}_{mois_nom}_ca_moyen"
            with col1:
                jours_travailles = st.number_input(
                    f"{mois_nom} - Nombre de jours travaillés",
                    min_value=0,
                    key=key_jours,
                    value=chiffre_affaires_dict.get(key_jours, 0)
                )
                chiffre_affaires_dict[key_jours] = jours_travailles
            with col2:
                ca_moyen_jour = st.number_input(
                    f"{mois_nom} - Chiffre d'affaires moyen / jour (€)",
                    min_value=0.0,
                    key=key_ca_moyen,
                    value=chiffre_affaires_dict.get(key_ca_moyen, 0.0)
                )
                chiffre_affaires_dict[key_ca_moyen] = ca_moyen_jour
            ca_mensuel = jours_travailles * ca_moyen_jour
            chiffre_affaires_dict[f"{nom_vente}_{mois_nom}_ca"] = ca_mensuel
            data_ca.append({
                "mois": mois_nom,
                "jours_travailles": jours_travailles,
                "ca_moyen_jour": ca_moyen_jour,
                "ca_mensuel": ca_mensuel
            })
            with col3:
                st.write(f"CA mensuel: {ca_mensuel:.2f} €")
        
        df_ca = pd.DataFrame(data_ca)
        total_ca_annee1 = df_ca["ca_mensuel"].sum()
        chiffre_affaires_dict[f"total_ca_{nom_vente}_annee1"] = total_ca_annee1
        
        st.write("---")
        st.markdown(f"**Total Chiffre d'Affaires Année 1 ({nom_vente}) :** {total_ca_annee1:.2f} €")
        
        # Pourcentages d'augmentation
        key_aug_annee2 = f"{nom_vente}_augmentation_annee2"
        key_aug_annee3 = f"{nom_vente}_augmentation_annee3"
        pourcentage_augmentation_annee2 = st.number_input(
            f"Pourcentage d'augmentation du CA entre l'année 1 et l'année 2 (%) ({nom_vente})",
            min_value=0.0,
            key=key_aug_annee2,
            value=chiffre_affaires_dict.get(key_aug_annee2, 0.0)
        )
        chiffre_affaires_dict[key_aug_annee2] = pourcentage_augmentation_annee2
        pourcentage_augmentation_annee3 = st.number_input(
            f"Pourcentage d'augmentation du CA entre l'année 2 et l'année 3 (%) ({nom_vente})",
            min_value=0.0,
            key=key_aug_annee3,
            value=chiffre_affaires_dict.get(key_aug_annee3, 0.0)
        )
        chiffre_affaires_dict[key_aug_annee3] = pourcentage_augmentation_annee3
        
        total_ca_annee2 = total_ca_annee1 * (1 + pourcentage_augmentation_annee2 / 100)
        total_ca_annee3 = total_ca_annee2 * (1 + pourcentage_augmentation_annee3 / 100)
        
        chiffre_affaires_dict[f"total_ca_{nom_vente}_annee2"] = total_ca_annee2
        chiffre_affaires_dict[f"total_ca_{nom_vente}_annee3"] = total_ca_annee3
        
        st.markdown(f"**Total Chiffre d'Affaires Année 2 ({nom_vente}) :** {total_ca_annee2:.2f} €")
        st.markdown(f"**Total Chiffre d'Affaires Année 3 ({nom_vente}) :** {total_ca_annee3:.2f} €")
    
    if type_vente in ["Marchandises", "Mixte"]:
        calcul_chiffre_affaires("Marchandises")
    if type_vente in ["Services", "Mixte"]:
        calcul_chiffre_affaires("Services")
    
    # Calcul du total CA toutes ventes
    total_ca_annee1 = sum(
        chiffre_affaires_dict.get(f"total_ca_{type}_annee1", 0.0) for type in ["Marchandises", "Services"]
    )
    total_ca_annee2 = sum(
        chiffre_affaires_dict.get(f"total_ca_{type}_annee2", 0.0) for type in ["Marchandises", "Services"]
    )
    total_ca_annee3 = sum(
        chiffre_affaires_dict.get(f"total_ca_{type}_annee3", 0.0) for type in ["Marchandises", "Services"]
    )
    
    data["total_chiffre_affaires_annee1"] = total_ca_annee1
    data["total_chiffre_affaires_annee2"] = total_ca_annee2
    data["total_chiffre_affaires_annee3"] = total_ca_annee3
    
    st.write("---")
    st.markdown(f"**Total Chiffre d'Affaires Année 1 (toutes ventes) :** {total_ca_annee1:.2f} €")
    st.markdown(f"**Total Chiffre d'Affaires Année 2 (toutes ventes) :** {total_ca_annee2:.2f} €")
    st.markdown(f"**Total Chiffre d'Affaires Année 3 (toutes ventes) :** {total_ca_annee3:.2f} €")
    
    st.session_state["data"] = data

# Section 6 : Charges Variables
def page_charges_variables():
    st.title("Charges Variables")
    
    data = st.session_state["data"]
    type_vente = data["informations_generales"].get("type_vente", "Marchandises")
    
    if type_vente in ["Marchandises", "Mixte"]:
        st.markdown("""
        ### Vos charges variables
        Les charges variables sont liées au niveau d’activité ou à la production. 
        Il s’agit des achats de marchandises destinées à être revendues, des achats de matières destinées à être transformées, 
        des commissions versées à des agents commerciaux.
        """)
        
        data["charges_variables"] = data.get("charges_variables", {})
        charges_variables = data["charges_variables"]
        
        # Coût d'achat des marchandises en %
        cout_achat_marchandises_pct = st.number_input(
            "Quel est, en % du prix de vente, le coût d'achat de vos marchandises ? (concerne uniquement le chiffre d'affaires vente de marchandises)",
            min_value=0.0,
            max_value=100.0,
            format="%.2f",
            key="cout_achat_marchandises_pct",
            value=charges_variables.get("cout_achat_marchandises_pct", 0.0)
        )
        charges_variables["cout_achat_marchandises_pct"] = cout_achat_marchandises_pct
        
        st.write(f"Coût d'achat des marchandises : {cout_achat_marchandises_pct:.2f}% du prix de vente")
        
        total_ca_marchandises_annee1 = data["chiffre_affaires"].get("total_ca_Marchandises_annee1", 0.0)
        total_charges_variables = total_ca_marchandises_annee1 * cout_achat_marchandises_pct / 100.0
        
        data["total_charges_variables"] = total_charges_variables
        
        st.write(f"Total des Charges Variables Année 1 : {total_charges_variables:.2f} €")
        
    else:
        st.info("Cette section est uniquement applicable si vous vendez des marchandises ou des services mixtes.")
        data["total_charges_variables"] = 0.0
    
    st.session_state["data"] = data

# Section 7 : Fonds de Roulement
def page_fonds_roulement():
    st.title("Votre Besoin en Fonds de Roulement")
    
    data = st.session_state["data"]
    
    st.markdown("""
    ### Déterminez votre besoin en fonds de roulement
    Le fonds de roulement représente le montant nécessaire pour financer le cycle d'exploitation de votre entreprise.
    """)
    
    data["fonds_roulement"] = data.get("fonds_roulement", {})
    fonds_roulement = data["fonds_roulement"]
    
    duree_credits_clients = st.number_input(
        "Durée moyenne des crédits accordés aux clients (en jours) :",
        min_value=0,
        help="Temps moyen qu'un client met pour vous payer.",
        key="duree_credits_clients",
        value=fonds_roulement.get("duree_credits_clients", 0)
    )
    fonds_roulement["duree_credits_clients"] = duree_credits_clients
    
    duree_dettes_fournisseurs = st.number_input(
        "Durée moyenne des dettes fournisseurs (en jours) :",
        min_value=0,
        help="Temps moyen que vous mettez pour payer vos fournisseurs.",
        key="duree_dettes_fournisseurs",
        value=fonds_roulement.get("duree_dettes_fournisseurs", 0)
    )
    fonds_roulement["duree_dettes_fournisseurs"] = duree_dettes_fournisseurs
    
    total_ca_annee1 = data.get("total_chiffre_affaires_annee1", 0.0)
    total_charges_variables = data.get("total_charges_variables", 0.0)
    
    bfr = (total_ca_annee1 * duree_credits_clients / 360) - (total_charges_variables * duree_dettes_fournisseurs / 360)
    fonds_roulement["bfr"] = bfr
    
    st.write("---")
    st.markdown(f"**Durée moyenne des crédits clients :** {duree_credits_clients} jours")
    st.markdown(f"**Durée moyenne des dettes fournisseurs :** {duree_dettes_fournisseurs} jours")
    st.markdown(f"**Besoin en Fonds de Roulement (BFR) Année 1 :** {bfr:.2f} €")
    
    st.session_state["data"] = data

# Section 8 : Salaires
def page_salaires():
    st.title("Salaires Employés et Rémunération Chef d'Entreprise")
    
    data = st.session_state["data"]
    data["salaires"] = data.get("salaires", {"employes": {}, "dirigeants": {}})
    salaires = data["salaires"]
    
    st.markdown("""
    ### Saisissez les salaires et rémunérations pour les 3 années
    Veuillez entrer les chiffres annuels pour les salaires des employés et la rémunération nette des dirigeants.
    """)
    
    st.subheader("Salaires Employés (Net)")
    for annee in range(1, 4):
        key = f"annee{annee}"
        salaires["employes"][key] = st.number_input(
            f"Salaires Employés Année {annee} (€)",
            min_value=0.0,
            key=f"salaires_employes_annee_{annee}",
            value=salaires["employes"].get(key, 0.0)
        )
    
    st.subheader("Rémunération Nette Dirigeant(s)")
    for annee in range(1, 4):
        key = f"annee{annee}"
        salaires["dirigeants"][key] = st.number_input(
            f"Rémunération Dirigeant Année {annee} (€)",
            min_value=0.0,
            key=f"remuneration_dirigeant_annee_{annee}",
            value=salaires["dirigeants"].get(key, 0.0)
        )
    
    st.write("---")
    accre = st.selectbox(
        "Le(s) dirigeant(s) bénéficient-ils de l'ACRE ?",
        options=["Oui", "Non"],
        key="accre",
        index=["Oui", "Non"].index(data.get("accre", "Non")),
        help="Veuillez sélectionner 'Oui' si les dirigeants bénéficient de l'ACRE. Cette question est obligatoire."
    )
    data["accre"] = accre
    
    total_salaires_annee1 = salaires["employes"]["annee1"] + salaires["dirigeants"]["annee1"]
    total_salaires_annee2 = salaires["employes"]["annee2"] + salaires["dirigeants"]["annee2"]
    total_salaires_annee3 = salaires["employes"]["annee3"] + salaires["dirigeants"]["annee3"]
    
    data["total_salaires_annee1"] = total_salaires_annee1
    data["total_salaires_annee2"] = total_salaires_annee2
    data["total_salaires_annee3"] = total_salaires_annee3
    
    st.write("---")
    st.markdown(f"**Total Salaires et Rémunération Année 1 :** {total_salaires_annee1:.2f} €")
    st.markdown(f"**Total Salaires et Rémunération Année 2 :** {total_salaires_annee2:.2f} €")
    st.markdown(f"**Total Salaires et Rémunération Année 3 :** {total_salaires_annee3:.2f} €")
    
    st.session_state["data"] = data

# Section 9 : Contrôle de Rentabilité
def page_rentabilite():
    st.title("Contrôle de Rentabilité")
    
    data = st.session_state["data"]
    
    total_charges_fixes_annee1 = data.get("total_charges_fixes_annee1", 0.0)
    total_charges_variables = data.get("total_charges_variables", 0.0)
    total_chiffre_affaires = data.get("total_chiffre_affaires_annee1", 0.0)
    total_salaires_annee1 = data.get("total_salaires_annee1", 0.0)
    
    if total_chiffre_affaires > 0:
        marge_brute = ((total_chiffre_affaires - total_charges_variables) / total_chiffre_affaires) * 100.0
    else:
        marge_brute = 0.0
    
    charges_fixes_totales = total_charges_fixes_annee1 + total_salaires_annee1
    if marge_brute > 0:
        seuil_rentabilite = charges_fixes_totales / (marge_brute / 100.0)
    else:
        seuil_rentabilite = 0.0
    
    if total_chiffre_affaires >= seuil_rentabilite:
        rentabilite = "Rentable"
        message_rentabilite = "L'entreprise est rentable."
        couleur_rentabilite = "green"
    else:
        rentabilite = "Non rentable"
        message_rentabilite = "L'entreprise n'est pas rentable. Il faut augmenter le chiffre d'affaires ou réduire les charges."
        couleur_rentabilite = "red"
    
    data["marge_brute"] = marge_brute
    data["seuil_rentabilite"] = seuil_rentabilite
    data["rentabilite"] = rentabilite
    
    st.write("---")
    st.markdown(f"**Marge Brute :** {marge_brute:.2f} %")
    st.markdown(f"**Seuil de Rentabilité :** {seuil_rentabilite:.2f} €")
    st.markdown(f"<div style='background-color:{couleur_rentabilite}; color:white; padding:10px; border-radius:5px; text-align:center;'>"
                f"<strong>{rentabilite}</strong> - {message_rentabilite}</div>", unsafe_allow_html=True)
    
    st.session_state["data"] = data

# Section 10 : Trésorerie de Départ
def page_tresorerie():
    st.title("Contrôle du Niveau de votre Trésorerie de Départ")
    data = st.session_state["data"]
    besoins_demarrage=data.get("besoins_demarrage", 0.0)
    tresorerie_depart1 = besoins_demarrage.get("Trésorerie de départ", 0.0)
    
    total_charges_fixes_annee1 = data.get("total_charges_fixes_annee1", 0.0)
    tresorerie_depart = st.number_input(
        "Montant de la trésorerie initiale (€)",
        min_value=0.0,
        key="tresorerie_depart",
        value=data.get("tresorerie_depart", tresorerie_depart1 )
    )
    data["tresorerie_depart"] = tresorerie_depart
    
    seuil_tresorerie = total_charges_fixes_annee1 / 4.0  # 3 mois de charges fixes
    if tresorerie_depart >= seuil_tresorerie:
        niveau_tresorerie = "Adéquate"
        message_tresorerie = "Votre trésorerie de départ est adéquate pour couvrir les charges initiales."
        couleur_tresorerie = "green"
    else:
        niveau_tresorerie = "Trop faible"
        message_tresorerie = "Votre trésorerie de départ est trop faible. Prévoyez plus de trésorerie pour couvrir les charges."
        couleur_tresorerie = "red"
    
    data["niveau_tresorerie"] = niveau_tresorerie
    
    st.write("---")
    st.markdown(f"### Résultat pour la 1ère année :")
    st.markdown(f"<div style='background-color:{couleur_tresorerie}; color:white; padding:10px; border-radius:5px; text-align:center;'>"
                f"<strong>{niveau_tresorerie}</strong> - {message_tresorerie}</div>", unsafe_allow_html=True)
    
    st.session_state["data"] = data

# Section 11 : Récapitulatif
def page_recapitulatif():
    st.title("Récapitulatif Complet des Données")
    
    data = st.session_state["data"]
    
    st.subheader("1. Informations Générales")
    info = data.get("informations_generales", {})
    st.write(f"Prénom, nom : {info.get('prenom_nom', '')}")
    st.write(f"Intitulé du projet : {info.get('intitule_projet', '')}")
    st.write(f"Statut juridique : {info.get('statut_juridique', '')}")
    st.write(f"Téléphone : {info.get('telephone', '')}")
    st.write(f"Email : {info.get('email', '')}")
    st.write(f"Ville : {info.get('ville', '')}")
    st.write(f"Type de vente : {info.get('type_vente', '')}")
    
    st.subheader("2. Besoins de Démarrage")
    besoins = data.get("besoins_demarrage", {})
    total_besoins = data.get("total_besoins", 0.0)
    for besoin, montant in besoins.items():
        st.write(f"{besoin} : {montant:.2f} €")
    st.write(f"**Total des Besoins de Démarrage : {total_besoins:.2f} €**")
    
    st.title("Récapitulatif des Financements")
    data = st.session_state.get("data", {})
    financements_dict = data.get("financements", {})
    total_financement = data.get("total_financement", 0.0)
    st.subheader("Financements")
    for financement, details in financements_dict.items():
        if isinstance(details, dict):
            montant = details.get("montant", 0.0)
            st.write(f"{details.get('nom', financement)} : {montant:.2f} €")
        else:
            montant = details
            st.write(f"{financement} : {montant:.2f} €")
    
    st.markdown(f"**Total des Financements :** {total_financement:.2f} €")
    
    st.subheader("4. Charges Fixes sur 3 Années")
    charges_fixes_dict = data.get("charges_fixes", {"annee1": {}, "annee2": {}, "annee3": {}})
    total_annee1 = data.get("total_charges_fixes_annee1", 0.0)
    total_annee2 = data.get("total_charges_fixes_annee2", 0.0)
    total_annee3 = data.get("total_charges_fixes_annee3", 0.0)
    charges_supp = data.get("charges_supplementaires", [])
    
    for charge in charges_fixes_dict["annee1"].keys():
        montant1 = charges_fixes_dict["annee1"].get(charge, 0.0)
        montant2 = charges_fixes_dict["annee2"].get(charge, 0.0)
        montant3 = charges_fixes_dict["annee3"].get(charge, 0.0)
        st.write(f"{charge} - Année 1 : {montant1:.2f} €, Année 2 : {montant2:.2f} €, Année 3 : {montant3:.2f} €")
    
    st.write(f"**Total Charges Fixes Année 1 : {total_annee1:.2f} €**")
    st.write(f"**Total Charges Fixes Année 2 : {total_annee2:.2f} €**")
    st.write(f"**Total Charges Fixes Année 3 : {total_annee3:.2f} €**")
    
    st.subheader("5. Chiffre d'Affaires Prévisionnel")
    total_ca_annee1 = data.get("total_chiffre_affaires_annee1", 0.0)
    total_ca_annee2 = data.get("total_chiffre_affaires_annee2", 0.0)
    total_ca_annee3 = data.get("total_chiffre_affaires_annee3", 0.0)
    
    st.write(f"Total Chiffre d'Affaires Année 1 : {total_ca_annee1:.2f} €")
    st.write(f"Total Chiffre d'Affaires Année 2 : {total_ca_annee2:.2f} €")
    st.write(f"Total Chiffre d'Affaires Année 3 : {total_ca_annee3:.2f} €")
    
    st.subheader("6. Charges Variables")
    cout_achat_marchandises_pct = data.get("charges_variables", {}).get("cout_achat_marchandises_pct", 0.0)
    total_charges_variables = data.get("total_charges_variables", 0.0)
    st.write(f"Coût d'achat des marchandises (% du CA) : {cout_achat_marchandises_pct:.2f} %")
    st.write(f"Total Charges Variables Année 1 : {total_charges_variables:.2f} €")
    
    st.subheader("7. Fonds de Roulement")
    fonds_roulement = data.get("fonds_roulement", {})
    duree_credits_clients = fonds_roulement.get("duree_credits_clients", 0)
    duree_dettes_fournisseurs = fonds_roulement.get("duree_dettes_fournisseurs", 0)
    bfr = fonds_roulement.get("bfr", 0.0)
    st.write(f"Durée moyenne des crédits clients : {duree_credits_clients} jours")
    st.write(f"Durée moyenne des dettes fournisseurs : {duree_dettes_fournisseurs} jours")
    st.write(f"Besoin en Fonds de Roulement (BFR) Année 1 : {bfr:.2f} €")
    
    st.subheader("8. Salaires et Rémunération")
    salaires = data.get("salaires", {})
    for annee in range(1, 4):
        key = f"annee{annee}"
        salaires_employes = salaires.get("employes", {}).get(key, 0.0)
        remuneration_dirigeants = salaires.get("dirigeants", {}).get(key, 0.0)
        st.write(f"Année {annee} : Salaires employés : {salaires_employes:.2f} €, Rémunération dirigeants : {remuneration_dirigeants:.2f} €")
        st.write(f"Total Salaires Année {annee} : {(salaires_employes + remuneration_dirigeants):.2f} €")
    
    st.subheader("9. Rentabilité")
    marge_brute = data.get("marge_brute", 0.0)
    seuil_rentabilite = data.get("seuil_rentabilite", 0.0)
    rentabilite = data.get("rentabilite", "Non rentable")
    st.write(f"Marge Brute : {marge_brute:.2f} %")
    st.write(f"Seuil de Rentabilité : {seuil_rentabilite:.2f} €")
    st.write(f"Rentabilité : {rentabilite}")
    
    st.subheader("10. Trésorerie de Départ")
    tresorerie_depart = data.get("tresorerie_depart", 0.0)
    niveau_tresorerie = data.get("niveau_tresorerie", "Trop faible")
    st.write(f"Montant de la Trésorerie Initiale : {tresorerie_depart:.2f} €")
    st.write(f"Niveau de Trésorerie : {niveau_tresorerie}")
    
    st.session_state["data"] = data
    
    
    
    
import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def page_investissements_et_financements(): 
    st.title("Investissements et Financements")
    
    # Initialiser la clé 'export_data' dans session_state si elle n'existe pas
    if 'export_data' not in st.session_state:
        st.session_state['export_data'] = {}
    
    # Récupérer les données de la session
    data = st.session_state.get("data", {})
    
    # Récupérer les informations du projet
    projet = data.get("informations_generales", {}).get("intitule_projet", "N/A")
    porteur_projet = data.get("informations_generales", {}).get("prenom_nom", "N/A")
    
    st.write(f"**Projet :** {projet}")
    st.write(f"**Porteur de projet :** {porteur_projet}")
    
    st.write("---")
    
    # Initialiser une liste pour stocker toutes les lignes du tableau
    table_data = []
    
    # Immobilisations Incorporelles
    immobilisations_incorporelles = {
        "Frais d’établissement": data.get("besoins_demarrage", {}).get("Frais d’établissement", 0.0),
        "Frais d’ouverture de compteurs": data.get("besoins_demarrage", {}).get("Frais d’ouverture de compteurs", 0.0),
        "Logiciels, formations": data.get("besoins_demarrage", {}).get("Logiciels, formations", 0.0),
        "Dépôt de marque": data.get("besoins_demarrage", {}).get("Dépôt de marque", 0.0),
        "Droits d’entrée": data.get("besoins_demarrage", {}).get("Droits d’entrée", 0.0),
        "Achat fonds de commerce ou parts": data.get("besoins_demarrage", {}).get("Achat fonds de commerce ou parts", 0.0),
        "Droit au bail": data.get("besoins_demarrage", {}).get("Droit au bail", 0.0),
        "Caution ou dépôt de garantie": data.get("besoins_demarrage", {}).get("Caution ou dépôt de garantie", 0.0),
        "Frais de dossier": data.get("besoins_demarrage", {}).get("Frais de dossier", 0.0),
        "Frais de notaire": data.get("besoins_demarrage", {}).get("Frais de notaire", 0.0),
    }
    total_incorporelles = sum(immobilisations_incorporelles.values())
    table_data.append({
        "Investissements": "Immobilisations incorporelles",
        "Taux (%)": "",
        "Durée (mois)": "",
        "Montant (€)": f"{total_incorporelles:.2f}"
    })
    for desc, montant in immobilisations_incorporelles.items():
        table_data.append({
            "Investissements": desc,
            "Taux (%)": "",
            "Durée (mois)": "",
            "Montant (€)": f"{montant:.2f}"
        })
    
    # Immobilisations Corporelles
    immobilisations_corporelles = {
        "Enseigne et éléments de communication": data.get("besoins_demarrage", {}).get("Enseigne et éléments de communication", 0.0),
        "Véhicule": data.get("besoins_demarrage", {}).get("Véhicule", 0.0),
        "Matériel professionnel": data.get("besoins_demarrage", {}).get("Matériel professionnel", 0.0),
        "Matériel autre": data.get("besoins_demarrage", {}).get("Matériel autre", 0.0),
        "Matériel de bureau": data.get("besoins_demarrage", {}).get("Matériel de bureau", 0.0),
    }
    total_corporelles = sum(immobilisations_corporelles.values())
    table_data.append({
        "Investissements": "Immobilisations corporelles",
        "Taux (%)": "",
        "Durée (mois)": "",
        "Montant (€)": f"{total_corporelles:.2f}"
    })
    for desc, montant in immobilisations_corporelles.items():
        table_data.append({
            "Investissements": desc,
            "Taux (%)": "",
            "Durée (mois)": "",
            "Montant (€)": f"{montant:.2f}"
        })
    
    # Autres Investissements
    autres_investissements = {
        "Stock de matières et produits": data.get("besoins_demarrage", {}).get("Stock de matières et produits", 0.0),
        "Trésorerie de départ": data.get("besoins_demarrage", {}).get("Trésorerie de départ", 0.0)
    }
    total_autres = sum(autres_investissements.values())
    table_data.append({
        "Investissements": "Autres investissements",
        "Taux (%)": "",
        "Durée (mois)": "",
        "Montant (€)": f"{total_autres:.2f}"
    })
    for desc, montant in autres_investissements.items():
        table_data.append({
            "Investissements": desc,
            "Taux (%)": "",
            "Durée (mois)": "",
            "Montant (€)": f"{montant:.2f}"
        })
    
    # TOTAL BESOINS
    total_besoins = total_incorporelles + total_corporelles + total_autres
    table_data.append({
        "Investissements": "TOTAL BESOINS",
        "Taux (%)": "",
        "Durée (mois)": "",
        "Montant (€)": f"{total_besoins:.2f}"
    })
    
    # Section FINANCEMENT DES INVESTISSEMENTS
    table_data.append({
        "Investissements": "FINANCEMENT DES INVESTISSEMENTS",
        "Taux (%)": "",
        "Durée (mois)": "",
        "Montant (€)": ""
    })
    table_data.append({
        "Investissements": "Montant € hors taxes",
        "Taux (%)": "",
        "Durée (mois)": "",
        "Montant (€)": ""
    })
    
    # Apport Personnel
    financements = data.get("financements", {})
    apport_personnel = {
        "Apport personnel ou familial": financements.get("Apport personnel ou familial", 0.0),
        "Apports en nature (en valeur)": financements.get("Apports en nature (en valeur)", 0.0),
    }
    total_apport_personnel = sum(apport_personnel.values())
    table_data.append({
        "Investissements": "Apport personnel",
        "Taux (%)": "",
        "Durée (mois)": "",
        "Montant (€)": f"{total_apport_personnel:.2f}"
    })
    for desc, montant in apport_personnel.items():
        table_data.append({
            "Investissements": desc,
            "Taux (%)": "",
            "Durée (mois)": "",
            "Montant (€)": f"{montant:.2f}"
        })
    
    # Emprunts Dynamiques
    emprunts_keys = ["Prêt 1", "Prêt 2", "Prêt 3"]
    emprunts_list = []
    total_emprunts = 0.0

    # Collecter les détails des emprunts
    for i, key in enumerate(emprunts_keys, start=1):
        pret = financements.get(key, {})
        nom = pret.get("nom", "")
        taux = pret.get("taux", 0.0)
        duree = pret.get("duree", 0)
        montant = pret.get("montant", 0.0)
        
        # Définir le nom de l'emprunt
        emprunt_nom = nom if nom else f"Prêt {i}"
        
        # Ajouter les détails du prêt
        if montant > 0:
            emprunts_list.append({
                "Investissements": emprunt_nom,
                "Taux (%)": f"{taux:.2f}%",
                "Durée (mois)": duree,
                "Montant (€)": f"{montant:.2f}"
            })
            total_emprunts += montant
        else:
            emprunts_list.append({
                "Investissements": emprunt_nom,
                "Taux (%)": "-",
                "Durée (mois)": "-",
                "Montant (€)": "0.00"
            })

    # TOTAL EMPRUNTS - placé avant les emprunts individuels
    table_data.append({
        "Investissements": "TOTAL EMPRUNTS",
        "Taux (%)": "",
        "Durée (mois)": "",
        "Montant (€)": f"{total_emprunts:.2f}"
    })

    # Ajouter les emprunts individuels après le total
    for emprunt in emprunts_list:
        table_data.append(emprunt)

     # Subventions Dynamiques
    subventions_keys = ["Subvention 1", "Subvention 2"]
    subventions_list = []
    total_subventions = 0.0
    
    # Calculer le total des subventions d'abord
    for i, key in enumerate(subventions_keys, start=1):
        subv = financements.get(key, {})
        nom = subv.get("nom", "")
        montant = subv.get("montant", 0.0)
        
        # Définir le nom de la subvention
        subvention_nom = nom if nom else f"Subvention {i}"
        
        # Ajouter les détails de la subvention
        if montant > 0:
            subventions_list.append({
                "Investissements": subvention_nom,
                "Taux (%)": "",
                "Durée (mois)": "",
                "Montant (€)": f"{montant:.2f}"
            })
            total_subventions += montant
        else:
            subventions_list.append({
                "Investissements": subvention_nom,
                "Taux (%)": "",
                "Durée (mois)": "",
                "Montant (€)": "0.00"
            })
    
    # TOTAL SUBVENTIONS - placé avant les subventions individuelles
    table_data.append({
        "Investissements": "TOTAL SUBVENTIONS",
        "Taux (%)": "",
        "Durée (mois)": "",
        "Montant (€)": f"{total_subventions:.2f}"
    })
    
    # Ajouter les subventions individuelles après le total
    for subv in subventions_list:
        table_data.append(subv)
    
    # Autre Financement
    autre_financement = financements.get("Autre financement", 0.0)
    table_data.append({
        "Investissements": "Autre financement",
        "Taux (%)": "",
        "Durée (mois)": "",
        "Montant (€)": f"{autre_financement:.2f}"
    })
    
    # TOTAL RESSOURCES
    total_ressources = total_apport_personnel + total_emprunts + total_subventions + autre_financement
    table_data.append({
        "Investissements": "TOTAL RESSOURCES",
        "Taux (%)": "",
        "Durée (mois)": "",
        "Montant (€)": f"{total_ressources:.2f}"
    })
    
    # Vérification de l'équilibre
    if total_ressources == total_besoins:
        equilibrium_message = "Le total des ressources couvre exactement les besoins."
        equilibrium_type = "success"
    elif total_ressources > total_besoins:
        surplus = total_ressources - total_besoins
        equilibrium_message = f"Les ressources dépassent les besoins de {surplus:.2f} €."
        equilibrium_type = "info"
    else:
        deficit = total_besoins - total_ressources
        equilibrium_message = f"Il manque {deficit:.2f} € pour couvrir les besoins."
        equilibrium_type = "warning"
    
    if equilibrium_type == "success":
        st.success(equilibrium_message)
    elif equilibrium_type == "info":
        st.info(equilibrium_message)
    else:
        st.warning(equilibrium_message)
    
    st.write("---")
    
    # Créer le DataFrame unique avec les quatre colonnes
    df_unique = pd.DataFrame(table_data, columns=["Investissements", "Taux (%)", "Durée (mois)", "Montant (€)"])
    
    # Afficher le tableau dans Streamlit
    st.dataframe(df_unique.style.apply(lambda x: ['background-color: #f0f0f0' if pd.isna(v) else '' for v in x], axis=1))
    
    # Stocker les totaux dans la session
    data["total_investissements"] = total_besoins
    data["total_financements"] = total_ressources
    
    st.session_state["data"] = data
    
    # Stocker les données d'exportation dans la nouvelle session
    st.session_state['export_data_investissements'] = {
        "projet": projet,
        "porteur_projet": porteur_projet,
        "table_data": table_data,
        "equilibre": {
            "type": equilibrium_type,
            "message": equilibrium_message
        }
    }
    
    # Section Export
    st.header("Exporter les données")
    
    # Bouton pour télécharger le Markdown
    if st.button("Télécharger en Markdown"):
        markdown_content = f"# Investissements et Financements\n\n**Projet :** {projet}\n\n**Porteur de projet :** {porteur_projet}\n\n"
        
        # Convertir le DataFrame en Markdown
        markdown_content += df_unique.to_markdown(index=False)
        markdown_content += f"\n\n---\n\n{equilibrium_message}\n"
        
        markdown_bytes = markdown_content.encode('utf-8')
        st.download_button(
            label="Télécharger le Markdown",
            data=markdown_bytes,
            file_name="investissements_et_financements.md",
            mime="text/markdown"
        )
    
    # Bouton pour télécharger le fichier Word
    if st.button("Télécharger en Word"):
        export_data = st.session_state.get('export_data', {})
        if not export_data or "table_data" not in export_data:
            st.error("Aucune donnée disponible pour l'exportation.")
        else:
            doc = Document()
            doc.add_heading('Investissements et Financements', level=1)
            doc.add_paragraph(f"**Projet :** {export_data['projet']}")
            doc.add_paragraph(f"**Porteur de projet :** {export_data['porteur_projet']}")
            doc.add_page_break()
            
            # Créer le DataFrame pour Word
            df_word = pd.DataFrame(export_data['table_data'], columns=["Investissements", "Taux (%)", "Durée (mois)", "Montant (€)"])
            
            # Ajouter le tableau au document Word
            table = doc.add_table(rows=1, cols=len(df_word.columns))
            table.style = 'Light List Accent 1'  # Choisissez un style approprié
            hdr_cells = table.rows[0].cells
            for i, column in enumerate(df_word.columns):
                hdr_cells[i].text = column
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for index, row in df_word.iterrows():
                row_cells = table.add_row().cells
                for i, item in enumerate(row):
                    row_cells[i].text = str(item)
                    # Mettre en gras les catégories principales et les totaux
                    if row["Investissements"] in ["INVESTISSEMENTS", "Montant € hors taxes",
                                                 "Immobilisations incorporelles", "Immobilisations corporelles",
                                                 "Autres investissements", "TOTAL BESOINS",
                                                 "FINANCEMENT DES INVESTISSEMENTS", "Apport personnel",
                                                 "Emprunt", "TOTAL EMPRUNTS", "Subvention",
                                                 "TOTAL SUBVENTIONS", "Autre financement", "TOTAL RESSOURCES"]:
                        run = row_cells[i].paragraphs[0].runs
                        if run:
                            run[0].font.bold = True
            doc.add_paragraph()
            doc.add_paragraph(export_data['equilibre']['message'])
            
            # Enregistrer le document dans un buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            st.download_button(
                label="Télécharger le fichier Word",
                data=buffer,
                file_name="investissements_et_financements.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )



import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT

def page_salaires_charges_sociales():
    st.title("Salaires et Charges Sociales")
    
    # Initialiser la clé 'export_data_salaires_charges_sociales' dans session_state si elle n'existe pas
    if 'export_data_salaires_charges_sociales' not in st.session_state:
        st.session_state['export_data_salaires_charges_sociales'] = {}
    
    data = st.session_state.get("data", {})
    
    # Récupérer les informations du projet
    projet = data.get("informations_generales", {}).get("intitule_projet", "")
    porteur_projet = data.get("informations_generales", {}).get("prenom_nom", "")
    statut_juridique = data.get("informations_generales", {}).get("statut_juridique", "")
    benefice_accre = data.get("accre", "Non")  # Assurez-vous que cette information est bien stockée dans data
    
    # Déterminer le statut social du dirigeant en fonction du statut juridique
    if statut_juridique in ["Entreprise individuelle", "EURL (IS)", "EIRL (IS)", "Micro-entreprise"]:
        statut_social_dirigeant = "Travailleur Non Salarié (TNS)"
    elif statut_juridique in ["SARL (IS)", "SAS (IS)", "SASU (IS)"]:
        statut_social_dirigeant = "Assimilé Salarié"
    else:
        statut_social_dirigeant = "Autre"
    
    st.write(f"**Projet :** {projet}")
    st.write(f"**Porteur de projet :** {porteur_projet}")
    st.write(f"**Statut juridique :** {statut_juridique}")
    st.write(f"**Bénéfice de l'ACRE :** {benefice_accre}")
    st.write(f"**Statut social du (des) dirigeant(s) :** {statut_social_dirigeant}")
    
    st.write("---")
    
    # Récupérer les données de salaires
    salaires = data.get("salaires", {})
    salaires_dirigeant = salaires.get("dirigeants", {})
    salaires_employes = salaires.get("employes", {})
    
    # Définir les taux de charges sociales en fonction du statut juridique et de l'ACCRE
    taux_charges_dirigeant = {
        # Sans ACCRE
        "Sans ACCRE": {
            "Travailleur Non Salarié (TNS)": 0.45,
            "Assimilé Salarié": 0.80,  # Taux approximatif pour les assimilés salariés
        },
        # Avec ACCRE
        "Avec ACCRE": {
            "Travailleur Non Salarié (TNS)": 0.22,
            "Assimilé Salarié": 0.50,  # Taux réduit pour les assimilés salariés avec ACRE
        }
    }
    
    # Sélection du taux approprié pour le dirigeant
    if benefice_accre.lower() == "oui":
        taux_dirigeant = taux_charges_dirigeant["Avec ACCRE"].get(statut_social_dirigeant, 0.45)
    else:
        taux_dirigeant = taux_charges_dirigeant["Sans ACCRE"].get(statut_social_dirigeant, 0.45)
    
    # Taux de charges sociales pour les employés
    taux_charges_employe = 0.72  # Comme indiqué, multiplier par 0.72 qu'il ait ACCRE ou pas
    
    # Préparation des données pour le tableau
    annees = ["Année 1", "Année 2", "Année 3"]
    remuneration_dirigeant = []
    augmentation_dirigeant = []
    charges_sociales_dirigeant = []
    remuneration_employes = []
    augmentation_employes = []
    charges_sociales_employes = []
    
    for i, annee in enumerate(annees):
        annee_key = f"annee{i+1}"
        # Rémunération du (des) dirigeants
        remu_dirigeant = salaires_dirigeant.get(annee_key, 0.0)
        remuneration_dirigeant.append(remu_dirigeant)
        # % augmentation dirigeant
        if i == 0:
            aug_dirigeant = "-"
        else:
            previous_remu_dirigeant = remuneration_dirigeant[i-1]
            if previous_remu_dirigeant != 0:
                aug_dirigeant_value = ((remu_dirigeant - previous_remu_dirigeant) / previous_remu_dirigeant) * 100
                aug_dirigeant = f"{aug_dirigeant_value:.2f}%"
            else:
                aug_dirigeant = "-"
        augmentation_dirigeant.append(aug_dirigeant)
        # Charges sociales du (des) dirigeant(s)
        charge_sociale_dirigeant = remu_dirigeant * taux_dirigeant
        charges_sociales_dirigeant.append(charge_sociale_dirigeant)
        
        # Salaires des employés
        remu_employes = salaires_employes.get(annee_key, 0.0)
        remuneration_employes.append(remu_employes)
        # % augmentation employés
        if i == 0:
            aug_employes = "-"
        else:
            previous_remu_employes = remuneration_employes[i-1]
            if previous_remu_employes != 0:
                aug_employes_value = ((remu_employes - previous_remu_employes) / previous_remu_employes) * 100
                aug_employes = f"{aug_employes_value:.2f}%"
            else:
                aug_employes = "-"
        augmentation_employes.append(aug_employes)
        # Charges sociales employés
        charge_sociale_employes = remu_employes * taux_charges_employe
        charges_sociales_employes.append(charge_sociale_employes)
    
    # Création du DataFrame pour l'affichage
    df = pd.DataFrame({
        "": ["Rémunération du (des) dirigeants", "% augmentation", "Charges sociales du (des) dirigeant(s)",
             "Salaires des employés", "% augmentation", "Charges sociales employés"],
        "Année 1": [f"{remuneration_dirigeant[0]:.2f} €", augmentation_dirigeant[0], f"{charges_sociales_dirigeant[0]:.2f} €",
                    f"{remuneration_employes[0]:.2f} €", augmentation_employes[0], f"{charges_sociales_employes[0]:.2f} €"],
        "Année 2": [f"{remuneration_dirigeant[1]:.2f} €", augmentation_dirigeant[1], f"{charges_sociales_dirigeant[1]:.2f} €",
                    f"{remuneration_employes[1]:.2f} €", augmentation_employes[1], f"{charges_sociales_employes[1]:.2f} €"],
        "Année 3": [f"{remuneration_dirigeant[2]:.2f} €", augmentation_dirigeant[2], f"{charges_sociales_dirigeant[2]:.2f} €",
                    f"{remuneration_employes[2]:.2f} €", augmentation_employes[2], f"{charges_sociales_employes[2]:.2f} €"]
    })
    
    st.table(df)
    
    # Stocker les charges sociales dans les données pour exportation
    data["charges_sociales"] = {
        "dirigeants": {
            "annee1": charges_sociales_dirigeant[0],
            "annee2": charges_sociales_dirigeant[1],
            "annee3": charges_sociales_dirigeant[2]
        },
        "employes": {
            "annee1": charges_sociales_employes[0],
            "annee2": charges_sociales_employes[1],
            "annee3": charges_sociales_employes[2]
        }
    }
    
    st.session_state["data"] = data
    
    # Préparer les données d'exportation pour Salaires et Charges Sociales
    export_table_data = []
    
    # Ajouter les lignes du tableau
    for index, row in df.iterrows():
        export_table_data.append({
            "Description": row[""],
            "Année 1": row["Année 1"],
            "Année 2": row["Année 2"],
            "Année 3": row["Année 3"]
        })
    
    # Stocker les données d'exportation dans la session
    st.session_state['export_data_salaires_charges_sociales'] = {
        "projet": projet,
        "porteur_projet": porteur_projet,
        "statut_juridique": statut_juridique,
        "benefice_accre": benefice_accre,
        "statut_social_dirigeant": statut_social_dirigeant,
        "table_data": export_table_data
    }
    
    # Section Export
    st.header("Exporter les données")
    
    # Bouton pour télécharger le Markdown
    if st.button("Télécharger Salaires en Markdown"):
        export_data = st.session_state.get('export_data_salaires_charges_sociales', {})
        if not export_data or "table_data" not in export_data:
            st.error("Aucune donnée disponible pour l'exportation.")
        else:
            # Construire le contenu Markdown
            markdown_content = f"# Salaires et Charges Sociales\n\n**Projet :** {export_data['projet']}\n\n"
            markdown_content += f"**Porteur de projet :** {export_data['porteur_projet']}\n\n"
            markdown_content += f"**Statut juridique :** {export_data['statut_juridique']}\n\n"
            markdown_content += f"**Bénéfice de l'ACRE :** {export_data['benefice_accre']}\n\n"
            markdown_content += f"**Statut social du (des) dirigeant(s) :** {export_data['statut_social_dirigeant']}\n\n"
            markdown_content += "---\n\n"
            
            # Créer un DataFrame pour Markdown
            df_markdown = pd.DataFrame(export_data['table_data'])
            markdown_content += df_markdown.to_markdown(index=False)
            
            markdown_content += f"\n\n---\n\n"
            
            markdown_bytes = markdown_content.encode('utf-8')
            st.download_button(
                label="Télécharger le Markdown",
                data=markdown_bytes,
                file_name="salaires_charges_sociales.md",
                mime="text/markdown"
            )
    
    # Bouton pour télécharger le fichier Word
    if st.button("Télécharger Salaires en Word"):
        export_data_salaires = st.session_state.get('export_data_salaires_charges_sociales', {})
        export_data_investissements = st.session_state.get('export_data_investissements', {})
        
        if not export_data_salaires or "table_data" not in export_data_salaires:
            st.error("Aucune donnée disponible pour l'exportation des Salaires et Charges Sociales.")
            return
        
        if not export_data_investissements or "table_data" not in export_data_investissements:
            st.error("Aucune donnée disponible pour l'exportation des Investissements et Financements.")
            return
        
        doc = Document()
        
        # Ajouter la première table : Investissements et Financements
        doc.add_heading('Investissements et Financements', level=1)
        doc.add_paragraph(f"**Projet :** {export_data_investissements['projet']}")
        doc.add_paragraph(f"**Porteur de projet :** {export_data_investissements['porteur_projet']}")
        doc.add_paragraph(f"**Equilibre :** {export_data_investissements['equilibre']['message']}")
        doc.add_page_break()
        
        # Créer le tableau Investissements et Financements dans Word
        table_word_inv = doc.add_table(rows=1, cols=4)
        table_word_inv.style = 'Light List Accent 1'
        table_word_inv.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells_inv = table_word_inv.rows[0].cells
        headers_inv = ["Investissements", "Taux (%)", "Durée (mois)", "Montant (€)"]
        for i, header in enumerate(headers_inv):
            hdr_cells_inv[i].text = header
            # Mettre en gras les en-têtes
            for paragraph in hdr_cells_inv[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            hdr_cells_inv[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Ajouter les données Investissements et Financements au tableau
        for row in export_data_investissements['table_data']:
            row_cells_inv = table_word_inv.add_row().cells
            row_cells_inv[0].text = row["Investissements"]
            row_cells_inv[1].text = row["Taux (%)"]
            row_cells_inv[2].text = str(row["Durée (mois)"]) if row["Durée (mois)"] != "-" else "-"
            row_cells_inv[3].text = row["Montant (€)"]
            
            # Mise en forme des lignes spécifiques
            if row["Investissements"] in ["INVESTISSEMENTS", "FINANCEMENT DES INVESTISSEMENTS", "TOTAL SUBVENTIONS", "TOTAL EMPRUNTS"]:
                for cell in row_cells_inv:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
            elif "TOTAL" in row["Investissements"]:
                for cell in row_cells_inv:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
            else:
                pass  # Aucune mise en forme supplémentaire
            
            # Alignement des cellules
            row_cells_inv[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_cells_inv[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_cells_inv[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        # Ajouter la deuxième table : Salaires et Charges Sociales
        doc.add_heading('Salaires et Charges Sociales', level=1)
        doc.add_paragraph(f"**Projet :** {export_data_salaires['projet']}")
        doc.add_paragraph(f"**Porteur de projet :** {export_data_salaires['porteur_projet']}")
        doc.add_paragraph(f"**Statut juridique :** {export_data_salaires['statut_juridique']}")
        doc.add_paragraph(f"**Bénéfice de l'ACRE :** {export_data_salaires['benefice_accre']}")
        doc.add_paragraph(f"**Statut social du (des) dirigeant(s) :** {export_data_salaires['statut_social_dirigeant']}")
        doc.add_paragraph("---")
        
        # Créer le tableau Salaires et Charges Sociales dans Word
        table_word_sal = doc.add_table(rows=1, cols=4)
        table_word_sal.style = 'Light List Accent 1'
        table_word_sal.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells_sal = table_word_sal.rows[0].cells
        headers_sal = ["Description", "Année 1", "Année 2", "Année 3"]
        for i, header in enumerate(headers_sal):
            hdr_cells_sal[i].text = header
            # Mettre en gras les en-têtes
            for paragraph in hdr_cells_sal[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            hdr_cells_sal[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Ajouter les données Salaires et Charges Sociales au tableau
        for row in export_data_salaires['table_data']:
            row_cells_sal = table_word_sal.add_row().cells
            row_cells_sal[0].text = row["Description"]
            row_cells_sal[1].text = row["Année 1"]
            row_cells_sal[2].text = row["Année 2"]
            row_cells_sal[3].text = row["Année 3"]
            
            # Mise en forme des lignes spécifiques
            # Vous pouvez ajouter des conditions ici pour mettre en forme certaines lignes si nécessaire
            
            # Alignement des cellules
            row_cells_sal[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row_cells_sal[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row_cells_sal[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        # Ajouter des informations supplémentaires si nécessaire
        doc.add_paragraph()
        doc.add_paragraph("Les charges sociales sont calculées en fonction des taux applicables.")
        
        # Enregistrer le document dans un buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        st.download_button(
            label="Télécharger le fichier Word Complet",
            data=buffer,
            file_name="investissements_et_salaires_charges_sociales.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )


def page_detail_amortissements():
    st.title("Détail des Amortissements")
    
    # Initialiser la clé 'export_data_detail_amortissements' dans session_state si elle n'existe pas
    if 'export_data_detail_amortissements' not in st.session_state:
        st.session_state['export_data_detail_amortissements'] = {}
    
    data = st.session_state.get("data", {})
    
    st.write("---")
    
    # Récupérer la durée d'amortissement
    duree_amortissement = data.get("duree_amortissement", 0)
    if duree_amortissement <= 0:
        st.warning("La durée d'amortissement doit être supérieure à zéro.")
        return
    
    # Récupérer les montants des investissements
    besoins_demarrage = data.get("besoins_demarrage", {})
    
    # Incorporels
    incorporels_items = [
        "Frais d’établissement",
        "Logiciels, formations",
        "Droits d’entrée",
        "Frais de dossier",
        "Frais de notaire"
    ]
    incorporels_amortissements = {}
    total_incorporels_amort = [0.0, 0.0, 0.0]
    
    for item in incorporels_items:
        amount = besoins_demarrage.get(item, 0.0)
        annual_depreciation = amount / duree_amortissement if duree_amortissement > 0 else 0.0
        amortization_years = [0.0, 0.0, 0.0]
        for year in range(3):
            if year < duree_amortissement:
                amortization_years[year] = annual_depreciation
                total_incorporels_amort[year] += annual_depreciation
        incorporels_amortissements[item] = amortization_years
    
    # Corporels
    corporels_items = [
        "Enseigne et éléments de communication",
        "Véhicule",
        "Matériel professionnel",
        "Matériel autre",
        "Matériel de bureau"
    ]
    corporels_amortissements = {}
    total_corporels_amort = [0.0, 0.0, 0.0]
    
    for item in corporels_items:
        amount = besoins_demarrage.get(item, 0.0)
        annual_depreciation = amount / duree_amortissement if duree_amortissement > 0 else 0.0
        amortization_years = [0.0, 0.0, 0.0]
        for year in range(3):
            if year < duree_amortissement:
                amortization_years[year] = annual_depreciation
                total_corporels_amort[year] += annual_depreciation
        corporels_amortissements[item] = amortization_years
    
    # Total amortissements par année
    total_amortissements = [
        total_incorporels_amort[year] + total_corporels_amort[year] for year in range(3)
    ]
    
    # Création d'un tableau unique
    st.subheader("Amortissements")
    amortissements_data = []
    for item in incorporels_items:
        amortization_years = incorporels_amortissements.get(item, [0.0, 0.0, 0.0])
        amortissements_data.append({
            "Amortissement": item,
            "Année 1": f"{amortization_years[0]:.2f}",
            "Année 2": f"{amortization_years[1]:.2f}",
            "Année 3": f"{amortization_years[2]:.2f}"
        })
    for item in corporels_items:
        amortization_years = corporels_amortissements.get(item, [0.0, 0.0, 0.0])
        amortissements_data.append({
            "Amortissement": item,
            "Année 1": f"{amortization_years[0]:.2f}",
            "Année 2": f"{amortization_years[1]:.2f}",
            "Année 3": f"{amortization_years[2]:.2f}"
        })
    # Total amortissements
    amortissements_data.append({
        "Amortissement": "Total Amortissements",
        "Année 1": f"{total_amortissements[0]:.2f}",
        "Année 2": f"{total_amortissements[1]:.2f}",
        "Année 3": f"{total_amortissements[2]:.2f}"
    })
    df_amortissements = pd.DataFrame(amortissements_data)
    st.table(df_amortissements)
    
    # Stocker les amortissements dans les données pour exportation
    data["amortissements"] = {
        "incorporels": {
            "annee1": total_incorporels_amort[0],
            "annee2": total_incorporels_amort[1],
            "annee3": total_incorporels_amort[2]
        },
        "corporels": {
            "annee1": total_corporels_amort[0],
            "annee2": total_corporels_amort[1],
            "annee3": total_corporels_amort[2]
        },
        "total": {
            "annee1": total_amortissements[0],
            "annee2": total_amortissements[1],
            "annee3": total_amortissements[2]
        }
    }
    
    st.session_state["data"] = data
    
    # Préparer les données d'exportation pour Détail des Amortissements
    export_table_amortissements = []
    for row in amortissements_data:
        export_table_amortissements.append({
            "Amortissement": row["Amortissement"],
            "Année 1": row["Année 1"],
            "Année 2": row["Année 2"],
            "Année 3": row["Année 3"]
        })
    
    # Stocker les données d'exportation dans la session
    st.session_state['export_data_detail_amortissements'] = {
        "amortissements": export_table_amortissements
    }
    
    # Section Export
    st.header("Exporter les données")
    
    # Bouton pour télécharger le Markdown
    if st.button("Télécharger Amortissements en Markdown"):
        export_data = st.session_state.get('export_data_detail_amortissements', {})
        if not export_data:
            st.error("Aucune donnée disponible pour l'exportation.")
        else:
            # Construire le contenu Markdown
            markdown_content = f"# Détail des Amortissements\n\n"
            markdown_content += "---\n\n"
            
            # Amortissements
            markdown_content += "## Amortissements\n\n"
            df_amortissements_md = pd.DataFrame(export_data['amortissements'])
            markdown_content += df_amortissements_md.to_markdown(index=False)
            markdown_content += "\n\n"
            
            markdown_bytes = markdown_content.encode('utf-8')
            st.download_button(
                label="Télécharger le Markdown",
                data=markdown_bytes,
                file_name="detail_amortissements.md",
                mime="text/markdown"
            )
    
    # Bouton pour télécharger le fichier Word
    if st.button("Télécharger Amortissements en Word"):
        export_data_amortissements = st.session_state.get('export_data_detail_amortissements', {})
        export_data_investissements = st.session_state.get('export_data_investissements', {})
        export_data_salaires = st.session_state.get('export_data_salaires_charges_sociales', {})
        
        if not export_data_amortissements or "amortissements" not in export_data_amortissements:
            st.error("Aucune donnée disponible pour l'exportation des Amortissements.")
            return
        
        if not export_data_investissements or "table_data" not in export_data_investissements:
            st.error("Aucune donnée disponible pour l'exportation des Investissements et Financements.")
            return
        
        if not export_data_salaires or "table_data" not in export_data_salaires:
            st.error("Aucune donnée disponible pour l'exportation des Salaires et Charges Sociales.")
            return
        
        doc = Document()
        
        # Ajouter la première table : Investissements et Financements
        doc.add_heading('Investissements et Financements', level=1)
        doc.add_paragraph(f"**Projet :** {export_data_investissements['projet']}")
        doc.add_paragraph(f"**Porteur de projet :** {export_data_investissements['porteur_projet']}")
        doc.add_paragraph(f"**Equilibre :** {export_data_investissements['equilibre']['message']}")
        doc.add_page_break()
        
        # Créer le tableau Investissements et Financements dans Word
        table_word_inv = doc.add_table(rows=1, cols=4)
        table_word_inv.style = 'Light List Accent 1'
        table_word_inv.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells_inv = table_word_inv.rows[0].cells
        headers_inv = ["Investissements", "Taux (%)", "Durée (mois)", "Montant (€)"]
        for i, header in enumerate(headers_inv):
            hdr_cells_inv[i].text = header
            # Mettre en gras les en-têtes
            for paragraph in hdr_cells_inv[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            hdr_cells_inv[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Ajouter les données Investissements et Financements au tableau
        for row in export_data_investissements['table_data']:
            row_cells_inv = table_word_inv.add_row().cells
            row_cells_inv[0].text = row["Investissements"]
            row_cells_inv[1].text = row["Taux (%)"]
            row_cells_inv[2].text = str(row["Durée (mois)"]) if row["Durée (mois)"] != "-" else "-"
            row_cells_inv[3].text = row["Montant (€)"]
            
            # Mise en forme des lignes spécifiques
            if row["Investissements"] in ["INVESTISSEMENTS", "FINANCEMENT DES INVESTISSEMENTS", "TOTAL SUBVENTIONS", "TOTAL EMPRUNTS"]:
                for cell in row_cells_inv:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
            elif "TOTAL" in row["Investissements"]:
                for cell in row_cells_inv:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
            else:
                pass  # Aucune mise en forme supplémentaire
            
            # Alignement des cellules
            row_cells_inv[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_cells_inv[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_cells_inv[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        # Ajouter la deuxième table : Salaires et Charges Sociales
        doc.add_heading('Salaires et Charges Sociales', level=1)
        doc.add_paragraph(f"**Projet :** {export_data_salaires['projet']}")
        doc.add_paragraph(f"**Porteur de projet :** {export_data_salaires['porteur_projet']}")
        doc.add_paragraph(f"**Statut juridique :** {export_data_salaires['statut_juridique']}")
        doc.add_paragraph(f"**Bénéfice de l'ACRE :** {export_data_salaires['benefice_accre']}")
        doc.add_paragraph(f"**Statut social du (des) dirigeant(s) :** {export_data_salaires['statut_social_dirigeant']}")
        doc.add_paragraph("---")
        
        # Créer le tableau Salaires et Charges Sociales dans Word
        table_word_sal = doc.add_table(rows=1, cols=4)
        table_word_sal.style = 'Light List Accent 1'
        table_word_sal.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells_sal = table_word_sal.rows[0].cells
        headers_sal = ["Description", "Année 1", "Année 2", "Année 3"]
        for i, header in enumerate(headers_sal):
            hdr_cells_sal[i].text = header
            # Mettre en gras les en-têtes
            for paragraph in hdr_cells_sal[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            hdr_cells_sal[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Ajouter les données Salaires et Charges Sociales au tableau
        for row in export_data_salaires['table_data']:
            row_cells_sal = table_word_sal.add_row().cells
            row_cells_sal[0].text = row["Description"]
            row_cells_sal[1].text = row["Année 1"]
            row_cells_sal[2].text = row["Année 2"]
            row_cells_sal[3].text = row["Année 3"]
            
            # Mise en forme des lignes spécifiques
            # Vous pouvez ajouter des conditions ici pour mettre en forme certaines lignes si nécessaire
            
            # Alignement des cellules
            row_cells_sal[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row_cells_sal[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row_cells_sal[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        # Ajouter la troisième table : Détail des Amortissements
        doc.add_heading('Détail des Amortissements', level=1)
        
        export_data_amortissements = st.session_state.get('export_data_detail_amortissements', {})
        
        # Vérifier si les données d'amortissements sont disponibles
        if not export_data_amortissements or "amortissements" not in export_data_amortissements:
            st.error("Aucune donnée disponible pour l'exportation des Amortissements.")
            return
        
        # Créer le tableau Amortissements dans Word
        doc.add_heading('Amortissements', level=2)
        table_word_amort = doc.add_table(rows=1, cols=4)
        table_word_amort.style = 'Light List Accent 1'
        table_word_amort.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells_amort = table_word_amort.rows[0].cells
        headers_amort = ["Amortissement", "Année 1", "Année 2", "Année 3"]
        for i, header in enumerate(headers_amort):
            hdr_cells_amort[i].text = header
            # Mettre en gras les en-têtes
            for paragraph in hdr_cells_amort[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            hdr_cells_amort[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Ajouter les données Amortissements au tableau
        for row in export_data_amortissements['amortissements']:
            row_cells_amort = table_word_amort.add_row().cells
            row_cells_amort[0].text = row["Amortissement"]
            row_cells_amort[1].text = row["Année 1"]
            row_cells_amort[2].text = row["Année 2"]
            row_cells_amort[3].text = row["Année 3"]
        
        # Ajouter des informations supplémentaires si nécessaire
        doc.add_paragraph()
        doc.add_paragraph("Les amortissements sont calculés en fonction de la durée d'amortissement spécifiée.")
        
        # Enregistrer le document dans un buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        st.download_button(
            label="Télécharger le fichier Word Complet",
            data=buffer,
            file_name="document_complet_financier.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT

def telecharger_document_complet():
    export_data_investissements = st.session_state.get('export_data_investissements', {})
    export_data_salaires = st.session_state.get('export_data_salaires_charges_sociales', {})
    export_data_amortissements = st.session_state.get('export_data_detail_amortissements', {})
    export_data_compte = st.session_state.get('export_data_compte_resultats_previsionnel', {})
    
    # Vérifiez que toutes les données sont présentes
    if not all([export_data_investissements.get("table_data"),
                export_data_salaires.get("table_data"),
                export_data_amortissements.get("amortissements"),
                export_data_compte.get("table_data")]):
        st.error("Toutes les sections doivent être remplies avant de télécharger le document complet.")
        return
    
    # Créer un document Word
    doc = Document()
    
    ### 1. Ajouter la section Investissements et Financements ###
    doc.add_heading('Investissements et Financements', level=1)
    doc.add_paragraph(f"**Projet :** {export_data_investissements.get('projet', 'N/A')}")
    doc.add_paragraph(f"**Porteur de projet :** {export_data_investissements.get('porteur_projet', 'N/A')}")
    doc.add_paragraph(f"**Equilibre :** {export_data_investissements.get('equilibre', {}).get('message', '')}")
    doc.add_page_break()
    
    # Créer le tableau Investissements et Financements
    table_inv = doc.add_table(rows=1, cols=4)
    table_inv.style = 'Light List Accent 1'
    table_inv.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_inv = table_inv.rows[0].cells
    headers_inv = ["Investissements", "Taux (%)", "Durée (mois)", "Montant (€)"]
    for i, header in enumerate(headers_inv):
        hdr_cells_inv[i].text = header
        for paragraph in hdr_cells_inv[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_inv[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_investissements['table_data']:
        row_cells = table_inv.add_row().cells
        row_cells[0].text = row.get("Investissements", "")
        row_cells[1].text = row.get("Taux (%)", "")
        row_cells[2].text = str(row.get("Durée (mois)", "")) if row.get("Durée (mois)", "") != "-" else "-"
        row_cells[3].text = row.get("Montant (€)", "")
        
        # Mise en forme des lignes spécifiques
        if row["Investissements"] in ["INVESTISSEMENTS", "FINANCEMENT DES INVESTISSEMENTS", "TOTAL SUBVENTIONS", "TOTAL EMPRUNTS"]:
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        elif "TOTAL" in row["Investissements"]:
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        
        # Alignement des cellules
        row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    ### 2. Ajouter la section Salaires et Charges Sociales ###
    doc.add_heading('Salaires et Charges Sociales', level=1)
    doc.add_paragraph(f"**Projet :** {export_data_salaires.get('projet', 'N/A')}")
    doc.add_paragraph(f"**Porteur de projet :** {export_data_salaires.get('porteur_projet', 'N/A')}")
    doc.add_paragraph(f"**Statut juridique :** {export_data_salaires.get('statut_juridique', 'N/A')}")
    doc.add_paragraph(f"**Bénéfice de l'ACRE :** {export_data_salaires.get('benefice_accre', 'N/A')}")
    doc.add_paragraph(f"**Statut social du (des) dirigeant(s) :** {export_data_salaires.get('statut_social_dirigeant', 'N/A')}")
    doc.add_paragraph("---")
    
    # Créer le tableau Salaires et Charges Sociales
    table_sal = doc.add_table(rows=1, cols=4)
    table_sal.style = 'Light List Accent 1'
    table_sal.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_sal = table_sal.rows[0].cells
    headers_sal = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_sal):
        hdr_cells_sal[i].text = header
        for paragraph in hdr_cells_sal[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_sal[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_salaires['table_data']:
        row_cells = table_sal.add_row().cells
        row_cells[0].text = row.get("Description", "")
        row_cells[1].text = row.get("Année 1", "")
        row_cells[2].text = row.get("Année 2", "")
        row_cells[3].text = row.get("Année 3", "")
        
        # Alignement des cellules
        row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    ### 3. Ajouter la section Détail des Amortissements ###
    doc.add_heading('Détail des Amortissements', level=1)
    
    # Créer le tableau Détail des Amortissements
    table_amort = doc.add_table(rows=1, cols=4)
    table_amort.style = 'Light List Accent 1'
    table_amort.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_amort = table_amort.rows[0].cells
    headers_amort = ["Amortissement", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_amort):
        hdr_cells_amort[i].text = header
        for paragraph in hdr_cells_amort[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_amort[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_amortissements['amortissements']:
        row_cells = table_amort.add_row().cells
        row_cells[0].text = row.get("Amortissement", "")
        row_cells[1].text = row.get("Année 1", "")
        row_cells[2].text = row.get("Année 2", "")
        row_cells[3].text = row.get("Année 3", "")
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les amortissements sont calculés en fonction de la durée d'amortissement spécifiée.")
    
    ### 4. Ajouter la section Compte de Résultats Prévisionnel ###
    doc.add_heading('Compte de Résultats Prévisionnel', level=1)
    
    # Créer le tableau Compte de Résultats Prévisionnel
    table_compte = doc.add_table(rows=1, cols=4)
    table_compte.style = 'Light List Accent 1'
    table_compte.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_compte = table_compte.rows[0].cells
    headers_compte = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_compte):
        hdr_cells_compte[i].text = header
        for paragraph in hdr_cells_compte[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_compte[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_compte['table_data']:
        row_cells = table_compte.add_row().cells
        row_cells[0].text = row.get("Description", "")
        row_cells[1].text = row.get("Année 1", "")
        row_cells[2].text = row.get("Année 2", "")
        row_cells[3].text = row.get("Année 3", "")
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    # Enregistrer le document dans un buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Bouton de téléchargement
    st.download_button(
        label="Télécharger le fichier Word Complet",
        data=buffer,
        file_name="document_complet_financier.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )



def calculer_impot_societes(resultat_avant_impots):
    """
    Calcule l'Impôt sur les Sociétés (IS) selon la formule progressive.

    Args:
        resultat_avant_impots (float): Résultat avant impôts.

    Returns:
        float: Montant de l'IS.
    """
    if resultat_avant_impots < 0:
        return 0.0
    elif resultat_avant_impots > 38120:
        return 38120 * 0.15 + (resultat_avant_impots - 38120) * 0.28
    else:
        return resultat_avant_impots * 0.15

def page_compte_resultats_previsionnel():
    st.title("Compte de résultats prévisionnel sur 3 ans")
    
    # Initialiser la clé 'export_data_compte_resultats_previsionnel' dans session_state si elle n'existe pas
    if 'export_data_compte_resultats_previsionnel' not in st.session_state:
        st.session_state['export_data_compte_resultats_previsionnel'] = {}
    
    data = st.session_state.get("data", {})
    
    # Récupérer les informations du projet
    projet = data.get("informations_generales", {}).get("intitule_projet", "")
    porteur_projet = data.get("informations_generales", {}).get("prenom_nom", "")
    
    st.write(f"**Projet :** {projet}")
    st.write(f"**Porteur de projet :** {porteur_projet}")
    
    st.write("---")
    
    # Préparation des données
    # Chiffre d'affaires
    ca_marchandises = [
        data["chiffre_affaires"].get("total_ca_Marchandises_annee1", 0.0),
        data["chiffre_affaires"].get("total_ca_Marchandises_annee2", 0.0),
        data["chiffre_affaires"].get("total_ca_Marchandises_annee3", 0.0)
    ]
    ca_services = [
        data["chiffre_affaires"].get("total_ca_Services_annee1", 0.0),
        data["chiffre_affaires"].get("total_ca_Services_annee2", 0.0),
        data["chiffre_affaires"].get("total_ca_Services_annee3", 0.0)
    ]
    total_ca = [
        ca_marchandises[0] + ca_services[0],
        ca_marchandises[1] + ca_services[1],
        ca_marchandises[2] + ca_services[2]
    ]
    
    # Achats consommés (charges variables) - Supposés nuls si pas de marchandises vendues
    charges_variables = [0.0, 0.0, 0.0]
    
    # Marge brute = Total CA - Achats consommés
    marge_brute = [
        total_ca[0] - charges_variables[0],
        total_ca[1] - charges_variables[1],
        total_ca[2] - charges_variables[2]
    ]
    
    # Charges externes (charges fixes)
    charges_fixes_data = data.get("charges_fixes", {})
    charges_fixes_annee1 = charges_fixes_data.get("annee1", {})
    charges_fixes_annee2 = charges_fixes_data.get("annee2", {})
    charges_fixes_annee3 = charges_fixes_data.get("annee3", {})
    
    # Liste des charges externes détaillées
    liste_charges = [
        "Assurances véhicule et RC pro", "Téléphone, internet", "Autres abonnements",
        "Carburant", "Frais de déplacement / hébergement", "Eau, électricité, gaz",
        "Mutuelle", "Fournitures diverses", "Entretien Moto livraison et matériel",
        "Nettoyage des locaux", "Budget publicité et communication", "Emplacements",
        "Expert comptable, avocats", "Markting"
    ]
    
    # Récupération des montants pour chaque charge
    charges_detaillees = {}
    total_charges_fixes = [0.0, 0.0, 0.0]
    for charge in liste_charges:
        montant_annee1 = charges_fixes_annee1.get(charge, 0.0)
        montant_annee2 = charges_fixes_annee2.get(charge, 0.0)
        montant_annee3 = charges_fixes_annee3.get(charge, 0.0)
        charges_detaillees[charge] = [montant_annee1, montant_annee2, montant_annee3]
        total_charges_fixes[0] += montant_annee1
        total_charges_fixes[1] += montant_annee2
        total_charges_fixes[2] += montant_annee3
    
    # Valeur ajoutée = Marge brute - Charges externes
    valeur_ajoutee = [
        marge_brute[0] - total_charges_fixes[0],
        marge_brute[1] - total_charges_fixes[1],
        marge_brute[2] - total_charges_fixes[2]
    ]
    
    # Impôts et taxes (ajouter d'autres impôts si nécessaire)
    impots_et_taxes = [
        charges_fixes_annee1.get("Taxes, CFE", 0.0),
        charges_fixes_annee2.get("Taxes, CFE", 0.0),
        charges_fixes_annee3.get("Taxes, CFE", 0.0)
    ]
    
    # Salaires employés
    salaires_employes = [
        data["salaires"]["employes"].get("annee1", 0.0),
        data["salaires"]["employes"].get("annee2", 0.0),
        data["salaires"]["employes"].get("annee3", 0.0)
    ]
    
    # Charges sociales employés
    charges_sociales_employes = [
        data["charges_sociales"]["employes"].get("annee1", 0.0),
        data["charges_sociales"]["employes"].get("annee2", 0.0),
        data["charges_sociales"]["employes"].get("annee3", 0.0)
    ]
    
    # Prélèvement dirigeant(s)
    salaires_dirigeants = [
        data["salaires"]["dirigeants"].get("annee1", 0.0),
        data["salaires"]["dirigeants"].get("annee2", 0.0),
        data["salaires"]["dirigeants"].get("annee3", 0.0)
    ]
    
    # Charges sociales dirigeant(s)
    charges_sociales_dirigeants = [
        data["charges_sociales"]["dirigeants"].get("annee1", 0.0),
        data["charges_sociales"]["dirigeants"].get("annee2", 0.0),
        data["charges_sociales"]["dirigeants"].get("annee3", 0.0)
    ]
    
    # Excédent brut d'exploitation = Valeur ajoutée - Impôts et taxes - Salaires - Charges sociales
    ebe = [
        valeur_ajoutee[0] - impots_et_taxes[0] - salaires_employes[0] - charges_sociales_employes[0] - salaires_dirigeants[0] - charges_sociales_dirigeants[0],
        valeur_ajoutee[1] - impots_et_taxes[1] - salaires_employes[1] - charges_sociales_employes[1] - salaires_dirigeants[1] - charges_sociales_dirigeants[1],
        valeur_ajoutee[2] - impots_et_taxes[2] - salaires_employes[2] - charges_sociales_employes[2] - salaires_dirigeants[2] - charges_sociales_dirigeants[2]
    ]
    
    # Frais bancaires, charges financières
    frais_bancaires = [
        charges_fixes_annee1.get("Frais bancaires et terminal carte bleue", 0.0),
        charges_fixes_annee2.get("Frais bancaires et terminal carte bleue", 0.0),
        charges_fixes_annee3.get("Frais bancaires et terminal carte bleue", 0.0)
    ]
    
    # Intérêts des prêts
    interets_prets = data.get("interets_prets", {
        "annee1": 0.0,
        "annee2": 0.0,
        "annee3": 0.0
    })
    
    # Ajouter les intérêts des prêts aux autres frais financiers
    frais_financiers = [
        interets_prets.get("annee1", 0.0),
        interets_prets.get("annee2", 0.0),
        interets_prets.get("annee3", 0.0)
    ]
    
    # Total des frais bancaires et charges financières
    total_frais_financiers = [
        frais_bancaires[0] + frais_financiers[0],
        frais_bancaires[1] + frais_financiers[1],
        frais_bancaires[2] + frais_financiers[2]
    ]
    
    # Dotations aux amortissements (supposées nulles si non fournies)
    amortissements = [0.0, 0.0, 0.0]
    
    # Résultat avant impôts = EBE - Frais bancaires - Amortissements
    resultat_avant_impots = [
        ebe[0] - total_frais_financiers[0] - amortissements[0],
        ebe[1] - total_frais_financiers[1] - amortissements[1],
        ebe[2] - total_frais_financiers[2] - amortissements[2]
    ]
    
    # Impôt sur les sociétés (selon la formule progressive)
    impot_societes = [
        calculer_impot_societes(resultat_avant_impots[0]),
        calculer_impot_societes(resultat_avant_impots[1]),
        calculer_impot_societes(resultat_avant_impots[2])
    ]
    
    # Résultat net comptable (résultat de l'exercice)
    resultat_net = [
        resultat_avant_impots[0] - impot_societes[0],
        resultat_avant_impots[1] - impot_societes[1],
        resultat_avant_impots[2] - impot_societes[2]
    ]
    
    # Préparation des données pour le tableau
    tableau = {
        "": [
            "Produits d'exploitation",
            "Chiffre d'affaires HT vente de marchandises",
            "Chiffre d'affaires HT services",
            "",
            "Charges d'exploitation",
            "Achats consommés",
            "",
            "Marge brute",
            "Charges externes",
            ""
        ],
        "Année 1": [
            f"{total_ca[0]:,.2f} €",
            f"{ca_marchandises[0]:,.2f} €",
            f"{ca_services[0]:,.2f} €",
            "",
            "",
            f"{charges_variables[0]:,.2f} €",
            "",
            f"{marge_brute[0]:,.2f} €",
            "",
            ""
        ],
        "Année 2": [
            f"{total_ca[1]:,.2f} €",
            f"{ca_marchandises[1]:,.2f} €",
            f"{ca_services[1]:,.2f} €",
            "",
            "",
            f"{charges_variables[1]:,.2f} €",
            "",
            f"{marge_brute[1]:,.2f} €",
            "",
            ""
        ],
        "Année 3": [
            f"{total_ca[2]:,.2f} €",
            f"{ca_marchandises[2]:,.2f} €",
            f"{ca_services[2]:,.2f} €",
            "",
            "",
            f"{charges_variables[2]:,.2f} €",
            "",
            f"{marge_brute[2]:,.2f} €",
            "",
            ""
        ]
    }
    
    # Ajouter les charges détaillées au tableau
    for charge in liste_charges:
        tableau[""].append(charge)
        tableau["Année 1"].append(f"{charges_detaillees[charge][0]:,.2f} €")
        tableau["Année 2"].append(f"{charges_detaillees[charge][1]:,.2f} €")
        tableau["Année 3"].append(f"{charges_detaillees[charge][2]:,.2f} €")
    
    # Ajouter le total des charges externes
    tableau[""].append("Total Charges externes")
    tableau["Année 1"].append(f"{total_charges_fixes[0]:,.2f} €")
    tableau["Année 2"].append(f"{total_charges_fixes[1]:,.2f} €")
    tableau["Année 3"].append(f"{total_charges_fixes[2]:,.2f} €")
    
    # Continuer à remplir le tableau
    additional_rows = {
        "Valeur ajoutée": valeur_ajoutee,
        "Impôts et taxes": impots_et_taxes,
        "Salaires employés": salaires_employes,
        "Charges sociales employés": charges_sociales_employes,
        "Prélèvement dirigeant(s)": salaires_dirigeants,
        "Charges sociales dirigeant(s)": charges_sociales_dirigeants,
        "Excédent brut d'exploitation": ebe,
        "Frais bancaires, charges financières": total_frais_financiers,
        "Dotations aux amortissements": amortissements,
        "Résultat avant impôts": resultat_avant_impots,
        "Impôt sur les sociétés": impot_societes,
        "Résultat net comptable (résultat de l'exercice)": resultat_net
    }
    
    for key, values in additional_rows.items():
        tableau[""].append(key)
        tableau["Année 1"].append(f"{values[0]:,.2f} €")
        tableau["Année 2"].append(f"{values[1]:,.2f} €")
        tableau["Année 3"].append(f"{values[2]:,.2f} €")
    
    # Créer le DataFrame
    df_resultats = pd.DataFrame(tableau)
    
    # Afficher le tableau
    st.table(df_resultats)
    
    # Ajouter les variables calculées au dictionnaire 'data'
    data["compte_de_resultat"] = {
        "total_ca": total_ca,
        "ca_marchandises": ca_marchandises,
        "ca_services": ca_services,
        "charges_variables": charges_variables,
        "marge_brute": marge_brute,
        "charges_fixes": total_charges_fixes,
        "valeur_ajoutee": valeur_ajoutee,
        "impots_et_taxes": impots_et_taxes,
        "salaires_employes": salaires_employes,
        "charges_sociales_employes": charges_sociales_employes,
        "salaires_dirigeants": salaires_dirigeants,
        "charges_sociales_dirigeants": charges_sociales_dirigeants,
        "ebe": ebe,
        "frais_bancaires": frais_bancaires,
        "frais_financiers": frais_financiers,
        "total_frais_financiers": total_frais_financiers,
        "amortissements": amortissements,
        "resultat_avant_impots": resultat_avant_impots,
        "impot_societes": impot_societes,
        "resultat_net": resultat_net
    }
    
    # Enregistrer les données dans la session
    st.session_state["data"] = data
    
    # Préparer les données d'exportation pour Compte de résultats prévisionnel
    export_table_compte = []
    for index, row in df_resultats.iterrows():
        export_table_compte.append({
            "Description": row[""],
            "Année 1": row["Année 1"],
            "Année 2": row["Année 2"],
            "Année 3": row["Année 3"]
        })
    
    # Stocker les données d'exportation dans la session
    st.session_state['export_data_compte_resultats_previsionnel'] = {
        "table_data": export_table_compte
    }
    
    # Section Export
    st.header("Exporter les données")
    # Bouton pour télécharger le document complet
    st.button("Télécharger le Document Word Complet", on_click=telecharger_document_complet)
    


import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT

def page_soldes_intermediaires_de_gestion():
    st.title("Soldes intermédiaires de gestion")
    
    # Récupérer les données de la session
    data = st.session_state.get("data", {})
    
    # Récupérer les informations du projet
    info_generales = data.get("informations_generales", {})
    projet = info_generales.get("intitule_projet", "N/A")
    porteur_projet = info_generales.get("prenom_nom", "N/A")
    
    # Afficher les informations du projet
    st.write(f"**Projet :** {projet}")
    st.write(f"**Porteur de projet :** {porteur_projet}")
    
    st.write("---")
    
    # Récupérer les données nécessaires pour les calculs
    compte_resultat = data.get("compte_de_resultat", {})
    
    # Fonction pour assurer que les listes contiennent trois éléments et convertir en float
    def get_three_years_data(key):
        values = compte_resultat.get(key, [])
        processed_values = []
        for v in values:
            try:
                processed_values.append(float(v))
            except (ValueError, TypeError):
                processed_values.append(0.0)
        # Compléter avec 0.0 si moins de 3 éléments
        while len(processed_values) < 3:
            processed_values.append(0.0)
        return processed_values[:3]
    
    # Récupération des données avec validation
    total_ca = get_three_years_data("total_ca")
    ca_marchandises = get_three_years_data("ca_marchandises")
    ca_services = get_three_years_data("ca_services")
    achats_consommes = get_three_years_data("achats_consommes")  # Actuellement défini à [0.0, 0.0, 0.0]
    charges_fixes = get_three_years_data("charges_fixes")
    impot_societes = get_three_years_data("impot_societes")
    impots_et_taxes = get_three_years_data("impots_et_taxes")
    salaires_employes = get_three_years_data("salaires_employes")
    charges_sociales_employes = get_three_years_data("charges_sociales_employes")
    salaires_dirigeants = get_three_years_data("salaires_dirigeants")
    charges_sociales_dirigeants = get_three_years_data("charges_sociales_dirigeants")
    amortissements = get_three_years_data("amortissements")
    total_frais_financiers = get_three_years_data("total_frais_financiers")
    
    # Calcul des différents soldes intermédiaires
    ventes_production_reelle = [ca_marchandises[i] + ca_services[i] for i in range(3)]
    marge_globale = [ventes_production_reelle[i] - achats_consommes[i] for i in range(3)]
    valeur_ajoutee = [marge_globale[i] - charges_fixes[i] for i in range(3)]
    charges_personnel = [
        salaires_employes[i] + charges_sociales_employes[i] + salaires_dirigeants[i] + charges_sociales_dirigeants[i]
        for i in range(3)
    ]
    ebe = [valeur_ajoutee[i] - impots_et_taxes[i] - charges_personnel[i] for i in range(3)]
    resultat_exploitation = [ebe[i] - amortissements[i] for i in range(3)]
    resultat_financier = [-total_frais_financiers[i] for i in range(3)]
    resultat_courant = [resultat_exploitation[i] + resultat_financier[i] for i in range(3)]
    resultat_exercice = [resultat_courant[i] - impot_societes[i] for i in range(3)]
    capacite_autofinancement = [resultat_exercice[i] + amortissements[i] for i in range(3)]
    
    # Fonction de calcul des pourcentages avec gestion de la division par zéro
    def calculate_percentage(value, ca):
        return (value / ca * 100) if ca != 0 else 0.0
    
    # Préparation des données pour le tableau
    soldes = [
        "Chiffre d'affaires",
        "Ventes + production réelle",
        "Achats consommés",
        "Marge globale",
        "Charges externes",
        "Valeur ajoutée",
        "Impôts et taxes",
        "Charges de personnel",
        "Excédent brut d'exploitation (EBE)",
        "Dotations aux amortissements",
        "Résultat d'exploitation",
        "Charges financières",
        "Résultat financier",
        "Résultat courant",
        "Résultat de l'exercice",
        "Capacité d'autofinancement"
    ]
    
    # Initialiser le data_table avec les soldes
    data_table = {"Soldes intermédiaires de gestion": soldes}
    
    # Ajouter les données pour chaque année et leurs pourcentages
    for year in range(3):
        data_table[f"Année {year+1}"] = [
            total_ca[year],
            ventes_production_reelle[year],
            achats_consommes[year],
            marge_globale[year],
            charges_fixes[year],
            valeur_ajoutee[year],
            impots_et_taxes[year],
            charges_personnel[year],
            ebe[year],
            amortissements[year],
            resultat_exploitation[year],
            total_frais_financiers[year],
            resultat_financier[year],
            resultat_courant[year],
            resultat_exercice[year],
            capacite_autofinancement[year]
        ]
        
        data_table[f"% Année {year+1}"] = [
            100.0,  # Chiffre d'affaires
            100.0,  # Ventes + production réelle
            calculate_percentage(achats_consommes[year], total_ca[year]),
            calculate_percentage(marge_globale[year], total_ca[year]),
            calculate_percentage(charges_fixes[year], total_ca[year]),
            calculate_percentage(valeur_ajoutee[year], total_ca[year]),
            calculate_percentage(impots_et_taxes[year], total_ca[year]),
            calculate_percentage(charges_personnel[year], total_ca[year]),
            calculate_percentage(ebe[year], total_ca[year]),
            calculate_percentage(amortissements[year], total_ca[year]),
            calculate_percentage(resultat_exploitation[year], total_ca[year]),
            calculate_percentage(total_frais_financiers[year], total_ca[year]),
            calculate_percentage(resultat_financier[year], total_ca[year]),
            calculate_percentage(resultat_courant[year], total_ca[year]),
            calculate_percentage(resultat_exercice[year], total_ca[year]),
            calculate_percentage(capacite_autofinancement[year], total_ca[year])
        ]
    
    # Créer le DataFrame avec les données
    df = pd.DataFrame(data_table)
    
    # Définir l'ordre des colonnes alternées entre "Année x" et "%"
    columns_order = ["Soldes intermédiaires de gestion"]
    for year in range(3):
        columns_order.append(f"Année {year+1}")
        columns_order.append(f"% Année {year+1}")
    df = df[columns_order]
    
    # Afficher le tableau avec une mise en forme améliorée
    st.dataframe(
        df.style.format({
            "Année 1": "{:,.2f} €",
            "Année 2": "{:,.2f} €",
            "Année 3": "{:,.2f} €",
            "% Année 1": "{:.2f}%",
            "% Année 2": "{:.2f}%",
            "% Année 3": "{:.2f}%"
        }).set_properties(**{
            'text-align': 'right'
        }).set_table_styles([{
            'selector': 'th',
            'props': [('text-align', 'center')]
        }])
    )
    
    # Stocker les résultats dans les données pour exportation
    data["soldes_intermediaires_de_gestion"] = {
        "ca": total_ca,
        "ventes_production_reelle": ventes_production_reelle,
        "achats_consommes": achats_consommes,
        "marge_globale": marge_globale,
        "charges_externes": charges_fixes,
        "valeur_ajoutee": valeur_ajoutee,
        "impots_et_taxes": impots_et_taxes,
        "charges_personnel": charges_personnel,
        "ebe": ebe,
        "dotations_aux_amortissements": amortissements,
        "resultat_exploitation": resultat_exploitation,
        "charges_financieres": total_frais_financiers,
        "resultat_financier": resultat_financier,
        "resultat_courant": resultat_courant,
        "resultat_exercice": resultat_exercice,
        "capacite_autofinancement": capacite_autofinancement
    }
    
    # Enregistrer les données mises à jour dans la session
    st.session_state["data"] = data
    
    # Préparer les données d'exportation pour Soldes Intermédiaires de Gestion avec % colonnes
    export_table_soldes = []
    for idx, solde in enumerate(soldes):
        export_table_soldes.append({
            "Description": solde,
            "Année 1": data_table["Année 1"][idx],
            "% Année 1": data_table["% Année 1"][idx],
            "Année 2": data_table["Année 2"][idx],
            "% Année 2": data_table["% Année 2"][idx],
            "Année 3": data_table["Année 3"][idx],
            "% Année 3": data_table["% Année 3"][idx]
        })
    
    # Stocker les données d'exportation dans la session
    st.session_state['export_data_soldes_intermediaires_de_gestion'] = {
        "projet": projet,
        "porteur_projet": porteur_projet,
        "table_data": export_table_soldes
    }
    
    # Section Export
    st.header("Exporter les données")
    
    # Bouton pour télécharger le Markdown
    if st.button("Télécharger Soldes Intermédiaires en Markdown"):
        export_data = st.session_state.get('export_data_soldes_intermediaires_de_gestion', {})
        if not export_data or "table_data" not in export_data:
            st.error("Aucune donnée disponible pour l'exportation.")
        else:
            # Construire le contenu Markdown
            markdown_content = f"# Soldes intermédiaires de gestion\n\n**Projet :** {export_data['projet']}\n\n**Porteur de projet :** {export_data['porteur_projet']}\n\n"
            markdown_content += "---\n\n"
            
            # Créer un DataFrame pour Markdown
            df_markdown = pd.DataFrame(export_data['table_data'])
            markdown_content += df_markdown.to_markdown(index=False)
            
            markdown_content += "\n\n---\n\n"
            
            markdown_bytes = markdown_content.encode('utf-8')
            st.download_button(
                label="Télécharger le Markdown",
                data=markdown_bytes,
                file_name="soldes_intermediaires_gestion.md",
                mime="text/markdown"
            )
    
    # Bouton pour télécharger le fichier Word
    if st.button("Télécharger Soldes Intermédiaires en Word"):
        export_data_soldes = st.session_state.get('export_data_soldes_intermediaires_de_gestion', {})
        export_data_compte = st.session_state.get('export_data_compte_resultats_previsionnel', {})
        export_data_amortissements = st.session_state.get('export_data_detail_amortissements', {})
        export_data_investissements = st.session_state.get('export_data_investissements', {})
        export_data_salaires = st.session_state.get('export_data_salaires_charges_sociales', {})
        
        if not export_data_soldes or "table_data" not in export_data_soldes:
            st.error("Aucune donnée disponible pour l'exportation des Soldes intermédiaires de gestion.")
            return
        
        # Vérifiez que toutes les autres sections sont également exportées
        if not all([
            export_data_investissements.get("table_data"),
            export_data_salaires.get("table_data"),
            export_data_amortissements.get("amortissements"),
            export_data_compte.get("table_data")
        ]):
            st.error("Toutes les sections doivent être remplies avant de télécharger le document complet.")
            return
        
        # Créer un document Word
        doc = Document()
        
        ### 1. Ajouter la section Investissements et Financements ###
        doc.add_heading('Investissements et Financements', level=1)
        doc.add_paragraph(f"**Projet :** {export_data_investissements.get('projet', 'N/A')}")
        doc.add_paragraph(f"**Porteur de projet :** {export_data_investissements.get('porteur_projet', 'N/A')}")
        doc.add_paragraph(f"**Equilibre :** {export_data_investissements.get('equilibre', {}).get('message', '')}")
        doc.add_page_break()
        
        # Créer le tableau Investissements et Financements
        table_inv = doc.add_table(rows=1, cols=4)
        table_inv.style = 'Light List Accent 1'
        table_inv.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells_inv = table_inv.rows[0].cells
        headers_inv = ["Investissements", "Taux (%)", "Durée (mois)", "Montant (€)"]
        for i, header in enumerate(headers_inv):
            hdr_cells_inv[i].text = header
            for paragraph in hdr_cells_inv[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            hdr_cells_inv[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        for row in export_data_investissements['table_data']:
            row_cells = table_inv.add_row().cells
            row_cells[0].text = row.get("Investissements", "")
            row_cells[1].text = row.get("Taux (%)", "")
            row_cells[2].text = str(row.get("Durée (mois)", "")) if row.get("Durée (mois)", "") != "-" else "-"
            row_cells[3].text = row.get("Montant (€)", "")
            
            # Mise en forme des lignes spécifiques
            if row["Investissements"] in ["INVESTISSEMENTS", "FINANCEMENT DES INVESTISSEMENTS", "TOTAL SUBVENTIONS", "TOTAL EMPRUNTS"]:
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
            elif "TOTAL" in row["Investissements"]:
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
            
            # Alignement des cellules
            row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        ### 2. Ajouter la section Salaires et Charges Sociales ###
        doc.add_heading('Salaires et Charges Sociales', level=1)
        doc.add_paragraph(f"**Projet :** {export_data_salaires.get('projet', 'N/A')}")
        doc.add_paragraph(f"**Porteur de projet :** {export_data_salaires.get('porteur_projet', 'N/A')}")
        doc.add_paragraph(f"**Statut juridique :** {export_data_salaires.get('statut_juridique', 'N/A')}")
        doc.add_paragraph(f"**Bénéfice de l'ACRE :** {export_data_salaires.get('benefice_accre', 'N/A')}")
        doc.add_paragraph(f"**Statut social du (des) dirigeant(s) :** {export_data_salaires.get('statut_social_dirigeant', 'N/A')}")
        doc.add_paragraph("---")
        
        # Créer le tableau Salaires et Charges Sociales
        table_sal = doc.add_table(rows=1, cols=4)
        table_sal.style = 'Light List Accent 1'
        table_sal.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells_sal = table_sal.rows[0].cells
        headers_sal = ["Description", "Année 1", "Année 2", "Année 3"]
        for i, header in enumerate(headers_sal):
            hdr_cells_sal[i].text = header
            for paragraph in hdr_cells_sal[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            hdr_cells_sal[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        for row in export_data_salaires['table_data']:
            row_cells = table_sal.add_row().cells
            row_cells[0].text = row.get("Description", "")
            row_cells[1].text = row.get("Année 1", "")
            row_cells[2].text = row.get("Année 2", "")
            row_cells[3].text = row.get("Année 3", "")
            
            # Alignement des cellules
            row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        ### 3. Ajouter la section Détail des Amortissements ###
        doc.add_heading('Détail des Amortissements', level=1)
        
        # Créer le tableau Détail des Amortissements
        table_amort = doc.add_table(rows=1, cols=4)
        table_amort.style = 'Light List Accent 1'
        table_amort.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells_amort = table_amort.rows[0].cells
        headers_amort = ["Amortissement", "Année 1", "Année 2", "Année 3"]
        for i, header in enumerate(headers_amort):
            hdr_cells_amort[i].text = header
            for paragraph in hdr_cells_amort[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            hdr_cells_amort[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        for row in export_data_amortissements['amortissements']:
            row_cells = table_amort.add_row().cells
            row_cells[0].text = row.get("Amortissement", "")
            row_cells[1].text = row.get("Année 1", "")
            row_cells[2].text = row.get("Année 2", "")
            row_cells[3].text = row.get("Année 3", "")
        
        # Ajouter une note
        doc.add_paragraph()
        doc.add_paragraph("Les amortissements sont calculés en fonction de la durée d'amortissement spécifiée.")
        
        ### 4. Ajouter la section Compte de Résultats Prévisionnel ###
        doc.add_heading('Compte de Résultats Prévisionnel', level=1)
        
        # Créer le tableau Compte de Résultats Prévisionnel
        table_compte = doc.add_table(rows=1, cols=4)
        table_compte.style = 'Light List Accent 1'
        table_compte.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells_compte = table_compte.rows[0].cells
        headers_compte = ["Description", "Année 1", "Année 2", "Année 3"]
        for i, header in enumerate(headers_compte):
            hdr_cells_compte[i].text = header
            for paragraph in hdr_cells_compte[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            hdr_cells_compte[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        for row in export_data_compte['table_data']:
            row_cells = table_compte.add_row().cells
            row_cells[0].text = row.get("Description", "")
            row_cells[1].text = row.get("Année 1", "")
            row_cells[2].text = row.get("Année 2", "")
            row_cells[3].text = row.get("Année 3", "")
        
        # Ajouter une note
        doc.add_paragraph()
        doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
        
        ### 5. Ajouter la section Soldes Intermédiaires de Gestion ###
        doc.add_heading('Soldes intermédiaires de gestion', level=1)
        
        # Créer le tableau Soldes intermédiaires de gestion avec 7 colonnes
        table_soldes = doc.add_table(rows=1, cols=7)
        table_soldes.style = 'Light List Accent 1'
        table_soldes.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells_soldes = table_soldes.rows[0].cells
        headers_soldes = ["Description", "Année 1", "% Année 1", "Année 2", "% Année 2", "Année 3", "% Année 3"]
        for i, header in enumerate(headers_soldes):
            hdr_cells_soldes[i].text = header
            for paragraph in hdr_cells_soldes[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            hdr_cells_soldes[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        for row in export_data_soldes['table_data']:
            row_cells = table_soldes.add_row().cells
            row_cells[0].text = row.get("Description", "")
            row_cells[1].text = f"{row.get('Année 1', 0.0):,.2f} €"
            row_cells[2].text = f"{row.get('% Année 1', 0.0):.2f}%"
            row_cells[3].text = f"{row.get('Année 2', 0.0):,.2f} €"
            row_cells[4].text = f"{row.get('% Année 2', 0.0):.2f}%"
            row_cells[5].text = f"{row.get('Année 3', 0.0):,.2f} €"
            row_cells[6].text = f"{row.get('% Année 3', 0.0):.2f}%"
            
            # Alignement des cellules de pourcentage
            row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_cells[6].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Ajouter une note
        doc.add_paragraph()
        doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
        
        # Enregistrer le document dans un buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Bouton de téléchargement
        st.download_button(
            label="Télécharger le Document Word Complet",
            data=buffer,
            file_name="document_complet_financier.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        # Message de confirmation
        st.success("Le document Word complet a été généré avec succès !")



    # Enregistrer les données mises à jour dans la session
    st.session_state["data"] = data
    
    
    
    
    

def telecharger_document_complets():
    # Récupérer les données exportées de toutes les sections
    export_data_investissements = st.session_state.get('export_data_investissements', {})
    export_data_salaires = st.session_state.get('export_data_salaires_charges_sociales', {})
    export_data_amortissements = st.session_state.get('export_data_detail_amortissements', {})
    export_data_compte = st.session_state.get('export_data_compte_resultats_previsionnel', {})
    export_data_soldes = st.session_state.get('export_data_soldes_intermediaires_de_gestion', {})
    export_data_capacite = st.session_state.get('export_data_capacite_autofinancement', {})
    export_data_seuil = st.session_state.get('export_data_seuil_rentabilite_economique', {})
    
    # Vérifiez que toutes les données sont présentes
    if not all([
        export_data_investissements.get("table_data"),
        export_data_salaires.get("table_data"),
        export_data_amortissements.get("amortissements"),
        export_data_compte.get("table_data"),
        export_data_soldes.get("table_data"),
        export_data_capacite.get("table_data"),
        export_data_seuil.get("table_data")
    ]):
        st.error("Toutes les sections doivent être remplies avant de télécharger le document complet.")
        return
    
    # Créer un document Word
    doc = Document()
    
    ### 1. Ajouter la section Investissements et Financements ###
    doc.add_heading('Investissements et Financements', level=1)
    doc.add_paragraph(f"**Projet :** {export_data_investissements.get('projet', 'N/A')}")
    doc.add_paragraph(f"**Porteur de projet :** {export_data_investissements.get('porteur_projet', 'N/A')}")
    doc.add_page_break()
    
    # Créer le tableau Investissements et Financements dans Word
    table_inv = doc.add_table(rows=1, cols=4)
    table_inv.style = 'Light List Accent 1'
    table_inv.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_inv = table_inv.rows[0].cells
    headers_inv = ["Investissements", "Taux (%)", "Durée (mois)", "Montant (€)"]
    for i, header in enumerate(headers_inv):
        hdr_cells_inv[i].text = header
        for paragraph in hdr_cells_inv[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_inv[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_investissements['table_data']:
        row_cells = table_inv.add_row().cells
        row_cells[0].text = row.get("Investissements", "")
        row_cells[1].text = row.get("Taux (%)", "")
        row_cells[2].text = str(row.get("Durée (mois)", "")) if row.get("Durée (mois)", "") != "-" else "-"
        row_cells[3].text = row.get("Montant (€)", "")
        
        # Mise en forme des lignes spécifiques
        if row["Investissements"] in ["INVESTISSEMENTS", "FINANCEMENT DES INVESTISSEMENTS", "TOTAL SUBVENTIONS", "TOTAL EMPRUNTS"]:
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        elif "TOTAL" in row["Investissements"]:
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        
        # Alignement des cellules
        row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    ### 2. Ajouter la section Salaires et Charges Sociales ###
    doc.add_heading('Salaires et Charges Sociales', level=1)
    doc.add_paragraph(f"**Projet :** {export_data_salaires.get('projet', 'N/A')}")
    doc.add_paragraph(f"**Porteur de projet :** {export_data_salaires.get('porteur_projet', 'N/A')}")
    doc.add_paragraph("---")
    
    # Créer le tableau Salaires et Charges Sociales dans Word
    table_sal = doc.add_table(rows=1, cols=4)
    table_sal.style = 'Light List Accent 1'
    table_sal.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_sal = table_sal.rows[0].cells
    headers_sal = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_sal):
        hdr_cells_sal[i].text = header
        for paragraph in hdr_cells_sal[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_sal[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_salaires['table_data']:
        row_cells = table_sal.add_row().cells
        row_cells[0].text = row.get("Description", "")
        row_cells[1].text = row.get("Année 1", "")
        row_cells[2].text = row.get("Année 2", "")
        row_cells[3].text = row.get("Année 3", "")
        
        # Alignement des cellules
        row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    ### 3. Ajouter la section Détail des Amortissements ###
    doc.add_heading('Détail des Amortissements', level=1)
    
    # Créer le tableau Détail des Amortissements dans Word
    table_amort = doc.add_table(rows=1, cols=4)
    table_amort.style = 'Light List Accent 1'
    table_amort.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_amort = table_amort.rows[0].cells
    headers_amort = ["Amortissement", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_amort):
        hdr_cells_amort[i].text = header
        for paragraph in hdr_cells_amort[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_amort[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Ajouter les données Amortissements au tableau
    for row in export_data_amortissements['amortissements']:
        row_cells_amort = table_amort.add_row().cells
        row_cells_amort[0].text = row.get("Amortissement", "")
        row_cells_amort[1].text = row.get("Année 1", "")
        row_cells_amort[2].text = row.get("Année 2", "")
        row_cells_amort[3].text = row.get("Année 3", "")
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les amortissements sont calculés en fonction de la durée d'amortissement spécifiée.")
    
    ### 4. Ajouter la section Compte de Résultats Prévisionnel ###
    doc.add_heading('Compte de Résultats Prévisionnel', level=1)
    
    # Créer le tableau Compte de Résultats Prévisionnel dans Word
    table_compte = doc.add_table(rows=1, cols=4)
    table_compte.style = 'Light List Accent 1'
    table_compte.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_compte = table_compte.rows[0].cells
    headers_compte = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_compte):
        hdr_cells_compte[i].text = header
        for paragraph in hdr_cells_compte[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_compte[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_compte['table_data']:
        row_cells_compte = table_compte.add_row().cells
        row_cells_compte[0].text = row.get("Description", "")
        row_cells_compte[1].text = row.get("Année 1", "")
        row_cells_compte[2].text = row.get("Année 2", "")
        row_cells_compte[3].text = row.get("Année 3", "")
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    ### 5. Ajouter la section Soldes Intermédiaires de Gestion ###
    doc.add_heading('Soldes intermédiaires de gestion', level=1)
    
    # Créer le tableau Soldes intermédiaires de gestion dans Word
    table_soldes = doc.add_table(rows=1, cols=4)
    table_soldes.style = 'Light List Accent 1'
    table_soldes.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_soldes = table_soldes.rows[0].cells
    headers_soldes = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_soldes):
        hdr_cells_soldes[i].text = header
        for paragraph in hdr_cells_soldes[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_soldes[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_soldes['table_data']:
        row_cells_soldes = table_soldes.add_row().cells
        row_cells_soldes[0].text = row.get("Description", "")
        row_cells_soldes[1].text = f"{row.get('Année 1', 0):,.2f} €"
        row_cells_soldes[2].text = f"{row.get('Année 2', 0):,.2f} €"
        row_cells_soldes[3].text = f"{row.get('Année 3', 0):,.2f} €"
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    ### 6. Ajouter la section Capacité d'Autofinancement ###
    doc.add_heading('Capacité d\'autofinancement', level=1)
    
    # Créer le tableau Capacité d'Autofinancement dans Word
    table_cap = doc.add_table(rows=1, cols=4)
    table_cap.style = 'Light List Accent 1'
    table_cap.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_cap = table_cap.rows[0].cells
    headers_cap = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_cap):
        hdr_cells_cap[i].text = header
        for paragraph in hdr_cells_cap[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_cap[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_capacite['table_data']:
        row_cells_cap = table_cap.add_row().cells
        row_cells_cap[0].text = row.get("Description", "")
        row_cells_cap[1].text = f"{row.get('Année 1', 0):,.2f} €"
        row_cells_cap[2].text = f"{row.get('Année 2', 0):,.2f} €"
        row_cells_cap[3].text = f"{row.get('Année 3', 0):,.2f} €"
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    ### 7. Ajouter la section Seuil de Rentabilité Économique ###
    doc.add_heading('Seuil de rentabilité économique', level=1)
    
    # Créer le tableau Seuil de Rentabilité Économique dans Word
    table_seuil = doc.add_table(rows=1, cols=4)
    table_seuil.style = 'Light List Accent 1'
    table_seuil.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_seuil = table_seuil.rows[0].cells
    headers_seuil = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_seuil):
        hdr_cells_seuil[i].text = header
        for paragraph in hdr_cells_seuil[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_seuil[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_seuil['table_data']:
        row_cells_seuil = table_seuil.add_row().cells
        row_cells_seuil[0].text = row.get("Description", "")
        row_cells_seuil[1].text = f"{row.get('Année 1', 0):,.0f} €"
        row_cells_seuil[2].text = f"{row.get('Année 2', 0):,.0f} €"
        row_cells_seuil[3].text = f"{row.get('Année 3', 0):,.0f} €"
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    # Enregistrer le document dans un buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Bouton de téléchargement
    st.download_button(
        label="Télécharger le Document Word Complet",
        data=buffer,
        file_name="document_complet_financier.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    
    # Message de confirmation
    st.success("Le document Word complet a été généré avec succès !")


def calculate_remboursements_emprunts(financements, years=3):
    """
    Votre fonction de calcul existante...
    """
    remboursements = [0.0 for _ in range(years)]  # Initialiser les remboursements pour chaque année

    for loan_name, loan_info in financements.items():
        # Vérifier que loan_info est un dictionnaire et commence par "Prêt "
        if isinstance(loan_info, dict) and loan_name.startswith("Prêt "):
            required_keys = {"montant", "duree", "taux"}
            if not required_keys.issubset(loan_info.keys()):
                st.warning(f"Le prêt '{loan_name}' est incomplet et sera ignoré.")
                continue  # Ignorer les financements incomplets

            montant = loan_info.get("montant", 0.0)
            duree_mois = loan_info.get("duree", 60)  # Par défaut 60 mois
            taux_annuel = loan_info.get("taux", 5.0)  # Par défaut 5%
            principal_mensuel =  montant / duree_mois if duree_mois > 0 else 0.0

            # Calcul des remboursements par année basés sur principal_mensuel
            # Principal Year 1
            if duree_mois > 12:
                principal_year1 = principal_mensuel * 12
            else:
                principal_year1 = principal_mensuel * duree_mois

            # Principal Year 2
            if duree_mois - 12 < 0:
                principal_year2 = 0.0
            elif duree_mois > 24:
                principal_year2 = principal_mensuel * 12
            else:
                principal_year2 = principal_mensuel * (duree_mois - 12)

            # Principal Year 3
            if duree_mois - 24 < 0:
                principal_year3 = 0.0
            elif duree_mois > 36:
                principal_year3 = principal_mensuel * 12
            else:
                principal_year3 = principal_mensuel * (duree_mois - 24)

            # Ajouter les remboursements principaux au total par année
            remboursements[0] += round(principal_year1, 2)
            if years >= 2:
                remboursements[1] += round(principal_year2, 2)
            if years >= 3:
                remboursements[2] += round(principal_year3, 2)
        else:
            # Ignorer les financements qui ne sont pas des prêts (e.g., Apports, Subventions)
            continue

    return remboursements

def page_capacite_autofinancement():
    """
    Affiche le tableau de Capacité d'Autofinancement en utilisant les données de la session.
    """
    st.title("Capacité d'autofinancement")
    
    # Vérifier si les données sont présentes dans la session
    if "data" not in st.session_state:
        st.error("Les données ne sont pas initialisées. Veuillez initialiser la session.")
        return
    
    # Récupérer les données de la session
    data = st.session_state["data"]
    
    # Récupérer les informations du projet
    info_generales = data.get("informations_generales", {})
    projet = info_generales.get("intitule_projet", "N/A")
    porteur_projet = info_generales.get("prenom_nom", "N/A")
    
    # Afficher les informations du projet
    st.write(f"**Projet :** {projet}")
    st.write(f"**Porteur de projet :** {porteur_projet}")
    
    st.write("---")
    
    # Récupérer les soldes intermédiaires de gestion
    soldes_intermediaires = data.get("soldes_intermediaires_de_gestion", {})
    
    # Fonction pour convertir les valeurs en float, remplacer les erreurs par 0.0
    def safe_float_conversion(values):
        return [float(x) if isinstance(x, (int, float)) else 0.0 for x in values]
    
    # Récupérer et convertir les données nécessaires
    resultat_exercice = safe_float_conversion(soldes_intermediaires.get("resultat_exercice", [0.0, 0.0, 0.0]))
    dotations_aux_amortissements = safe_float_conversion(soldes_intermediaires.get("dotations_aux_amortissements", [0.0, 0.0, 0.0]))
    capacite_autofinancement = safe_float_conversion(soldes_intermediaires.get("capacite_autofinancement", [0.0, 0.0, 0.0]))
    
    # Récupérer les financements
    financements = data.get("financements", {})
    
    # Filtrer uniquement les prêts (dictionnaires) nommés avec "Prêt " pour éviter les subventions
    pret_financements = {
        k: v for k, v in financements.items()
        if isinstance(v, dict) and k.startswith("Prêt ")
    }
    
    # Calculer les remboursements des emprunts
    remboursements_emprunts = calculate_remboursements_emprunts(pret_financements, years=3)
    
    # Autofinancement net = Capacité d'autofinancement - Remboursements des emprunts
    autofinancement_net = [
        capacite_autofinancement[i] - remboursements_emprunts[i]
        for i in range(3)
    ]
    
    # Préparer les valeurs monétaires
    values = {
        "Année 1": [
            resultat_exercice[0],
            dotations_aux_amortissements[0],
            capacite_autofinancement[0],
            remboursements_emprunts[0],
            autofinancement_net[0]
        ],
        "Année 2": [
            resultat_exercice[1],
            dotations_aux_amortissements[1],
            capacite_autofinancement[1],
            remboursements_emprunts[1],
            autofinancement_net[1]
        ],
        "Année 3": [
            resultat_exercice[2],
            dotations_aux_amortissements[2],
            capacite_autofinancement[2],
            remboursements_emprunts[2],
            autofinancement_net[2]
        ]
    }
    
    # Préparer le tableau final avec les labels
    capacite_fonc = [
        "Résultat de l'exercice",
        "+ Dotation aux amortissements",
        "Capacité d'autofinancement",
        "- Remboursements des emprunts",
        "Autofinancement net"
    ]
    
    data_table = {
        "Capacité d'autofinancement": capacite_fonc,
        "Année 1": values["Année 1"],
        "Année 2": values["Année 2"],
        "Année 3": values["Année 3"]
    }
    
    # Créer le DataFrame avec les données
    df = pd.DataFrame(data_table)
    
    # Définir l'ordre des colonnes
    columns_order = ["Capacité d'autofinancement",
                     "Année 1",
                     "Année 2",
                     "Année 3"]
    df = df[columns_order]
    
    # Définir la fonction de formatage
    def format_value(x):
        if x == 0.0:
            return "-"
        else:
            return f"{x:,.2f} €"
    
    # Afficher le tableau avec une mise en forme améliorée
    st.dataframe(
        df.style.format({
            "Année 1": format_value,
            "Année 2": format_value,
            "Année 3": format_value,
        }).set_properties(**{
            'text-align': 'right'
        }).set_table_styles([{
            'selector': 'th',
            'props': [('text-align', 'center')]
        }])
    )
    
    # Stocker les résultats dans les données
    data["capacite_autofinancement"] = {
        "resultat_exercice": resultat_exercice,
        "dotations_aux_amortissements": dotations_aux_amortissements,
        "capacite_autofinancement": capacite_autofinancement,
        "remboursements_emprunts": remboursements_emprunts,
        "autofinancement_net": autofinancement_net
    }
    
    # Enregistrer les données mises à jour dans la session
    st.session_state["data"] = data
    
    # Préparer les données d'exportation pour Capacité d'Autofinancement
    export_table_capacite = []
    for idx, label in enumerate(capacite_fonc):
        export_table_capacite.append({
            "Description": label,
            "Année 1": values["Année 1"][idx],
            "Année 2": values["Année 2"][idx],
            "Année 3": values["Année 3"][idx]
        })
    
    # Stocker les données d'exportation dans la session
    st.session_state['export_data_capacite_autofinancement'] = {
        "projet": projet,
        "porteur_projet": porteur_projet,
        "table_data": export_table_capacite
    }
    
    # Section Export
    st.header("Exporter les données")
    
    if st.button("Télécharger le Document Word Complet", key="download_word_complet_investissements_et_financements"):
        telecharger_document_complet()

    
import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT

def page_seuil_rentabilite_economique():
    st.title("Seuil de rentabilité économique")
    
    # Vérifier si les données sont présentes dans la session
    if "data" not in st.session_state:
        st.error("Les données ne sont pas initialisées. Veuillez initialiser la session.")
        return
    
    data = st.session_state["data"]

    # Récupérer les informations du projet
    info_generales = data.get("informations_generales", {})
    projet = info_generales.get("intitule_projet", "N/A")
    porteur_projet = info_generales.get("prenom_nom", "N/A")
    
    # Afficher les informations du projet
    st.write(f"**Projet :** {projet}")
    st.write(f"**Porteur de projet :** {porteur_projet}")
    
    st.write("---")
    
    # Récupérer les données nécessaires avec les clés exactes
    compte_resultat = data.get("compte_de_resultat", {})
    soldes_intermediaires = data.get("soldes_intermediaires_de_gestion", {})
    
    # Récupérer 'ventes_production_reelle' et 'achats_consommes' du 'soldes_intermediaires'
    ventes_production_reelle = soldes_intermediaires.get("ventes_production_reelle", [0.0, 0.0, 0.0])
    achats_consommes = soldes_intermediaires.get("achats_consommes", [0.0, 0.0, 0.0])
    
    # Récupérer les charges nécessaires du 'soldes_intermediaires'
    charges_externes = soldes_intermediaires.get("charges_externes", [0.0, 0.0, 0.0])
    impots_et_taxes = soldes_intermediaires.get("impots_et_taxes", [0.0, 0.0, 0.0])
    charges_personnel = soldes_intermediaires.get("charges_personnel", [0.0, 0.0, 0.0])
    dotations_aux_amortissements = soldes_intermediaires.get("dotations_aux_amortissements", [0.0, 0.0, 0.0])
    charges_financieres = soldes_intermediaires.get("charges_financieres", [0.0, 0.0, 0.0])
    
    # Récupérer 'resultat_avant_impots' du 'compte_resultat'
    resultat_avant_impots = compte_resultat.get("resultat_avant_impots", [0.0, 0.0, 0.0])
    
    # Vérifier si les listes ont 3 éléments
    if not (len(ventes_production_reelle) == len(achats_consommes) == len(charges_externes) == len(impots_et_taxes) == len(charges_personnel) == len(dotations_aux_amortissements) == len(charges_financieres) == len(resultat_avant_impots) == 3):
        st.error("Les listes de données ne contiennent pas exactement 3 éléments. Veuillez vérifier les données.")
        return
    
    # Calcul des Coûts fixes pour chaque année
    couts_fixes = []
    for i in range(3):
        cout_fix = (
            charges_externes[i] +
            impots_et_taxes[i] +
            charges_personnel[i] +
            dotations_aux_amortissements[i] +
            charges_financieres[i]
        )
        couts_fixes.append(cout_fix)
   
    # Total des coûts variables = Achats consommés
    total_couts_variables = achats_consommes.copy()
    
    # Marge sur coûts variables
    marge_sur_couts_variables = []
    for i in range(3):
        marge = ventes_production_reelle[i] - total_couts_variables[i]
        marge_sur_couts_variables.append(marge)
    
    # Taux de marge sur coûts variables
    taux_marge_sur_couts_variables = []
    for i in range(3):
        if ventes_production_reelle[i] != 0:
            taux_marge = marge_sur_couts_variables[i] / ventes_production_reelle[i]
        else:
            taux_marge = 0.0
        taux_marge_sur_couts_variables.append(taux_marge)
    
    # Total des charges
    total_charges = []
    for i in range(3):
        total_charge = couts_fixes[i] + total_couts_variables[i]
        total_charges.append(total_charge)
   
    # Seuil de rentabilité (CA)
    seuil_rentabilite_ca = []
    for i in range(3):
        if taux_marge_sur_couts_variables[i] != 0:
            seuil_ca = couts_fixes[i] / taux_marge_sur_couts_variables[i]
        else:
            seuil_ca = 0.0
        seuil_rentabilite_ca.append(seuil_ca)
    
    # Excédent / insuffisance
    excedent_insuffisance = []
    for i in range(3):
        excedent = ventes_production_reelle[i] - seuil_rentabilite_ca[i]
        excedent_insuffisance.append(excedent)
    
    # Point mort
    point_mort_ca_par_jour_ouvre = []
    for i in range(3):
        point_mort = seuil_rentabilite_ca[i] / 250
        point_mort_ca_par_jour_ouvre.append(point_mort)
    
    # Préparation des données pour le tableau
    data_table = {
        "Seuil de rentabilite_economique": [
            "Ventes + Production réelle",
            "Achats consommés",
            "Total des coûts variables",
            "Marge sur coûts variables",
            "Taux de marge sur coûts variables",
            "Coûts fixes",
            "Total des charges",
            "Résultat courant avant impôts",
            "Seuil de rentabilite (chiffre d'affaires)",
            "Excédent / insuffisance",
            "Point mort en chiffre d'affaires par jour ouvré"
        ],
        "Année 1": [
            ventes_production_reelle[0],
            achats_consommes[0],
            total_couts_variables[0],
            marge_sur_couts_variables[0],
            taux_marge_sur_couts_variables[0],
            couts_fixes[0],
            total_charges[0],
            resultat_avant_impots[0],
            seuil_rentabilite_ca[0],
            excedent_insuffisance[0],
            point_mort_ca_par_jour_ouvre[0]
        ],
        "Année 2": [
            ventes_production_reelle[1],
            achats_consommes[1],
            total_couts_variables[1],
            marge_sur_couts_variables[1],
            taux_marge_sur_couts_variables[1],
            couts_fixes[1],
            total_charges[1],
            resultat_avant_impots[1],
            seuil_rentabilite_ca[1],
            excedent_insuffisance[1],
            point_mort_ca_par_jour_ouvre[1]
        ],
        "Année 3": [
            ventes_production_reelle[2],
            achats_consommes[2],
            total_couts_variables[2],
            marge_sur_couts_variables[2],
            taux_marge_sur_couts_variables[2],
            couts_fixes[2],
            total_charges[2],
            resultat_avant_impots[2],
            seuil_rentabilite_ca[2],
            excedent_insuffisance[2],
            point_mort_ca_par_jour_ouvre[2]
        ]
    }
    
    # Créer le DataFrame
    df = pd.DataFrame(data_table)
    
    # Définir "Seuil de rentabilite_economique" comme index
    df.set_index("Seuil de rentabilite_economique", inplace=True)
    
    # Définir le formatage pour chaque colonne
    format_dict = {
        "Année 1": "{:,.0f} €",
        "Année 2": "{:,.0f} €",
        "Année 3": "{:,.0f} €",
        "Taux de marge sur coûts variables": "{:.0f} %",
        "Point mort en chiffre d'affaires par jour ouvré": "{:.0f} €"
    }
    
    # Appliquer le formatage et afficher le tableau
    st.table(
        df.style.format(format_dict)
          .set_properties(**{'text-align': 'right'})
          .set_table_styles([{
              'selector': 'th',
              'props': [('text-align', 'center')]
          }])
    )
    
    # Stocker les résultats dans les données
    data["seuil_rentabilite_economique"] = {
        "ventes_production_reelle": ventes_production_reelle,
        "achats_consommes": achats_consommes,
        "total_couts_variables": total_couts_variables,
        "marge_sur_couts_variables": marge_sur_couts_variables,
        "taux_marge_sur_couts_variables": taux_marge_sur_couts_variables,
        "couts_fixes": couts_fixes,
        "total_charges": total_charges,
        "resultat_courant_avant_impots": resultat_avant_impots,
        "seuil_rentabilite_ca": seuil_rentabilite_ca,
        "excedent_insuffisance": excedent_insuffisance,
        "point_mort_ca_par_jour_ouvre": point_mort_ca_par_jour_ouvre
    }
    
    # Enregistrer les données dans la session
    st.session_state["data"] = data
    
    # Préparer les données d'exportation pour Seuil de rentabilité économique
    export_table_seuil = []
    for idx, label in enumerate(data_table["Seuil de rentabilite_economique"]):
        export_table_seuil.append({
            "Description": label,
            "Année 1": data_table["Année 1"][idx],
            "Année 2": data_table["Année 2"][idx],
            "Année 3": data_table["Année 3"][idx]
        })
    
    # Stocker les données d'exportation dans la session
    st.session_state['export_data_seuil_rentabilite_economique'] = {
        "projet": projet,
        "porteur_projet": porteur_projet,
        "table_data": export_table_seuil
    }
    
    # Section Export
    st.header("Exporter les données")
    
    # Bouton pour télécharger le fichier Word complet contenant tous les tableaux avec une clé unique
    if st.button("Télécharger le Document Word Complet", key="download_word_complet_seuil_rentabilite_economique"):
        telecharger_document_complet()


def telecharger_document_complet():
    # Récupérer les données exportées de toutes les sections
    export_data_investissements = st.session_state.get('export_data_investissements', {})
    export_data_salaires = st.session_state.get('export_data_salaires_charges_sociales', {})
    export_data_amortissements = st.session_state.get('export_data_detail_amortissements', {})
    export_data_compte = st.session_state.get('export_data_compte_resultats_previsionnel', {})
    export_data_soldes = st.session_state.get('export_data_soldes_intermediaires_de_gestion', {})
    export_data_capacite = st.session_state.get('export_data_capacite_autofinancement', {})
    export_data_seuil = st.session_state.get('export_data_seuil_rentabilite_economique', {})
    export_data_bfr = st.session_state.get('export_data_besoin_fonds_roulement', {})
    
    # Vérifiez que toutes les données sont présentes
    if not all([
        export_data_investissements.get("table_data"),
        export_data_salaires.get("table_data"),
        export_data_amortissements.get("amortissements"),
        export_data_compte.get("table_data"),
        export_data_soldes.get("table_data"),
        export_data_capacite.get("table_data"),
        export_data_seuil.get("table_data"),
        export_data_bfr.get("table_data")
    ]):
        st.error("Toutes les sections doivent être remplies avant de télécharger le document complet.")
        return
    
    # Créer un document Word
    doc = Document()
    
    ### 1. Ajouter la section Investissements et Financements ###
    doc.add_heading('Investissements et Financements', level=1)
    doc.add_paragraph(f"**Projet :** {export_data_investissements.get('projet', 'N/A')}")
    doc.add_paragraph(f"**Porteur de projet :** {export_data_investissements.get('porteur_projet', 'N/A')}")
    doc.add_page_break()
    
    # Créer le tableau Investissements et Financements dans Word
    table_inv = doc.add_table(rows=1, cols=4)
    table_inv.style = 'Light List Accent 1'
    table_inv.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_inv = table_inv.rows[0].cells
    headers_inv = ["Investissements", "Taux (%)", "Durée (mois)", "Montant (€)"]
    for i, header in enumerate(headers_inv):
        hdr_cells_inv[i].text = header
        for paragraph in hdr_cells_inv[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_inv[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_investissements['table_data']:
        row_cells = table_inv.add_row().cells
        row_cells[0].text = row.get("Investissements", "")
        row_cells[1].text = row.get("Taux (%)", "")
        row_cells[2].text = str(row.get("Durée (mois)", "")) if row.get("Durée (mois)", "") != "-" else "-"
        row_cells[3].text = row.get("Montant (€)", "")
        
        # Mise en forme des lignes spécifiques
        if row["Investissements"] in ["INVESTISSEMENTS", "FINANCEMENT DES INVESTISSEMENTS", "TOTAL SUBVENTIONS", "TOTAL EMPRUNTS"]:
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        elif "TOTAL" in row["Investissements"]:
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        
        # Alignement des cellules
        row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    ### 2. Ajouter la section Salaires et Charges Sociales ###
    doc.add_heading('Salaires et Charges Sociales', level=1)
    doc.add_paragraph(f"**Projet :** {export_data_salaires.get('projet', 'N/A')}")
    doc.add_paragraph(f"**Porteur de projet :** {export_data_salaires.get('porteur_projet', 'N/A')}")
    doc.add_paragraph("---")
    
    # Créer le tableau Salaires et Charges Sociales dans Word
    table_sal = doc.add_table(rows=1, cols=4)
    table_sal.style = 'Light List Accent 1'
    table_sal.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_sal = table_sal.rows[0].cells
    headers_sal = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_sal):
        hdr_cells_sal[i].text = header
        for paragraph in hdr_cells_sal[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_sal[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_salaires['table_data']:
        row_cells = table_sal.add_row().cells
        row_cells[0].text = row.get("Description", "")
        row_cells[1].text = row.get("Année 1", "")
        row_cells[2].text = row.get("Année 2", "")
        row_cells[3].text = row.get("Année 3", "")
        
        # Alignement des cellules
        row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    ### 3. Ajouter la section Détail des Amortissements ###
    doc.add_heading('Détail des Amortissements', level=1)
    
    # Créer le tableau Détail des Amortissements dans Word
    table_amort = doc.add_table(rows=1, cols=4)
    table_amort.style = 'Light List Accent 1'
    table_amort.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_amort = table_amort.rows[0].cells
    headers_amort = ["Amortissement", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_amort):
        hdr_cells_amort[i].text = header
        for paragraph in hdr_cells_amort[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_amort[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Ajouter les données Amortissements au tableau
    for row in export_data_amortissements['amortissements']:
        row_cells_amort = table_amort.add_row().cells
        row_cells_amort[0].text = row.get("Amortissement", "")
        row_cells_amort[1].text = row.get("Année 1", "")
        row_cells_amort[2].text = row.get("Année 2", "")
        row_cells_amort[3].text = row.get("Année 3", "")
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les amortissements sont calculés en fonction de la durée d'amortissement spécifiée.")
    
    ### 4. Ajouter la section Compte de Résultats Prévisionnel ###
    doc.add_heading('Compte de Résultats Prévisionnel', level=1)
    
    # Créer le tableau Compte de Résultats Prévisionnel dans Word
    table_compte = doc.add_table(rows=1, cols=4)
    table_compte.style = 'Light List Accent 1'
    table_compte.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_compte = table_compte.rows[0].cells
    headers_compte = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_compte):
        hdr_cells_compte[i].text = header
        for paragraph in hdr_cells_compte[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_compte[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_compte['table_data']:
        row_cells_compte = table_compte.add_row().cells
        row_cells_compte[0].text = row.get("Description", "")
        row_cells_compte[1].text = row.get("Année 1", "")
        row_cells_compte[2].text = row.get("Année 2", "")
        row_cells_compte[3].text = row.get("Année 3", "")
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    ### 5. Ajouter la section Soldes Intermédiaires de Gestion ###
    doc.add_heading('Soldes intermédiaires de gestion', level=1)
    
    # Créer le tableau Soldes intermédiaires de gestion dans Word
    table_soldes = doc.add_table(rows=1, cols=4)
    table_soldes.style = 'Light List Accent 1'
    table_soldes.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_soldes = table_soldes.rows[0].cells
    headers_soldes = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_soldes):
        hdr_cells_soldes[i].text = header
        for paragraph in hdr_cells_soldes[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_soldes[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_soldes['table_data']:
        row_cells_soldes = table_soldes.add_row().cells
        row_cells_soldes[0].text = row.get("Description", "")
        row_cells_soldes[1].text = f"{row.get('Année 1', 0):,.2f} €"
        row_cells_soldes[2].text = f"{row.get('Année 2', 0):,.2f} €"
        row_cells_soldes[3].text = f"{row.get('Année 3', 0):,.2f} €"
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    ### 6. Ajouter la section Capacité d'Autofinancement ###
    doc.add_heading('Capacité d\'autofinancement', level=1)
    
    # Créer le tableau Capacité d'Autofinancement dans Word
    table_cap = doc.add_table(rows=1, cols=4)
    table_cap.style = 'Light List Accent 1'
    table_cap.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_cap = table_cap.rows[0].cells
    headers_cap = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_cap):
        hdr_cells_cap[i].text = header
        for paragraph in hdr_cells_cap[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_cap[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_capacite['table_data']:
        row_cells_cap = table_cap.add_row().cells
        row_cells_cap[0].text = row.get("Description", "")
        row_cells_cap[1].text = f"{row.get('Année 1', 0):,.2f} €"
        row_cells_cap[2].text = f"{row.get('Année 2', 0):,.2f} €"
        row_cells_cap[3].text = f"{row.get('Année 3', 0):,.2f} €"
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    ### 7. Ajouter la section Seuil de Rentabilité Économique ###
    doc.add_heading('Seuil de rentabilité économique', level=1)
    
    # Créer le tableau Seuil de Rentabilité Économique dans Word
    table_seuil = doc.add_table(rows=1, cols=4)
    table_seuil.style = 'Light List Accent 1'
    table_seuil.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_seuil = table_seuil.rows[0].cells
    headers_seuil = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_seuil):
        hdr_cells_seuil[i].text = header
        for paragraph in hdr_cells_seuil[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_seuil[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_seuil['table_data']:
        row_cells_seuil = table_seuil.add_row().cells
        row_cells_seuil[0].text = row.get("Description", "")
        row_cells_seuil[1].text = f"{row.get('Année 1', 0):,.0f} €"
        row_cells_seuil[2].text = f"{row.get('Année 2', 0):,.0f} €"
        row_cells_seuil[3].text = f"{row.get('Année 3', 0):,.0f} €"
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    ### 8. Ajouter la section Besoin en Fonds de Roulement ###
    doc.add_heading('Besoin en fonds de roulement', level=1)
    
    # Créer le tableau Besoin en Fonds de Roulement dans Word
    table_bfr = doc.add_table(rows=1, cols=5)
    table_bfr.style = 'Light List Accent 1'
    table_bfr.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_bfr = table_bfr.rows[0].cells
    headers_bfr = ["Analyse clients / fournisseurs", "Délai jours", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_bfr):
        hdr_cells_bfr[i].text = header
        for paragraph in hdr_cells_bfr[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_bfr[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_bfr['table_data']:
        row_cells_bfr = table_bfr.add_row().cells
        row_cells_bfr[0].text = row.get("Analyse clients / fournisseurs", "")
        row_cells_bfr[1].text = row.get("Délai jours", "")
        row_cells_bfr[2].text = row.get("Année 1", "")
        row_cells_bfr[3].text = row.get("Année 2", "")
        row_cells_bfr[4].text = row.get("Année 3", "")
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    # Enregistrer le document dans un buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Bouton de téléchargement
    st.download_button(
        label="Télécharger le Document Word Complet",
        data=buffer,
        file_name="document_complet_financier.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    
    # Message de confirmation
    st.success("Le document Word complet a été généré avec succès !")

def page_besoin_fonds_roulement():
    st.title("Besoin en fonds de roulement")
    
    # Vérifier si les données sont présentes dans la session
    if "data" not in st.session_state:
        st.error("Les données ne sont pas initialisées. Veuillez initialiser la session.")
        return
    
    data = st.session_state["data"]
    
    # Récupérer les informations du projet
    info_generales = data.get("informations_generales", {})
    projet = info_generales.get("intitule_projet", "N/A")
    porteur_projet = info_generales.get("prenom_nom", "N/A")
    
    # Afficher les informations du projet
    st.write(f"**Projet :** {projet}")
    st.write(f"**Porteur de projet :** {porteur_projet}")
    
    st.write("---")

    # Récupérer les délais clients et fournisseurs depuis "besoin_fonds_roulement"
    besoin_fonds = data.get("fonds_roulement", {})
    delai_clients = besoin_fonds.get("duree_credits_clients", 0)  # Durée moyenne des crédits accordés aux clients en jours
    delai_fournisseurs = besoin_fonds.get("duree_dettes_fournisseurs", 0)  # Durée moyenne des crédits accordés aux fournisseurs en jours

    st.write("---")
    
    # Récupérer "Ventes + Production réelle" et "Achats consommés" depuis "soldes_intermediaires_de_gestion"
    soldes_intermediaires = data.get("soldes_intermediaires_de_gestion", {})
    ventes_production_reelle = soldes_intermediaires.get("ventes_production_reelle", [0.0, 0.0, 0.0])
    achats_consommes = soldes_intermediaires.get("achats_consommes", [0.0, 0.0, 0.0])
    
    # Vérifier si les délais sont non nuls
    if delai_clients == 0 or delai_fournisseurs == 0:
        st.error("Les délais de paiement des clients et des fournisseurs ne sont pas renseignés. Veuillez les saisir dans la section 'Votre besoin en fonds de roulement'.")
        return
    
    # Vérifier si les listes ont 3 éléments
    if not (len(ventes_production_reelle) == len(achats_consommes) == 3):
        st.error("Les listes de 'Ventes + Production réelle' ou 'Achats consommés' ne contiennent pas exactement 3 éléments. Veuillez vérifier les données.")
        return
    
    # Calculer le Volume crédit client HT = Ventes + Production réelle / (delai_jours * 365)
    volume_credit_client_ht = []
    for i in range(3):
        vcc_ht = (ventes_production_reelle[i] * delai_clients) / 365
        volume_credit_client_ht.append(vcc_ht)
    
    # Calculer le Volume dettes fournisseurs HT = Achats consommés / (delai_jours * 365)
    volume_dettes_fournisseurs_ht = []
    for i in range(3):
        vdf_ht = (achats_consommes[i] * delai_fournisseurs) / 365
        volume_dettes_fournisseurs_ht.append(vdf_ht)
    
    # Calculer le Besoin en fonds de roulement (BFR) = Volume crédit client HT - Volume dettes fournisseurs HT
    bfr = [volume_credit_client_ht[i] - volume_dettes_fournisseurs_ht[i] for i in range(3)]
    
    # Afficher les résultats intermédiaires
    st.write("### Résultats des Calculs")
    st.write(f"**Volume crédit client HT Année 1** : {volume_credit_client_ht[0]:.2f} €")
    st.write(f"**Volume dettes fournisseurs HT Année 1** : {volume_dettes_fournisseurs_ht[0]:.2f} €")
    st.write(f"**Besoin en fonds de roulement Année 1** : {bfr[0]:.2f} €")
    st.write(f"**Volume crédit client HT Année 2** : {volume_credit_client_ht[1]:.2f} €")
    st.write(f"**Volume dettes fournisseurs HT Année 2** : {volume_dettes_fournisseurs_ht[1]:.2f} €")
    st.write(f"**Besoin en fonds de roulement Année 2** : {bfr[1]:.2f} €")
    st.write(f"**Volume crédit client HT Année 3** : {volume_credit_client_ht[2]:.2f} €")
    st.write(f"**Volume dettes fournisseurs HT Année 3** : {volume_dettes_fournisseurs_ht[2]:.2f} €")
    st.write(f"**Besoin en fonds de roulement Année 3** : {bfr[2]:.2f} €")
    
    # Préparer les données pour le tableau
    data_table = {
        "Analyse clients / fournisseurs": [
            "Besoins",
            "Volume crédit client HT",
            "Ressources",
            "Volume dettes fournisseurs HT",
            "Besoin en fonds de roulement"
        ],
        "Délai jours": [
            "",
            f"{delai_clients}",
            "",
            f"{delai_fournisseurs}",
            ""
        ],
        "Année 1": [
            "",
            f"{volume_credit_client_ht[0]:.2f} €",
            "",
            f"{volume_dettes_fournisseurs_ht[0]:.2f} €",
            f"{bfr[0]:.2f} €"
        ],
        "Année 2": [
            "",
            f"{volume_credit_client_ht[1]:.2f} €",
            "",
            f"{volume_dettes_fournisseurs_ht[1]:.2f} €",
            f"{bfr[1]:.2f} €"
        ],
        "Année 3": [
            "",
            f"{volume_credit_client_ht[2]:.2f} €",
            "",
            f"{volume_dettes_fournisseurs_ht[2]:.2f} €",
            f"{bfr[2]:.2f} €"
        ]
    }
    
    df = pd.DataFrame(data_table)
    
    # Afficher le tableau
    st.write("### Tableau du Besoin en fonds de roulement")
    st.table(df)
    
    # Stocker les résultats dans les données
    data["besoin_fonds_roulement"] = {
        "delai_clients": delai_clients,
        "delai_fournisseurs": delai_fournisseurs,
        "volume_credit_client_ht": volume_credit_client_ht,
        "volume_dettes_fournisseurs_ht": volume_dettes_fournisseurs_ht,
        "bfr": bfr
    }
    
    # Enregistrer les données dans la session
    st.session_state["data"] = data
    
    # Préparer les données d'exportation pour Besoin en fonds de roulement
    export_table_bfr = []
    for idx, label in enumerate(data_table["Analyse clients / fournisseurs"]):
        export_table_bfr.append({
            "Analyse clients / fournisseurs": label,
            "Délai jours": data_table["Délai jours"][idx],
            "Année 1": data_table["Année 1"][idx],
            "Année 2": data_table["Année 2"][idx],
            "Année 3": data_table["Année 3"][idx]
        })
    
    # Stocker les données d'exportation dans la session
    st.session_state['export_data_besoin_fonds_roulement'] = {
        "projet": projet,
        "porteur_projet": porteur_projet,
        "table_data": export_table_bfr
    }
    
    # Section Export
    st.header("Exporter les données")
    
    # Bouton pour télécharger le fichier Word complet contenant tous les tableaux avec une clé unique
    if st.button("Télécharger le Document Word Complet", key="download_word_complet_besoin_fonds_roulement"):
        telecharger_document_complet()


def telecharger_document_complet():
    # Récupérer les données exportées de toutes les sections
    export_data_investissements = st.session_state.get('export_data_investissements', {})
    export_data_salaires = st.session_state.get('export_data_salaires_charges_sociales', {})
    export_data_amortissements = st.session_state.get('export_data_detail_amortissements', {})
    export_data_compte = st.session_state.get('export_data_compte_resultats_previsionnel', {})
    export_data_soldes = st.session_state.get('export_data_soldes_intermediaires_de_gestion', {})
    export_data_capacite = st.session_state.get('export_data_capacite_autofinancement', {})
    export_data_seuil = st.session_state.get('export_data_seuil_rentabilite_economique', {})
    export_data_bfr = st.session_state.get('export_data_besoin_fonds_roulement', {})
    export_data_plan_financement = st.session_state.get('export_data_plan_financement_trois_ans', {})
    
    # Vérifiez que toutes les données sont présentes
    if not all([
        export_data_investissements.get("table_data"),
        export_data_salaires.get("table_data"),
        export_data_amortissements.get("amortissements"),
        export_data_compte.get("table_data"),
        export_data_soldes.get("table_data"),
        export_data_capacite.get("table_data"),
        export_data_seuil.get("table_data"),
        export_data_bfr.get("table_data"),
        export_data_plan_financement.get("table_data")
    ]):
        st.error("Toutes les sections doivent être remplies avant de télécharger le document complet.")
        return
    
    # Créer un document Word
    doc = Document()
    
    ### 1. Ajouter la section Investissements et Financements ###
    doc.add_heading('Investissements et Financements', level=1)
    doc.add_paragraph(f"**Projet :** {export_data_investissements.get('projet', 'N/A')}")
    doc.add_paragraph(f"**Porteur de projet :** {export_data_investissements.get('porteur_projet', 'N/A')}")
    doc.add_page_break()
    
    # Créer le tableau Investissements et Financements dans Word
    table_inv = doc.add_table(rows=1, cols=4)
    table_inv.style = 'Light List Accent 1'
    table_inv.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_inv = table_inv.rows[0].cells
    headers_inv = ["Investissements", "Taux (%)", "Durée (mois)", "Montant (€)"]
    for i, header in enumerate(headers_inv):
        hdr_cells_inv[i].text = header
        for paragraph in hdr_cells_inv[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_inv[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_investissements['table_data']:
        row_cells = table_inv.add_row().cells
        row_cells[0].text = row.get("Investissements", "")
        row_cells[1].text = row.get("Taux (%)", "")
        row_cells[2].text = str(row.get("Durée (mois)", "")) if row.get("Durée (mois)", "") != "-" else "-"
        row_cells[3].text = row.get("Montant (€)", "")
        
        # Mise en forme des lignes spécifiques
        if row["Investissements"] in ["INVESTISSEMENTS", "FINANCEMENT DES INVESTISSEMENTS", "TOTAL SUBVENTIONS", "TOTAL EMPRUNTS"]:
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        elif "TOTAL" in row["Investissements"]:
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        
        # Alignement des cellules
        row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    ### 2. Ajouter la section Salaires et Charges Sociales ###
    doc.add_heading('Salaires et Charges Sociales', level=1)
    doc.add_paragraph(f"**Projet :** {export_data_salaires.get('projet', 'N/A')}")
    doc.add_paragraph(f"**Porteur de projet :** {export_data_salaires.get('porteur_projet', 'N/A')}")
    doc.add_paragraph("---")
    
    # Créer le tableau Salaires et Charges Sociales dans Word
    table_sal = doc.add_table(rows=1, cols=4)
    table_sal.style = 'Light List Accent 1'
    table_sal.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_sal = table_sal.rows[0].cells
    headers_sal = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_sal):
        hdr_cells_sal[i].text = header
        for paragraph in hdr_cells_sal[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_sal[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_salaires['table_data']:
        row_cells = table_sal.add_row().cells
        row_cells[0].text = row.get("Description", "")
        row_cells[1].text = row.get("Année 1", "")
        row_cells[2].text = row.get("Année 2", "")
        row_cells[3].text = row.get("Année 3", "")
        
        # Alignement des cellules
        row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    ### 3. Ajouter la section Détail des Amortissements ###
    doc.add_heading('Détail des Amortissements', level=1)
    
    # Créer le tableau Détail des Amortissements dans Word
    table_amort = doc.add_table(rows=1, cols=4)
    table_amort.style = 'Light List Accent 1'
    table_amort.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_amort = table_amort.rows[0].cells
    headers_amort = ["Amortissement", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_amort):
        hdr_cells_amort[i].text = header
        for paragraph in hdr_cells_amort[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_amort[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Ajouter les données Amortissements au tableau
    for row in export_data_amortissements['amortissements']:
        row_cells_amort = table_amort.add_row().cells
        row_cells_amort[0].text = row.get("Amortissement", "")
        row_cells_amort[1].text = row.get("Année 1", "")
        row_cells_amort[2].text = row.get("Année 2", "")
        row_cells_amort[3].text = row.get("Année 3", "")
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les amortissements sont calculés en fonction de la durée d'amortissement spécifiée.")
    
    ### 4. Ajouter la section Compte de Résultats Prévisionnel ###
    doc.add_heading('Compte de Résultats Prévisionnel', level=1)
    
    # Créer le tableau Compte de Résultats Prévisionnel dans Word
    table_compte = doc.add_table(rows=1, cols=4)
    table_compte.style = 'Light List Accent 1'
    table_compte.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_compte = table_compte.rows[0].cells
    headers_compte = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_compte):
        hdr_cells_compte[i].text = header
        for paragraph in hdr_cells_compte[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_compte[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_compte['table_data']:
        row_cells_compte = table_compte.add_row().cells
        row_cells_compte[0].text = row.get("Description", "")
        row_cells_compte[1].text = row.get("Année 1", "")
        row_cells_compte[2].text = row.get("Année 2", "")
        row_cells_compte[3].text = row.get("Année 3", "")
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    ### 5. Ajouter la section Soldes Intermédiaires de Gestion ###
    doc.add_heading('Soldes intermédiaires de gestion', level=1)
    
    # Créer le tableau Soldes intermédiaires de gestion dans Word
    table_soldes = doc.add_table(rows=1, cols=4)
    table_soldes.style = 'Light List Accent 1'
    table_soldes.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_soldes = table_soldes.rows[0].cells
    headers_soldes = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_soldes):
        hdr_cells_soldes[i].text = header
        for paragraph in hdr_cells_soldes[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_soldes[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_soldes['table_data']:
        row_cells_soldes = table_soldes.add_row().cells
        row_cells_soldes[0].text = row.get("Description", "")
        row_cells_soldes[1].text = f"{row.get('Année 1', 0):,.2f} €"
        row_cells_soldes[2].text = f"{row.get('Année 2', 0):,.2f} €"
        row_cells_soldes[3].text = f"{row.get('Année 3', 0):,.2f} €"
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    ### 6. Ajouter la section Capacité d'Autofinancement ###
    doc.add_heading('Capacité d\'autofinancement', level=1)
    
    # Créer le tableau Capacité d'Autofinancement dans Word
    table_cap = doc.add_table(rows=1, cols=4)
    table_cap.style = 'Light List Accent 1'
    table_cap.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_cap = table_cap.rows[0].cells
    headers_cap = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_cap):
        hdr_cells_cap[i].text = header
        for paragraph in hdr_cells_cap[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_cap[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_capacite['table_data']:
        row_cells_cap = table_cap.add_row().cells
        row_cells_cap[0].text = row.get("Description", "")
        row_cells_cap[1].text = f"{row.get('Année 1', 0):,.2f} €"
        row_cells_cap[2].text = f"{row.get('Année 2', 0):,.2f} €"
        row_cells_cap[3].text = f"{row.get('Année 3', 0):,.2f} €"
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    ### 7. Ajouter la section Seuil de Rentabilité Économique ###
    doc.add_heading('Seuil de rentabilité économique', level=1)
    
    # Créer le tableau Seuil de Rentabilité Économique dans Word
    table_seuil = doc.add_table(rows=1, cols=4)
    table_seuil.style = 'Light List Accent 1'
    table_seuil.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_seuil = table_seuil.rows[0].cells
    headers_seuil = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_seuil):
        hdr_cells_seuil[i].text = header
        for paragraph in hdr_cells_seuil[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_seuil[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_seuil['table_data']:
        row_cells_seuil = table_seuil.add_row().cells
        row_cells_seuil[0].text = row.get("Description", "")
        row_cells_seuil[1].text = f"{row.get('Année 1', 0):,.0f} €"
        row_cells_seuil[2].text = f"{row.get('Année 2', 0):,.0f} €"
        row_cells_seuil[3].text = f"{row.get('Année 3', 0):,.0f} €"
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    ### 8. Ajouter la section Plan de Financement à Trois Ans ###
    doc.add_heading('Plan de financement à trois ans', level=1)
    
    # Créer le tableau Plan de Financement à Trois Ans dans Word
    table_plan = doc.add_table(rows=1, cols=5)
    table_plan.style = 'Light List Accent 1'
    table_plan.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_plan = table_plan.rows[0].cells
    headers_plan = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_plan):
        hdr_cells_plan[i].text = header
        for paragraph in hdr_cells_plan[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_plan[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_plan_financement['table_data']:
        row_cells_plan = table_plan.add_row().cells
        row_cells_plan[0].text = row.get("Plan de financement à trois ans", "")
        row_cells_plan[1].text = row.get("Année 1", "")
        row_cells_plan[2].text = row.get("Année 2", "")
        row_cells_plan[3].text = row.get("Année 3", "")
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    # Enregistrer le document dans un buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Bouton de téléchargement
    st.download_button(
        label="Télécharger le Document Word Complet",
        data=buffer,
        file_name="document_complet_financier.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    
    # Message de confirmation
    st.success("Le document Word complet a été généré avec succès !")




def page_plan_financement_trois_ans(): 
    st.title("Plan de financement à trois ans")
    
    # Vérifier si les données sont présentes dans la session
    if "data" not in st.session_state:
        st.error("Les données ne sont pas initialisées. Veuillez initialiser la session.")
        return
    
    data = st.session_state["data"]
    
    # Récupérer les informations du projet
    info_generales = data.get("informations_generales", {})
    projet = info_generales.get("intitule_projet", "N/A")
    porteur_projet = info_generales.get("prenom_nom", "N/A")
    
    # Afficher les informations du projet
    st.write(f"**Projet :** {projet}")
    st.write(f"**Porteur de projet :** {porteur_projet}")
    
    st.write("---")
    
    # Récupérer les besoins démarrage
    besoins_demarrage = data.get("besoins_demarrage", {})
    
    # Calcul des Immobilisations incorporelles et corporelles
    # Définissez quels éléments de "besoins_demarrage" correspondent à chaque catégorie
    immobilisations_inc = sum([
        besoins_demarrage.get("Frais d’établissement", 0),
        besoins_demarrage.get("Frais d’ouverture de compteurs", 0),
        besoins_demarrage.get("Logiciels, formations", 0),
        besoins_demarrage.get("Dépôt de marque", 0),
        besoins_demarrage.get("Droits d’entrée", 0),
        besoins_demarrage.get("Achat fonds de commerce ou parts", 0),
        besoins_demarrage.get("Droit au bail", 0),
        besoins_demarrage.get("Caution ou dépôt de garantie", 0),
        besoins_demarrage.get("Frais de dossier", 0),
        besoins_demarrage.get("Frais de notaire", 0),
        besoins_demarrage.get("Enseigne et éléments de communication", 0)
    ])
    
    immobilisations_corp = sum([
        besoins_demarrage.get("Enseigne et éléments de communication", 0),
        besoins_demarrage.get("Véhicule", 0),
        besoins_demarrage.get("Matériel professionnel", 0),
        besoins_demarrage.get("Matériel autre", 0),
        besoins_demarrage.get("Matériel de bureau", 0)
    ])
    
    immobilisations = [
        immobilisations_inc + immobilisations_corp,  # Année 1
        0.0,  # Année 2
        0.0   # Année 3
    ]
    
    # Acquisition des stocks
    acquisition_stocks = [
        besoins_demarrage.get("Stock de matières et produits", 0),
        0.0,  # Année 2
        0.0   # Année 3
    ]
    
    # Variation du Besoin en fonds de roulement (BFR)
    besoin_fonds = data.get("besoin_fonds_roulement", {})
    bfr = besoin_fonds.get("bfr", [0.0, 0.0, 0.0])
    
    # Variation BFR = BFR année n - BFR année n-1
    variation_bfr = [
        bfr[0],                    # Variation en année 1 (BFR année 1 - BFR année 0)
        bfr[1] - bfr[0],           # Variation en année 2
        bfr[2] - bfr[1]            # Variation en année 3
    ]
    
    # Remboursement d'emprunts
    capacite_autofinancement = data.get("capacite_autofinancement", {})
    remboursements_emprunts = capacite_autofinancement.get("remboursements_emprunts", [0.0, 0.0, 0.0])
    
    # Total des besoins
    total_besoins = [
        immobilisations[0] + acquisition_stocks[0] + variation_bfr[0] + remboursements_emprunts[0],
        immobilisations[1] + acquisition_stocks[1] + variation_bfr[1] + remboursements_emprunts[1],
        immobilisations[2] + acquisition_stocks[2] + variation_bfr[2] + remboursements_emprunts[2]
    ]
    
    # Apport personnel
    financements = data.get("financements", {})
    apport_personnel = financements.get("Apport personnel ou familial", 0.0)
    apports_nature = financements.get("Apports en nature (en valeur)", 0.0)
    apport_total = apport_personnel + apports_nature
    apport_personnel_list = [apport_total, 0.0, 0.0] 
    
    # Emprunts
    pret_1 = financements.get("Prêt 1", {}).get("montant", 0.0)
    pret_2 = financements.get("Prêt 2", {}).get("montant", 0.0)
    pret_3 = financements.get("Prêt 3", {}).get("montant", 0.0)
    total_emprunts = pret_1 + pret_2 + pret_3
    emprunts = [total_emprunts, 0.0, 0.0]  # Supposons que les emprunts sont en année 1
    
    # Subventions
    subvention_1 = financements.get("Subvention 1", {}).get("montant", 0.0)
    subvention_2 = financements.get("Subvention 2", {}).get("montant", 0.0)
    subventions = subvention_1 + subvention_2
    subventions_list = [subventions, 0.0, 0.0]  # Supposons que les subventions sont en année 1
    
    # Autres financements
    autres_financements = financements.get("Autre financement", 0.0)
    autres_financements_list = [autres_financements, 0.0, 0.0]  # Supposons que c'est en année 1
    
    # Capacité d'auto-financement
    capacite_autofinancement_values = capacite_autofinancement.get("capacite_autofinancement", [0.0, 0.0, 0.0])
    
    # Total des ressources
    total_ressources = [
        apport_personnel_list[0] + emprunts[0] + subventions_list[0] + autres_financements_list[0] + capacite_autofinancement_values[0],
        apport_personnel_list[1] + emprunts[1] + subventions_list[1] + autres_financements_list[1] + capacite_autofinancement_values[1],
        apport_personnel_list[2] + emprunts[2] + subventions_list[2] + autres_financements_list[2] + capacite_autofinancement_values[2]
    ]
    
    # Variation de trésorerie
    variation_tresorerie = [
        total_ressources[0] - total_besoins[0],
        total_ressources[1] - total_besoins[1],
        total_ressources[2] - total_besoins[2]
    ]
    
    # Excédent de trésorerie (cumulatif)
    excedent_tresorerie = []
    cumul_excedent = 0.0
    for i in range(3):
        cumul_excedent += variation_tresorerie[i]
        excedent_tresorerie.append(cumul_excedent)
    
    # Préparation des données pour le tableau
    data_table = {
        "Plan de financement à trois ans": [
            "Immobilisations",
            "Acquisition des stocks",
            "Variation du Besoin en fonds de roulement",
            "Remboursement d'emprunts",
            "Total des besoins",
            "Apport personnel",
            "Emprunts",
            "Subventions",
            "Autres financements",
            "Capacité d'auto-financement",
            "Total des ressources",
            "Variation de trésorerie",
            "Excédent de trésorerie"
        ],
        "Année 1": [
            f"{immobilisations[0]:,.2f} €",
            f"{acquisition_stocks[0]:,.2f} €",
            f"{variation_bfr[0]:,.2f} €",
            f"{remboursements_emprunts[0]:,.2f} €",
            f"{total_besoins[0]:,.2f} €",
            f"{apport_personnel_list[0]:,.2f} €",
            f"{emprunts[0]:,.2f} €",
            f"{subventions_list[0]:,.2f} €",
            f"{autres_financements_list[0]:,.2f} €",
            f"{capacite_autofinancement_values[0]:,.2f} €",
            f"{total_ressources[0]:,.2f} €",
            f"{variation_tresorerie[0]:,.2f} €",
            f"{excedent_tresorerie[0]:,.2f} €"
        ],
        "Année 2": [
            f"{immobilisations[1]:,.2f} €",
            f"{acquisition_stocks[1]:,.2f} €",
            f"{variation_bfr[1]:,.2f} €",
            f"{remboursements_emprunts[1]:,.2f} €",
            f"{total_besoins[1]:,.2f} €",
            f"{apport_personnel_list[1]:,.2f} €",
            f"{emprunts[1]:,.2f} €",
            f"{subventions_list[1]:,.2f} €",
            f"{autres_financements_list[1]:,.2f} €",
            f"{capacite_autofinancement_values[1]:,.2f} €",
            f"{total_ressources[1]:,.2f} €",
            f"{variation_tresorerie[1]:,.2f} €",
            f"{excedent_tresorerie[1]:,.2f} €"
        ],
        "Année 3": [
            f"{immobilisations[2]:,.2f} €",
            f"{acquisition_stocks[2]:,.2f} €",
            f"{variation_bfr[2]:,.2f} €",
            f"{remboursements_emprunts[2]:,.2f} €",
            f"{total_besoins[2]:,.2f} €",
            f"{apport_personnel_list[2]:,.2f} €",
            f"{emprunts[2]:,.2f} €",
            f"{subventions_list[2]:,.2f} €",
            f"{autres_financements_list[2]:,.2f} €",
            f"{capacite_autofinancement_values[2]:,.2f} €",
            f"{total_ressources[2]:,.2f} €",
            f"{variation_tresorerie[2]:,.2f} €",
            f"{excedent_tresorerie[2]:,.2f} €"
        ]
    }
    
    df = pd.DataFrame(data_table)
    st.write("### Tableau du Plan de financement à trois ans")
    st.table(df)
    
    # Stocker les résultats dans les données
    data["plan_financement"] = {
        "immobilisations": immobilisations,
        "acquisition_stocks": acquisition_stocks,
        "variation_bfr": variation_bfr,
        "remboursements_emprunts": remboursements_emprunts,
        "total_besoins": total_besoins,
        "apport_personnel": apport_personnel_list,
        "emprunts": emprunts,
        "subventions": subventions_list,
        "autres_financements": autres_financements_list,
        "capacite_autofinancement": capacite_autofinancement_values,
        "total_ressources": total_ressources,
        "variation_tresorerie": variation_tresorerie,
        "excedent_tresorerie": excedent_tresorerie
    }
    
    # Enregistrer les données dans la session
    st.session_state["data"] = data   
    
    # Préparer les données d'exportation pour Plan de Financement à Trois Ans
    export_table_plan_financement = []
    for idx, label in enumerate(data_table["Plan de financement à trois ans"]):
        export_table_plan_financement.append({
            "Plan de financement à trois ans": label,
            "Année 1": data_table["Année 1"][idx],
            "Année 2": data_table["Année 2"][idx],
            "Année 3": data_table["Année 3"][idx]
        })
    
    # Stocker les données d'exportation dans la session
    st.session_state['export_data_plan_financement_trois_ans'] = {
        "projet": projet,
        "porteur_projet": porteur_projet,
        "table_data": export_table_plan_financement
    }
    
    # Section Export
    st.header("Exporter les données")
    
    # Bouton pour télécharger le fichier Word complet contenant tous les tableaux avec une clé unique
    if st.button("Télécharger le Document Word Complet", key="download_word_complet_plan_financement_trois_ans"):
        telecharger_document_complet()
        


def telecharger_document_complet():
    # Récupérer les données exportées de toutes les sections
    export_data_investissements = st.session_state.get('export_data_investissements', {})
    export_data_salaires = st.session_state.get('export_data_salaires_charges_sociales', {})
    export_data_amortissements = st.session_state.get('export_data_detail_amortissements', {})
    export_data_compte = st.session_state.get('export_data_compte_resultats_previsionnel', {})
    export_data_soldes = st.session_state.get('export_data_soldes_intermediaires_de_gestion', {})
    export_data_capacite = st.session_state.get('export_data_capacite_autofinancement', {})
    export_data_seuil = st.session_state.get('export_data_seuil_rentabilite_economique', {})
    export_data_bfr = st.session_state.get('export_data_besoin_fonds_roulement', {})
    export_data_plan_financement = st.session_state.get('export_data_plan_financement_trois_ans', {})
    export_data_budget_tresorerie_part1 = st.session_state.get('export_data_budget_previsionnel_tresorerie_part1', {})
    export_data_budget_tresorerie_part2 = st.session_state.get('export_data_budget_previsionnel_tresorerie_part2', {})
    
    # Vérifiez que toutes les données nécessaires sont présentes
    if not all([
        export_data_investissements.get("table_data"),
        export_data_salaires.get("table_data"),
        export_data_amortissements.get("amortissements"),
        export_data_compte.get("table_data"),
        export_data_soldes.get("table_data"),
        export_data_capacite.get("table_data"),
        export_data_seuil.get("table_data"),
        export_data_bfr.get("table_data"),
        export_data_plan_financement.get("table_data"),
        export_data_budget_tresorerie_part1.get("table_data"),
        export_data_budget_tresorerie_part2.get("table_data")
    ]):
        st.error("Toutes les sections doivent être remplies avant de télécharger le document complet.")
        return
    
    # Créer un document Word
    doc = Document()
    
    ### 1. Ajouter la section Investissements et Financements ###
    doc.add_heading('Investissements et Financements', level=1)
    doc.add_paragraph(f"**Projet :** {export_data_investissements.get('projet', 'N/A')}")
    doc.add_paragraph(f"**Porteur de projet :** {export_data_investissements.get('porteur_projet', 'N/A')}")
    
    # Créer le tableau Investissements et Financements dans Word
    table_inv = doc.add_table(rows=1, cols=4)
    table_inv.style = 'Light List Accent 1'
    table_inv.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_inv = table_inv.rows[0].cells
    headers_inv = ["Investissements", "Taux (%)", "Durée (mois)", "Montant (€)"]
    for i, header in enumerate(headers_inv):
        hdr_cells_inv[i].text = header
        for paragraph in hdr_cells_inv[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_inv[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_investissements['table_data']:
        row_cells = table_inv.add_row().cells
        row_cells[0].text = row.get("Investissements", "")
        row_cells[1].text = row.get("Taux (%)", "")
        row_cells[2].text = str(row.get("Durée (mois)", "")) if row.get("Durée (mois)", "") != "-" else "-"
        row_cells[3].text = row.get("Montant (€)", "")
        
        # Mise en forme des lignes spécifiques
        if row["Investissements"] in ["INVESTISSEMENTS", "FINANCEMENT DES INVESTISSEMENTS", "TOTAL SUBVENTIONS", "TOTAL EMPRUNTS"]:
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        elif "TOTAL" in row["Investissements"]:
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        
        # Alignement des cellules
        row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    ### 2. Ajouter la section Salaires et Charges Sociales ###
    doc.add_heading('Salaires et Charges Sociales', level=1)
    doc.add_paragraph(f"**Projet :** {export_data_salaires.get('projet', 'N/A')}")
    doc.add_paragraph(f"**Porteur de projet :** {export_data_salaires.get('porteur_projet', 'N/A')}")
    doc.add_paragraph("---")
    
    # Créer le tableau Salaires et Charges Sociales dans Word
    table_sal = doc.add_table(rows=1, cols=4)
    table_sal.style = 'Light List Accent 1'
    table_sal.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_sal = table_sal.rows[0].cells
    headers_sal = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_sal):
        hdr_cells_sal[i].text = header
        for paragraph in hdr_cells_sal[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_sal[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_salaires['table_data']:
        row_cells = table_sal.add_row().cells
        row_cells[0].text = row.get("Description", "")
        row_cells[1].text = row.get("Année 1", "")
        row_cells[2].text = row.get("Année 2", "")
        row_cells[3].text = row.get("Année 3", "")
        
        # Alignement des cellules
        row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        row_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    ### 3. Ajouter la section Détail des Amortissements ###
    doc.add_heading('Détail des Amortissements', level=1)
    
    # Créer le tableau Détail des Amortissements dans Word
    table_amort = doc.add_table(rows=1, cols=4)
    table_amort.style = 'Light List Accent 1'
    table_amort.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_amort = table_amort.rows[0].cells
    headers_amort = ["Amortissement", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_amort):
        hdr_cells_amort[i].text = header
        for paragraph in hdr_cells_amort[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_amort[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Ajouter les données Amortissements au tableau
    for row in export_data_amortissements['amortissements']:
        row_cells_amort = table_amort.add_row().cells
        row_cells_amort[0].text = row.get("Amortissement", "")
        row_cells_amort[1].text = row.get("Année 1", "")
        row_cells_amort[2].text = row.get("Année 2", "")
        row_cells_amort[3].text = row.get("Année 3", "")
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les amortissements sont calculés en fonction de la durée d'amortissement spécifiée.")
    
    ### 4. Ajouter la section Compte de Résultats Prévisionnel ###
    doc.add_heading('Compte de Résultats Prévisionnel', level=1)
    
    # Créer le tableau Compte de Résultats Prévisionnel dans Word
    table_compte = doc.add_table(rows=1, cols=4)
    table_compte.style = 'Light List Accent 1'
    table_compte.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_compte = table_compte.rows[0].cells
    headers_compte = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_compte):
        hdr_cells_compte[i].text = header
        for paragraph in hdr_cells_compte[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_compte[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_compte['table_data']:
        row_cells_compte = table_compte.add_row().cells
        row_cells_compte[0].text = row.get("Description", "")
        row_cells_compte[1].text = row.get("Année 1", "")
        row_cells_compte[2].text = row.get("Année 2", "")
        row_cells_compte[3].text = row.get("Année 3", "")
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    ### 5. Ajouter la section Soldes Intermédiaires de Gestion ###
    doc.add_heading('Soldes intermédiaires de gestion', level=1)
    
    # Créer le tableau Soldes intermédiaires de gestion dans Word
    table_soldes = doc.add_table(rows=1, cols=4)
    table_soldes.style = 'Light List Accent 1'
    table_soldes.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_soldes = table_soldes.rows[0].cells
    headers_soldes = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_soldes):
        hdr_cells_soldes[i].text = header
        for paragraph in hdr_cells_soldes[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_soldes[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_soldes['table_data']:
        row_cells_soldes = table_soldes.add_row().cells
        row_cells_soldes[0].text = row.get("Description", "")
        row_cells_soldes[1].text = f"{row.get('Année 1', 0):,.2f} €"
        row_cells_soldes[2].text = f"{row.get('Année 2', 0):,.2f} €"
        row_cells_soldes[3].text = f"{row.get('Année 3', 0):,.2f} €"
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    ### 6. Ajouter la section Capacité d'Autofinancement ###
    doc.add_heading('Capacité d\'autofinancement', level=1)
    
    # Créer le tableau Capacité d'Autofinancement dans Word
    table_cap = doc.add_table(rows=1, cols=4)
    table_cap.style = 'Light List Accent 1'
    table_cap.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_cap = table_cap.rows[0].cells
    headers_cap = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_cap):
        hdr_cells_cap[i].text = header
        for paragraph in hdr_cells_cap[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_cap[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_capacite['table_data']:
        row_cells_cap = table_cap.add_row().cells
        row_cells_cap[0].text = row.get("Description", "")
        row_cells_cap[1].text = f"{row.get('Année 1', 0):,.2f} €"
        row_cells_cap[2].text = f"{row.get('Année 2', 0):,.2f} €"
        row_cells_cap[3].text = f"{row.get('Année 3', 0):,.2f} €"
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    ### 7. Ajouter la section Seuil de Rentabilité Économique ###
    doc.add_heading('Seuil de rentabilité économique', level=1)
    
    # Créer le tableau Seuil de Rentabilité Économique dans Word
    table_seuil = doc.add_table(rows=1, cols=4)
    table_seuil.style = 'Light List Accent 1'
    table_seuil.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_seuil = table_seuil.rows[0].cells
    headers_seuil = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_seuil):
        hdr_cells_seuil[i].text = header
        for paragraph in hdr_cells_seuil[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_seuil[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_seuil['table_data']:
        row_cells_seuil = table_seuil.add_row().cells
        row_cells_seuil[0].text = row.get("Description", "")
        row_cells_seuil[1].text = f"{row.get('Année 1', 0):,.0f} €"
        row_cells_seuil[2].text = f"{row.get('Année 2', 0):,.0f} €"
        row_cells_seuil[3].text = f"{row.get('Année 3', 0):,.0f} €"
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    ### 8. Ajouter la section Plan de Financement à Trois Ans ###
    doc.add_heading('Plan de financement à trois ans', level=1)
    
    # Créer le tableau Plan de Financement à Trois Ans dans Word
    table_plan = doc.add_table(rows=1, cols=4)
    table_plan.style = 'Light List Accent 1'
    table_plan.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells_plan = table_plan.rows[0].cells
    headers_plan = ["Description", "Année 1", "Année 2", "Année 3"]
    for i, header in enumerate(headers_plan):
        hdr_cells_plan[i].text = header
        for paragraph in hdr_cells_plan[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_plan[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for row in export_data_plan_financement['table_data']:
        row_cells_plan = table_plan.add_row().cells
        row_cells_plan[0].text = row.get("Description", "")
        row_cells_plan[1].text = row.get("Année 1", "")
        row_cells_plan[2].text = row.get("Année 2", "")
        row_cells_plan[3].text = row.get("Année 3", "")
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les résultats sont calculés selon les données fournies.")
    
    ### 9. Ajouter la section Budget Prévisionnel de Trésorerie Partie 1 ###
    doc.add_heading('Budget prévisionnel de trésorerie - Partie 1', level=1)
    
    # Créer le premier tableau Budget prévisionnel de trésorerie
    table_budget_part1 = doc.add_table(rows=1, cols=len(export_data_budget_tresorerie_part1['table_data'][0]))
    table_budget_part1.style = 'Light List Accent 1'
    table_budget_part1.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Ajouter les en-têtes
    headers_budget_part1 = export_data_budget_tresorerie_part1['table_data'][0].keys()
    hdr_cells_budget_part1 = table_budget_part1.rows[0].cells
    for i, header in enumerate(headers_budget_part1):
        hdr_cells_budget_part1[i].text = header
        for paragraph in hdr_cells_budget_part1[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_budget_part1[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Ajouter les données du premier tableau
    for row in export_data_budget_tresorerie_part1['table_data'][1:]:
        row_cells = table_budget_part1.add_row().cells
        for i, value in enumerate(row.values()):
            row_cells[i].text = value
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les données du budget prévisionnel de trésorerie - Partie 1 sont basées sur les estimations fournies.")
    
    ### 10. Ajouter la section Budget Prévisionnel de Trésorerie Partie 2 ###
    doc.add_heading('Budget prévisionnel de trésorerie - Partie 2', level=1)
    
    # Créer le deuxième tableau Budget prévisionnel de trésorerie
    table_budget_part2 = doc.add_table(rows=1, cols=len(export_data_budget_tresorerie_part2['table_data'][0]))
    table_budget_part2.style = 'Light List Accent 1'
    table_budget_part2.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Ajouter les en-têtes
    headers_budget_part2 = export_data_budget_tresorerie_part2['table_data'][0].keys()
    hdr_cells_budget_part2 = table_budget_part2.rows[0].cells
    for i, header in enumerate(headers_budget_part2):
        hdr_cells_budget_part2[i].text = header
        for paragraph in hdr_cells_budget_part2[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
        hdr_cells_budget_part2[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Ajouter les données du deuxième tableau
    for row in export_data_budget_tresorerie_part2['table_data'][1:]:
        row_cells = table_budget_part2.add_row().cells
        for i, value in enumerate(row.values()):
            row_cells[i].text = value
    
    # Ajouter une note
    doc.add_paragraph()
    doc.add_paragraph("Les données du budget prévisionnel de trésorerie - Partie 2 sont basées sur les estimations fournies.")
    
    # Enregistrer le document dans un buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Bouton de téléchargement
    st.download_button(
        label="Télécharger le Document Word Complet",
        data=buffer,
        file_name="document_complet_financier.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    
    # Message de confirmation
    st.success("Le document Word complet a été généré avec succès !")


import streamlit as st
import pandas as pd

def page_budget_previsionnel_tresorerie():
    st.title("Budget prévisionnel de trésorerie")
    
    data = st.session_state.get("data", {})
    
    if not data:
        st.error("Les données ne sont pas initialisées. Veuillez initialiser la session.")
        return
    
    # Récupérer les informations du projet
    projet = data.get("informations_generales", {}).get("intitule_projet", "N/A")
    porteur_projet = data.get("informations_generales", {}).get("prenom_nom", "N/A")
    
    st.write(f"**Projet :** {projet} &nbsp;&nbsp;&nbsp;&nbsp;&nbsp; **(Hors TVA)**")
    st.write(f"**Porteur de projet :** {porteur_projet}")
    
    st.write("---")
    
    # Création du budget prévisionnel pour la première année, mois par mois
    months = [f"Mois {i+1}" for i in range(12)] + ["TOTAL"]
    
    # Initialisation des structures de données
    encaissements = {}
    decaissements = {}
    solde_precedent = [0.0] * 12
    solde_mois = [0.0] * 12
    solde_tresorerie_cumul = [0.0] * 12
    
    # ----------------------------
    # Encaissements
    # ----------------------------
    
    # Récupérer les encaissements depuis "Plan de financement à trois ans"
    plan_financement = data.get("plan_financement", {})
    apport_personnel = plan_financement.get("apport_personnel", [0.0, 0.0, 0.0])[0]
    emprunts = plan_financement.get("emprunts", [0.0, 0.0, 0.0])[0]
    subventions = plan_financement.get("subventions", [0.0, 0.0, 0.0])[0]
    autres_financements = plan_financement.get("autres_financements", [0.0, 0.0, 0.0])[0]
    
    encaissements["Apport personnel"] = [apport_personnel] + [0.0]*11
    encaissements["Emprunts"] = [emprunts] + [0.0]*11
    encaissements["Subventions"] = [subventions] + [0.0]*11
    encaissements["Autres financements"] = [autres_financements] + [0.0]*11
    
    # Récupérer les ventes depuis "Chiffre d'Affaires Prévisionnel"
    chiffre_affaires = data.get("chiffre_affaires", {})
    
    # Initialiser les listes de ventes mensuelles
    vente_marchandises_mensuel = []
    vente_services_mensuel = []
    
    # Remplir les ventes mensuelles de Marchandises
    for i in range(1, 13):
        key_ca = f"Marchandises_Mois {i}_ca"
        ca = chiffre_affaires.get(key_ca, 0.0)
        vente_marchandises_mensuel.append(ca)
    
    # Remplir les ventes mensuelles de Services
    for i in range(1, 13):
        key_ca = f"Services_Mois {i}_ca"
        ca = chiffre_affaires.get(key_ca, 0.0)
        vente_services_mensuel.append(ca)
    
    encaissements["Vente de marchandises"] = vente_marchandises_mensuel
    encaissements["Vente de services"] = vente_services_mensuel
    encaissements["Chiffre d'affaires (total)"] = [vente_marchandises_mensuel[i] + vente_services_mensuel[i] for i in range(12)]
    
    # Total des encaissements
    total_encaissements = []
    for i in range(12):
        total = (
            encaissements["Apport personnel"][i] +
            encaissements["Emprunts"][i] +
            encaissements["Subventions"][i] +
            encaissements["Autres financements"][i] +
            encaissements["Vente de marchandises"][i] +
            encaissements["Vente de services"][i]
        )
        total_encaissements.append(total)
    total_total_encaissements = sum(total_encaissements)
    total_encaissements.append(total_total_encaissements)
    
    # ----------------------------
    # Décaissements
    # ----------------------------
    
    # Récupérer les données nécessaires pour les décaissements
    besoins_demarrage = data.get("besoins_demarrage", {})
    charges_variables = data.get("charges_variables", {})
    compte_resultat = data.get("compte_de_resultat", {})
    soldes_intermediaires = data.get("soldes_intermediaires_de_gestion", {})
    capacite_autofinancement = data.get("capacite_autofinancement", {})
    
    # Immobilisations incorporelles et corporelles depuis "besoins_demarrage"
    immobilisations_incorporelles = sum([
        besoins_demarrage.get("Frais d’établissement", 0.0),
        besoins_demarrage.get("Frais d’ouverture de compteurs", 0.0),
        besoins_demarrage.get("Logiciels, formations", 0.0),
        besoins_demarrage.get("Dépôt de marque", 0.0),
        besoins_demarrage.get("Droits d’entrée", 0.0),
        besoins_demarrage.get("Achat fonds de commerce ou parts", 0.0),
        besoins_demarrage.get("Droit au bail", 0.0),
        besoins_demarrage.get("Caution ou dépôt de garantie", 0.0),
        besoins_demarrage.get("Frais de dossier", 0.0),
        besoins_demarrage.get("Frais de notaire", 0.0),
    ])
    
    immobilisations_corporelles = sum([
        besoins_demarrage.get("Enseigne et éléments de communication", 0.0),
        besoins_demarrage.get("Véhicule", 0.0),
        besoins_demarrage.get("Matériel professionnel", 0.0),
        besoins_demarrage.get("Matériel autre", 0.0),
        besoins_demarrage.get("Matériel de bureau", 0.0)
    ])
    
    immobilisations_total = immobilisations_incorporelles + immobilisations_corporelles
    
    decaissements["Immobilisations incorporelles"] = [immobilisations_incorporelles] + [0.0]*11
    decaissements["Immobilisations corporelles"] = [immobilisations_corporelles] + [0.0]*11
    decaissements["Immobilisations (total)"] = [immobilisations_total] + [0.0]*11
    
    # Acquisition des stocks depuis "Stock de matières et produits"
    acquisition_stocks = besoins_demarrage.get("Stock de matières et produits", 0.0)
    decaissements["Acquisition stocks"] = [acquisition_stocks] + [0.0]*11
    
    # Échéances emprunt : "Principal année 1" divisé par 12
    remboursements_emprunts = capacite_autofinancement.get("remboursements_emprunts", [0.0, 0.0, 0.0])
    principal_annee1 = remboursements_emprunts[0]
    echeances_emprunt_mensuel = principal_annee1 / 12.0 if principal_annee1 > 0 else 0.0
    decaissements["Échéances emprunt"] = [echeances_emprunt_mensuel] * 12
    
    # Achats de marchandises : "Vente de marchandises" * "le coût d'achat de vos marchandises" de "Charges Variables"
    cout_achat_marchandises_pct = charges_variables.get("cout_achat_marchandises_pct", 100.0)
    if cout_achat_marchandises_pct == 0.0:
        cout_achat_marchandises_pct = 100.0  # Supposer 100% si non renseigné
    
    achats_marchandises_mensuel = [vente_marchandises_mensuel[i] * cout_achat_marchandises_pct / 100.0 for i in range(12)]
    decaissements["Achats de marchandises"] = achats_marchandises_mensuel
    
    # Charges externes : Récupérer depuis "soldes_intermediaires_de_gestion"
    charges_externes_annee1 = soldes_intermediaires.get("charges_externes", [0.0, 0.0, 0.0])[0]
    charges_externes_mensuel = charges_externes_annee1 / 12.0 if charges_externes_annee1 > 0 else 0.0
    decaissements["Charges externes"] = [charges_externes_mensuel] * 12
    
    # Impôts et taxes
    impots_et_taxes_annee1 = compte_resultat.get("impots_et_taxes", [0.0, 0.0, 0.0])[0]
    impots_et_taxes_mensuel = impots_et_taxes_annee1 / 12.0 if impots_et_taxes_annee1 > 0 else 0.0
    decaissements["Impôts et taxes"] = [impots_et_taxes_mensuel] * 12
    
    # Salaires employés, Charges sociales employés, Prélèvement dirigeant(s), Charges sociales dirigeant(s), Frais bancaires, charges financières
    salaires_employes_annee1 = compte_resultat.get("salaires_employes", [0.0, 0.0, 0.0])[0]
    charges_sociales_employes_annee1 = compte_resultat.get("charges_sociales_employes", [0.0, 0.0, 0.0])[0]
    prelevement_dirigeants_annee1 = compte_resultat.get("salaires_dirigeants", [0.0, 0.0, 0.0])[0]
    charges_sociales_dirigeants_annee1 = compte_resultat.get("charges_sociales_dirigeants", [0.0, 0.0, 0.0])[0]
    frais_bancaires_annuels = compte_resultat.get("total_frais_financiers", [0.0, 0.0, 0.0])[0]
    
    salaires_employes_mensuel = [salaires_employes_annee1 / 12.0] * 12
    charges_sociales_employes_mensuel = [charges_sociales_employes_annee1 / 12.0] * 12
    prelevement_dirigeants_mensuel = [prelevement_dirigeants_annee1 / 12.0] * 12
    charges_sociales_dirigeants_mensuel = [charges_sociales_dirigeants_annee1 / 12.0] * 12
    frais_bancaires_mensuel = [frais_bancaires_annuels / 12.0] * 12 if frais_bancaires_annuels > 0 else [0.0] * 12
    
    decaissements["Salaires employés"] = salaires_employes_mensuel
    decaissements["Charges sociales employés"] = charges_sociales_employes_mensuel
    decaissements["Prélèvement dirigeant(s)"] = prelevement_dirigeants_mensuel
    decaissements["Charges sociales dirigeant(s)"] = charges_sociales_dirigeants_mensuel
    decaissements["Frais bancaires, charges financières"] = frais_bancaires_mensuel
    
    # ----------------------------
    # Total charges de personnel
    # ----------------------------
    # Calculer le total des charges de personnel pour chaque mois
    total_charges_personnel_mensuel = [
        salaires_employes_mensuel[i] + charges_sociales_employes_mensuel[i] +
        prelevement_dirigeants_mensuel[i] + charges_sociales_dirigeants_mensuel[i]
        for i in range(12)
    ]
    decaissements["Total charges de personnel"] = total_charges_personnel_mensuel
    
    # ----------------------------
    # Total des décaissements
    # ----------------------------
    
    # Définir les clés à inclure dans le total des décaissements
    decaissements_keys = [
        "Immobilisations (total)",
        "Acquisition stocks",
        "Échéances emprunt",
        "Achats de marchandises",
        "Charges externes",
        "Impôts et taxes",
        "Total charges de personnel",
        "Frais bancaires, charges financières"
    ]
    
    total_decaissements = []
    for i in range(12):
        total = sum([decaissements[key][i] for key in decaissements_keys])
        total_decaissements.append(total)
    total_total_decaissements = sum(total_decaissements)
    total_decaissements.append(total_total_decaissements)
    
    # ----------------------------
    # Calcul des Soldes
    # ----------------------------
    
    for i in range(12):
        solde_mois[i] = total_encaissements[i] - total_decaissements[i]
        solde_tresorerie_cumul[i] = solde_tresorerie_cumul[i - 1] + solde_mois[i] if i > 0 else solde_mois[i]
        solde_precedent[i] = solde_tresorerie_cumul[i - 1] if i > 0 else 0.0
    
    # Append totals to solde_mois, solde_precedent, solde_tresorerie_cumul
    total_solde_mois = sum(solde_mois)
    solde_mois.append(total_solde_mois)
    
    # Pour solde_precedent, le total n'est pas significatif, on peut ajouter une chaîne vide
    solde_precedent.append("")
    
    # Pour solde_tresorerie_cumul, on peut ajouter la dernière valeur cumulative
    solde_tresorerie_cumul.append(solde_tresorerie_cumul[-1])
    
    # ----------------------------
    # Préparation des données pour le tableau
    # ----------------------------
    
    table_data = {"Description": months}
    
    # Encaissements
    for key in encaissements:
        amounts = encaissements[key]
        total = sum(amounts)
        amounts_with_total = amounts + [total]
        table_data[key] = [f"{value:,.2f} €" if value != 0 else "-" for value in amounts_with_total]
    
    # Décaissements
    for key in decaissements:
        # Inclure toutes les lignes de décaissements
        amounts = decaissements[key]
        total = sum(amounts)
        # Remplacer 0 par '-' si nécessaire
        amounts_with_total = [f"{value:,.2f} €" if value != 0 else "-" for value in amounts] + [f"{total:,.2f} €" if total != 0 else "-"]
        table_data[key] = amounts_with_total
    
    # Totaux et soldes
    table_data["Total des encaissements"] = [f"{value:,.2f} €" if value != 0 else "-" for value in total_encaissements]
    table_data["Total des décaissements"] = [f"{value:,.2f} €" if value != 0 else "-" for value in total_decaissements]
    solde_precedent_formatted = [f"{value:,.2f} €" if isinstance(value, (int, float)) and value != 0 else "-" for value in solde_precedent]
    table_data["Solde précédent"] = solde_precedent_formatted
    table_data["Solde du mois"] = [f"{value:,.2f} €" if value != 0 else "-" for value in solde_mois]
    table_data["Solde de trésorerie (cumul)"] = [f"{value:,.2f} €" if value != 0 else "-" for value in solde_tresorerie_cumul]
    
    # Assurer que toutes les listes ont la même longueur
    max_length = max(len(lst) for lst in table_data.values())
    for key in table_data:
        if len(table_data[key]) < max_length:
            difference = max_length - len(table_data[key])
            table_data[key] += [""] * difference  # Remplir avec des chaînes vides si nécessaire
        elif len(table_data[key]) > max_length:
            table_data[key] = table_data[key][:max_length]  # Tronquer si trop long
    
    # Création du DataFrame complet
    df_full = pd.DataFrame(table_data)
    df_full.set_index("Description", inplace=True)
    df_full = df_full.T  # Transposer pour avoir les mois comme colonnes
    
    # Séparation en deux tableaux
    # Tableau 1 : Mois 1 à Mois 5
    columns_part1 = ["Mois 1", "Mois 2", "Mois 3", "Mois 4", "Mois 5"]
    df_part1 = df_full[columns_part1]
    
    # Tableau 2 : Mois 6 à Mois 12 + TOTAL
    columns_part2 = ["Mois 6", "Mois 7", "Mois 8", "Mois 9", "Mois 10", "Mois 11", "Mois 12", "TOTAL"]
    df_part2 = df_full[columns_part2]
    
    ### 3. Ajouter la section Budget Prévisionnel de Trésorerie ###
    # (Les deux tableaux seront ajoutés dans telecharger_document_complet())
    
    ### 4. Affichage des tableaux séparés ###
    st.subheader("Budget prévisionnel de trésorerie")
    st.table(df_part1)
    
    st.subheader("Budget prévisionnel de trésorerie (suite)")
    st.table(df_part2)
    
    # ----------------------------
    # Stockage des résultats dans les données
    # ----------------------------
    
    data["budget_previsionnel_tresorerie"] = {
        "encaissements": encaissements,
        "decaissements": decaissements,
        "total_encaissements": total_encaissements,
        "total_decaissements": total_decaissements,
        "solde_precedent": solde_precedent,
        "solde_mois": solde_mois,
        "solde_tresorerie_cumul": solde_tresorerie_cumul
    }
    
    # Enregistrer les données dans la session
    st.session_state["data"] = data   
    
    # ----------------------------
    # Préparation des données d'exportation pour Budget Prévisionnel de Trésorerie Partie 1
    # ----------------------------
    
    export_table_budget_part1 = []
    headers_part1 = df_part1.columns.tolist()
    export_table_budget_part1.append(dict(zip(["Description"] + headers_part1, [""] + headers_part1)))
    for index, row in df_part1.iterrows():
        export_table_budget_part1.append(dict(zip(["Description"] + headers_part1, [index] + row.tolist())))
    
    # Stocker les données d'exportation pour Partie 1
    st.session_state['export_data_budget_previsionnel_tresorerie_part1'] = {
        "projet": projet,
        "porteur_projet": porteur_projet,
        "table_data": export_table_budget_part1
    }
    
    # ----------------------------
    # Préparation des données d'exportation pour Budget Prévisionnel de Trésorerie Partie 2
    # ----------------------------
    
    export_table_budget_part2 = []
    headers_part2 = df_part2.columns.tolist()
    export_table_budget_part2.append(dict(zip(["Description"] + headers_part2, [""] + headers_part2)))
    for index, row in df_part2.iterrows():
        export_table_budget_part2.append(dict(zip(["Description"] + headers_part2, [index] + row.tolist())))
    
    # Stocker les données d'exportation pour Partie 2
    st.session_state['export_data_budget_previsionnel_tresorerie_part2'] = {
        "projet": projet,
        "porteur_projet": porteur_projet,
        "table_data": export_table_budget_part2
    }
    
    # ----------------------------
    # Section Export
    # ----------------------------
    
    st.header("Exporter les données")
    
    # Bouton pour télécharger le fichier Word complet contenant tous les tableaux avec une clé unique
    if st.button("Télécharger le Document Word Complet", key="download_word_complet_budget_previsionnel_tresorerie"):
        telecharger_document_complet()

# Section 15 : Tableaux d'Analyse Financière
def page_douze_tableaux():
    st.title("Tableaux d'Analyse Financière")
    
    data = st.session_state["data"]
    
    st.markdown("""
    Cette section présente les principaux indicateurs financiers basés sur les données que vous avez saisies.
    """)
    
    total_ca_annee1 = data.get("total_chiffre_affaires_annee1", 0.0)
    total_charges_fixes_annee1 = data.get("total_charges_fixes_annee1", 0.0)
    total_charges_variables = data.get("total_charges_variables", 0.0)
    total_salaires_annee1 = data.get("total_salaires_annee1", 0.0)
    charges_sociales_dirigeant_annee1 = data.get("charges_sociales", {}).get("dirigeants", {}).get("annee1", 0.0)
    charges_sociales_employes_annee1 = data.get("charges_sociales", {}).get("employes", {}).get("annee1", 0.0)
    amortissements_annee1 = data.get("amortissements", {}).get("total", {}).get("annee1", 0.0)
    
    # Calcul du résultat net
    resultat_net = total_ca_annee1 - total_charges_fixes_annee1 - total_charges_variables - total_salaires_annee1 - charges_sociales_dirigeant_annee1 - charges_sociales_employes_annee1 - amortissements_annee1
    
    # Capacité d'autofinancement (simplifiée)
    capacite_autofinancement = resultat_net + amortissements_annee1  # Les amortissements sont réintégrés
    
    st.write(f"**Résultat Net Année 1 :** {resultat_net:.2f} €")
    st.write(f"**Capacité d'Autofinancement Année 1 :** {capacite_autofinancement:.2f} €")
    
    # Vous pouvez répéter les calculs pour les années 2 et 3 si nécessaire
    
    st.write("---")
    
    st.session_state["data"] = data
    
def load_and_split_documents(file_path):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=200)
    raw_documents = PyPDFLoader(file_path).load()
    return text_splitter.split_documents(raw_documents)

def create_faiss_db(documents):
    if not documents:
        raise ValueError("Aucun document trouvé pour créer la base de données FAISS.")
    embeddings = OpenAIEmbeddings(openai_api_key=api_key )
    return FAISS.from_documents(documents, embeddings)

def generate_section(system_message, query, documents, combined_content, tableau_financier):
    memory = ConversationBufferMemory(memory_key='chat_history', return_messages=True)
    llm = ChatOpenAI(openai_api_key=api_key )
    if documents:
        db = create_faiss_db(documents)
        qa_chain = ConversationalRetrievalChain.from_llm(llm, retriever=db.as_retriever(), memory=memory, verbose=True)
        combined_info = qa_chain.run({'question': query})
        full_content = combined_content + " " + combined_info + " " + query+ " "+tableau_financier
    else:
        full_content = combined_content + " " + query+ "Dans ce données où vous allez recuperer les informations generales de l'entreprises "+ tableau_financier+ "utiliser les données financier pour enrichir les arguments aussi sachez que le nom du projet  correspond nom de l'entreprise"
    completion = openai.ChatCompletion.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": system_message},
            {"role": "user", "content": full_content}
        ],
        temperature=0.9
    )
    return completion['choices'][0]['message']['content']

def extract_company_name(text):
    match = re.search(r"(nom de l'entreprise est|Nom de l'entreprise|La vision de) ([\w\s]+)", text, re.IGNORECASE)
    if match:
        return match.group(2).strip()
    return "Nom de l'entreprise non trouvé"

def generate_markdown(results):
    markdown_content = "# Business Plan\n\n"
    for sec_name, content in results.items():
        markdown_content += f"## {sec_name}\n\n"
        paragraphs = content.split('\n')
        for paragraph in paragraphs:
            if paragraph.startswith('- '):  # Points de liste
                markdown_content += f"- {paragraph[2:]}\n"
            elif re.match(r'^\d+\.\s', paragraph):  # Points numérotés
                markdown_content += f"{paragraph}\n"
            else:
                markdown_content += f"{paragraph}\n"
        markdown_content += "\n"

    return markdown_content

def convert_table_to_markdown(table_name, table_data):
    """
    Convertit les données d'une table en format Markdown.
    
    Args:
        table_name (str): Nom de la table.
        table_data (list of dict): Données de la table.
    
    Returns:
        str: Table au format Markdown.
    """
    if not table_data:
        return "Aucune donnée disponible."
    
    # Extraire les en-têtes de colonnes
    headers = list(table_data[0].keys())
    markdown_table = "| " + " | ".join(headers) + " |\n"
    markdown_table += "| " + " | ".join(['---'] * len(headers)) + " |\n"
    
    # Ajouter les lignes
    for row in table_data:
        markdown_table += "| " + " | ".join([str(row.get(header, "")) for header in headers]) + " |\n"
    
    return markdown_table

def convert_all_tables_to_markdown(tables):
    """
    Convertit toutes les tables en une seule chaîne de caractères au format Markdown.
    
    Args:
        tables (dict): Dictionnaire contenant les tables financières.
    
    Returns:
        str: Toutes les tables concaténées en Markdown.
    """
    markdown = ""
    for table_name, table_data in tables.items():
        markdown += f"### {table_name}\n\n"
        markdown += convert_table_to_markdown(table_name, table_data) + "\n\n"
    return markdown


def markdown_to_word_via_text(markdown_content):
    # Créer un nouveau document Word
    doc = Document()
    doc.add_heading('Business Plan', 0)

    # Diviser le contenu en lignes
    lines = markdown_content.split('\n')
    table_data = []
    inside_table = False
    plain_text_output = []  # Pour collecter le texte brut

    for line in lines:
        line = line.strip()
        if not line:
            # Si ligne vide et données de table en cours, ajouter le tableau au document
            if table_data:
                num_cols = len(table_data[0])
                table = doc.add_table(rows=len(table_data), cols=num_cols)
                for i, row in enumerate(table_data):
                    for j, cell in enumerate(row):
                        table.cell(i, j).text = cell.strip()
                table_data = []
                inside_table = False
            continue

        if line.startswith('## '):
            # Sous-titre
            doc.add_heading(line[3:], level=2)
            plain_text_output.append(line[3:])
        elif line.startswith('- '):
            # Liste à puces
            doc.add_paragraph(line[2:], style='List Bullet')
            plain_text_output.append(f"• {line[2:]}")
        elif re.match(r'^\d+\.\s', line):
            # Liste numérotée
            doc.add_paragraph(line, style='List Number')
            plain_text_output.append(line)
        elif line.startswith('|'):
            # Détection des lignes de tableau (évite les lignes de séparation)
            if re.match(r'\|?\s*[-:]+\s*\|', line):
                inside_table = True
                continue  # Ignorer les lignes de séparation
            else:
                inside_table = True
                table_data.append([cell.strip() for cell in line.split('|')[1:-1]])  # Enlever les bords vides et espaces
        elif line.startswith('**') and line.endswith('**'):
            # Texte en gras
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(line[2:-2])
            run.bold = True
            plain_text_output.append(line[2:-2])
        elif not inside_table:
            # Paragraphe normal
            doc.add_paragraph(line)
            plain_text_output.append(line)

    # Traiter les données de table restantes
    if table_data:
        num_cols = len(table_data[0])
        table = doc.add_table(rows=len(table_data), cols=num_cols)
        for i, row in enumerate(table_data):
            for j, cell in enumerate(row):
                table.cell(i, j).text = cell.strip()

    # Sauvegarder le document dans un buffer mémoire
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return "\n".join(plain_text_output), buffer

# Fonction pour convertir un dictionnaire en texte formaté
def format_table_data(data, title):
    if not data:
        return f"{title} : Aucune donnée disponible.\n"
    
    text = f"{title} :\n"
    for key, value in data.items():
        if isinstance(value, dict):
            text += f"  {key} :\n"
            for sub_key, sub_value in value.items():
                text += f"    {sub_key} : {sub_value}\n"
        elif isinstance(value, list):
            text += f"  {key} : {', '.join(map(str, value))}\n"
        else:
            text += f"  {key} : {value}\n"
    return text + "\n"

def page_generation_business_plan():
    st.title("Générateur de Business Plan")

    uploaded_file = st.file_uploader("Téléchargez votre fichier PDF", type="pdf")
    user_text_input = st.text_area("Entrez des informations supplémentaires ou un texte alternatif:", height=200)

    if uploaded_file or user_text_input:
        documents = []
        combined_content = user_text_input  

        if uploaded_file:
            file_path = "uploaded_document.pdf"
            with open(file_path, "wb") as f:
                f.write(uploaded_file.read())
            documents = load_and_split_documents(file_path)    

        # Créer un dictionnaire pour stocker les résultats
        results = {}
        
        # Messages système et requêtes pour chaque section
        system_messages = {
            "Couverture": """
                Générer cette section du business plan:
                Voici les textes à afficher sous forme :
                
                # Canevas de Plans d’Affaires

                Nom du projet ou entreprise
                
                 

            """,
            "Sommaire": """
                Générer cette section du business plan:
                Voici les textes à afficher sous forme de liste:
                ## Sommaire
                I. Résumé Exécutif « Executive Summary » / Pitch
                II. Présentation de votre entreprise/projet
                III. Présentation de l’offre de produit(s) et/ou service(s)  
                IV. Étude de marché
                V. Stratégie marketing, communication et politique commerciale
                VI. Moyens de production et organisation 
                VII. Étude des risques/hypothèses  
                VIII. Plan financier 
                
            """,
            "Résumé Exécutif": """
                Générer cette section du business plan:
                
                ## I. Résumé Exécutif « Executive Summary » / Pitch
                Générer deux grands paragraphes avec plusieurs lignes, l'objectif pour cette section est de :
                Attirer l'attention du lecteur en 5 minutes et lui donner envie d'en savoir plus.
                Décrire le projet en quelques phrases simples et impactantes.
                Ne pas essayer de tout couvrir, soyez concis et précis.

                Les elements clés à generer et qui doivent etre contenue dans les paragraphes:
                - **Présentation de la PME** : Nom de l’entreprise et brève description du service/produit fourni.
                - **Présentation des porteurs de projet** : Nom, prénom, coordonnées, situation de famille, formation et diplômes, expérience professionnelle, activités extra ou para-professionnelles (Joindre CV en annexe).
                - **Potentiel en termes de taille et de profit** : Démontrez par des calculs simples comment votre PME fera du profit.
                - **Votre besoin financier**.

            """,
            "Présentation de votre entreprise": """
                Générer cette section du business plan:

                ## II. Présentation de votre entreprise/projet

                Générer 6 grands paragraphes avec plusieurs lignes, l'objectif pour cette section est de :
                - Parler de votre entreprise/projet de manière plus détaillée.
                - Présenter l’équipe managériale clé.

                Les elements clés à generer et qui doivent etre contenue dans les paragraphes:
                - **Informations générales sur la PME** :
                - Forme juridique : Ets, Sarlu, Sarl, SAS, SA.
                - Siège social : Adresse juridique de l’entreprise.
                - Coordonnées bancaires : Numéro de compte (avec 23 chiffres) de l’entreprise ainsi que la banque où est logé le compte (joindre le Swift Copy).
                - Couverture géographique de l’entreprise et ses activités : lieu d’implantation de l’entreprise et différentes zones couvertes.
                - **Description détaillée de la PME et objectifs de son projet** : Présentez l’entreprise, son origine, introduisez ses atouts/opportunités et enfin décrivez le projet de l’entreprise.
                - **Stade d’avancement de l’entreprise ou du projet** :
                - Décrivez ce qui a été fait et les projets à mener dans le futur.
                - Parlez du niveau de maturité de la PME ou du projet.
                - Lister éventuellement les financements déjà acquis.
                - **Présentation de l’équipe managériale** : Décrivez l’organigramme et l’organisation des ressources humaines, présentez les associés de la PME ainsi que leurs parts sociales.
                - **Analyse SWOT** : Forces, faiblesses, opportunités, contraintes/menaces. de preference ca doit etre presenter sous forme de tableau.
                - **Business Modèle Canevas** : Insérer votre business modèle canevas avec les 9 rubriques bien remplies.

            """,
            "Présentation de l’offre de produit": """
                Générer cette section du business plan :

                ## III. Présentation de l’offre de produit(s) et/ou service(s)
                Générer 6 grands paragraphes avec plusieurs lignes, l'objectif pour cette section est de :
                - Parler de l’offre de produits/services de manière détaillée.
                - Présenter la proposition de valeur différenciante de la PME ou de son offre.

                Les elements clés à generer et qui doivent etre contenue dans les paragraphes:
                - **Noms du/des produit(s) ou service(s)**.
                - **Besoins identifiés** sur le marché auxquels répond votre offre.
                - **Description du/des produit(s) ou service(s)** répondant à ces besoins.
                - **Proposition de valeur unique**.
                - **Prise en compte de l’aspect genre** dans le fonctionnement de la PME ou du projet de l’entreprise.
                - **Prise en compte de l’environnement** :
                - Identification des impacts environnementaux et sociaux des activités de la PME.
                - Mise en place de mesures d’atténuation.
                - Existence d’un Plan de Gestion Environnemental et Social.

            """,
            "Étude de marché": """
                Générer cette section du business plan :

                ## IV. Étude de marché

                Générer 8 grands paragraphes avec plusieurs lignes, l'objectif pour cette section est de :
                - Expliquer la méthode utilisée pour la conduite de l’étude de marché.
                - Présenter les résultats de l’étude de marché.

                Les elements clés à generer et qui doivent etre contenue dans les paragraphes, les numeros doivent etre respecter:
                1. **Description des hypothèses et méthodes de l’étude de marché** :
                - Citer le produit ou service pré-ciblé.
                - Préciser le marché pré-ciblé : secteur d’activité dans lequel le produit s’inscrit.
                - Présenter les méthodes choisies pour réaliser l’étude de marché : questionnaire, étude documentaire, étude de concurrence, étude métier, etc.

                2. **Approche générale du marché (précisez les sources à chaque étape)** :
                - Décrire le marché, ses principales caractéristiques, historique et perspectives.
                - Présenter les résultats : marché cible, marché potentiel, marché réel.
                - Présenter les menaces et opportunités du marché.

                3. **Caractéristiques de la demande** :
                - Présenter le volume de la demande, l’évolution de la demande sur le marché ciblé et les tendances de consommation.
                - Détailler les différents types de clientèle (segmentation).
                - Lister les prescripteurs (partenaires qui peuvent apporter des clients).

                4. **Caractéristiques de l’offre** :
                - Présenter la concurrence directe et indirecte : lister les concurrents et décrire leur offre de services/produits.
                - Lister les points forts et les points faibles de la concurrence : avantages concurrentiels de la concurrence sur le marché.
                - Comment vous différenciez-vous de ces concurrents indirects ?

                5. **Caractéristiques de l’environnement** :
                - Décrire l’environnement des affaires relatif au développement de la PME/projet : le cadre légal, réglementaire, les facteurs externes au marché lui-même, l’évolution des technologies.
                - Lister les menaces et opportunités liées à l’environnement.

                6. **Partenariats** :
                - Préciser les partenariats stratégiques noués ou à mettre en place pour faire croître l’entreprise : il peut s’agir des acteurs en amont et en aval de votre chaîne de production/distribution (fournisseurs, distributeurs, partenaires commerciaux, etc.).

                7. **Création d’emplois** :
                - Démontrer l’impact de la PME/projet en termes d’emplois directs déjà créés ou à créer.

                8. **Chiffre d’affaires** :
                - Préciser la part de marché visée et le volume de chiffre d’affaires prévisible à horizon 1 an, 2 ans, 3 ans.

            """,
            "Stratégie Marketing":  """
                Générer cette section du business plan :

                ## V. Stratégie Marketing, Communication et Politique Commerciale

                Générer cette section, l'objectif pour cette section est de :
                - Présenter la stratégie marketing et commerciale à court et moyen terme.

                Les elements clés à generer et qui doivent etre contenue dans les paragraphes, les numeros doivent etre respecter:
                1. **Choix de segments de clientèle** :
                - Expliquer quels segments de clientèle vont constituer la cible de la PME/projet et pourquoi ce choix.
                - Expliquer dans les grandes lignes le positionnement stratégique.

                2. **Marketing-mix (4P : Produit – Prix – Place – Promotion)** :
                - Présenter la politique marketing générale :
                    - Choix du nom, du logo et des couleurs.
                    - Choix du message, du slogan.
                - Tableau synthétique des segments :

                    | Segment de clientèle | Produit proposé | Positionnement en termes de prix | Lieu de distribution | Style et mode de communication |
                    |-----------------------|-----------------|----------------------------------|-----------------------|---------------------------------|
                    | Segment 1            |                 |                                  |                       |                                 |
                    | Segment 2            |                 |                                  |                       |                                 |
                    | Segment 3            |                 |                                  |                       |                                 |

                3. **Plan Marketing et actions commerciales**  :
                - Présenter le plan marketing : lister les actions commerciales et actions de communication prévues ; inscrire leur coût si possible.

                    | Types d’actions       | Janvier | Février | Mars | ... | Décembre |
                    |-----------------------|---------|---------|------|-----|----------|
                    | Action 1             |         |         |      |     |          |
                    | Action 2             |         |         |      |     |          |

                4. **Moyens et partenaires sollicités** :
                - Lister les moyens à mettre en œuvre et les partenaires sollicités pour les actions commerciales et de communication.

            """,
            "Moyens de production et organisation": """
                Générer cette section du business plan:

                ## VI. Moyens de production et organisation

                Générer 4 grands paragraphes avec plusieurs lignes, l'objectif pour cette section est de :
                - Spécifier les moyens humains et matériels à disposition de la PME.

                Les elements clés à generer et qui doivent etre contenue dans les paragraphes:
                - **Locaux** :
                - Liste des locaux, bail de location, conditions négociées, coût, utilité.
                - **Matériel** :
                - Liste, mode d’acquisition ou de location, coût, utilité, renouvellement.
                - **Moyens humains** :
                - Personnel, plannings, horaires, coût, charges sociales ; indiquer une répartition claire des tâches.
                - **Fournisseurs et sous-traitants** :
                - Liste des fournisseurs et/ou sous-traitants, devis obtenus, tarifs, conditions négociées.

            """,
            "Étude des risques": """
                Générer cette section du business plan:

                ## VII. Étude des risques/hypothèses

                Générer cette section, l'objectif pour cette section est de :
                - Présenter la synthèse des risques et mesures d’atténuation identifiés quant au développement de la PME/projet.

                Les elements clés à generer et qui doivent etre contenue dans les paragraphes:
                - **Tableau des risques** :

                | Nature de risque          | Description              | Stratégie de traitement    |
                |---------------------------|--------------------------|----------------------------|
                | Risques liés à l’environnement général |                          |                            |
                | Risques liés au marché    |                          |                            |
                | Risques liés aux outils   |                          |                            |
                | Risques liés aux personnes |                          |                            |
                | Risques liés aux tiers    |                          |                            |
                | Autres risques (spécifiez) |                          |                            |

                Étude des risques/hypothèses:

            """,
            "Annexes": """
                Générer cette section du business plan:
                
                ## VII. Étude des risques/hypothèses

                ### Objectif
                - Présenter la synthèse des risques et mesures d’atténuation identifiés quant au développement de la PME/projet.

                ### Contenu attendu
                - **Tableau des risques** :

                | Nature de risque          | Description              | Stratégie de traitement    |
                |---------------------------|--------------------------|----------------------------|
                | Risques liés à l’environnement général |                          |                            |
                | Risques liés au marché    |                          |                            |
                | Risques liés aux outils   |                          |                            |
                | Risques liés aux personnes |                          |                            |
                | Risques liés aux tiers    |                          |                            |
                | Autres risques (spécifiez) |                          |                            |

            """,
            "Annexes": """
                Générer cette section du business plan:

                7 – ANNEXES
                Renvoyer en annexe les documents trop volumineux ou difficiles à lire : - - - -
                étude de marché complète,
                contrats,
                conditions

                Annexes du projet:

            """
        }

        queries = {
            "Couverture": "Afficher seulement le texte fournies",
            "Sommaire": "Afficher seulement le texte fournises",
            "Résumé Exécutif": "Décrire brièvement le projet, son potentiel de profit et les qualifications de l'équipe.",
            "Présentation de votre entreprise": "Fournir une analyse détaillée de l'entreprise, incluant son origine, ses objectifs et son organisation.",
            "Présentation de l’offre de produit": "Décrire les produits ou services, leur proposition de valeur unique, et les besoins du marché qu'ils adressent.",
            "Étude de marché": "Analyser le marché cible, les tendances de consommation, et la concurrence directe et indirecte.",
            "Stratégie Marketing": "Décrire la stratégie marketing, y compris les segments cibles, le positionnement, le mix marketing (Produit, Prix, Place, Promotion) et les actions commerciales prévues.",
            "Moyens de production et organisation": "Décrire les moyens humains et matériels, ainsi que l'organisation opérationnelle de l'entreprise.",
            "Étude des risques": "Identifier les risques potentiels et proposer des stratégies pour les atténuer.",
            "Annexes": "Inclure tous les documents annexes pertinents pour étayer le plan d'affaires."
        }

        # Espaces réservés pour chaque section
        placeholders = {name: st.empty() for name in system_messages.keys()}
        
        data = st.session_state.get("data", {})
        tables = data.get("tables", {})
        
        #st.write(data)
        # st.write(tables)
        
        
        
        # Récupérer les données exportées de toutes les sections
        # Récupérer les données exportées de toutes les sections 
        export_data_investissements = st.session_state.get('export_data_investissements', {})
        export_data_salaires = st.session_state.get('export_data_salaires_charges_sociales', {})
        export_data_amortissements = st.session_state.get('export_data_detail_amortissements', {})
        export_data_compte = st.session_state.get('export_data_compte_resultats_previsionnel', {})
        export_data_soldes = st.session_state.get('export_data_soldes_intermediaires_de_gestion', {})
        export_data_capacite = st.session_state.get('export_data_capacite_autofinancement', {})
        export_data_seuil = st.session_state.get('export_data_seuil_rentabilite_economique', {})
        export_data_bfr = st.session_state.get('export_data_besoin_fonds_roulement', {})
        export_data_plan_financement = st.session_state.get('export_data_plan_financement_trois_ans', {})
        export_data_budget_part1 = st.session_state.get('export_data_budget_previsionnel_tresorerie_part1', {})
        export_data_budget_part2 = st.session_state.get('export_data_budget_previsionnel_tresorerie_part2', {})

        # Concaténer toutes les sections
        final_text = ""
        final_text += format_table_data(export_data_investissements, "Investissements et financements")
        final_text += format_table_data(export_data_salaires, "Salaires et Charges Sociales")
        final_text += format_table_data(export_data_amortissements, "Détail des Amortissements")
        final_text += format_table_data(export_data_compte, "Compte de résultats prévisionnel")
        final_text += format_table_data(export_data_soldes, "Soldes intermédiaires de gestion")
        final_text += format_table_data(export_data_capacite, "Capacité d'autofinancement")
        final_text += format_table_data(export_data_seuil, "Seuil de rentabilité économique")
        final_text += format_table_data(export_data_bfr, "Besoin en fonds de roulement")

        # Ajouter les nouvelles sections
        final_text += format_table_data(export_data_plan_financement, "Plan de financement à trois ans")
        final_text += format_table_data(export_data_budget_part1, "Budget prévisionnel de trésorerie")
        final_text += format_table_data(export_data_budget_part2, "Budget prévisionnel de trésorerie(suite)")

        

        # Générer toutes les sections automatiquement
       # Génération du Business Plan et téléchargement des fichiers
        for section_name in system_messages.keys():
            with st.spinner(f"Génération de {section_name}..."):
                system_message = system_messages[section_name]
                query = queries[section_name]
                
                try:
                    # Vérifier si la section est "Couverture" ou "Sommaire"
                    if section_name in ["Couverture", "Sommaire"]:
                        results[section_name] = generate_section(system_message, query, documents, combined_content, "")
                    else:
                        results[section_name] = generate_section(system_message, query, documents, combined_content, final_text)
                except ValueError as e:
                    results[section_name] = f"Erreur: {str(e)}"
                combined_content += " " + results[section_name]
                placeholders[section_name].markdown(f"\n\n{results[section_name]}")

        # Extraction du nom de l'entreprise
        #first_section_content = results.get("Résumé Exécutif", "")
        #company_name = extract_company_name(first_section_content)

        # Créer le contenu Markdown principal
        markdown_content = generate_markdown(results)
        
        #content_result, word_buffer = markdown_to_word_via_text(arkdown_contentm)

        # Ajouter les tableaux au contenu Markdown
        #  markdown_content = append_tables_to_markdown(markdown_content)

        # Générer le PDF
        pdf = MarkdownPdf(toc_level=2)
        pdf.add_section(Section(markdown_content))
        pdf.meta["title"] = "Business Plan" 
        pdf_file_path = "business_plan.pdf"
        pdf.save(pdf_file_path)
        
        
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement


        from docx import Document
        import re

        def markdown_to_word_via_text(markdown_contents, doc):
            doc.add_heading('Business Plan', 0)

            # Diviser le contenu en lignes
            lines = markdown_contents.split('\n')
            table_data = []
            inside_table = False

            for line in lines:
                line = line.strip()
                if not line:
                    # Si ligne vide et données de table en cours, ajouter le tableau au document
                    if table_data:
                        add_table_with_borders(doc, table_data)
                        table_data = []
                        inside_table = False
                    continue

                if line.startswith('# '):  # Titre niveau 1
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):  # Titre niveau 2
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):  # Titre niveau 3
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('#### '):  # Titre niveau 4
                    doc.add_heading(line[5:], level=4)
                elif re.match(r'^\d+\.\s', line):  # Liste numérotée
                    # Vérifier s'il y a du texte en gras dans la liste numérotée
                    match = re.match(r'^(\d+\.\s)(\*\*.+?\*\*)', line)
                    if match:
                        paragraph = doc.add_paragraph(style='List Number')
                        paragraph.add_run(match.group(1))  # Numéro
                        bold_run = paragraph.add_run(match.group(2)[2:-2])  # Texte en gras sans `**`
                        bold_run.bold = True
                    else:
                        doc.add_paragraph(line, style='List Number')
                elif line.startswith('- ') or line.startswith('•'):  # Liste à puces
                    match = re.match(r'^(•|-)\s\*\*(.+?)\*\*(.*)', line)
                    if match:
                        paragraph = doc.add_paragraph(style='List Bullet')
                        bold_run = paragraph.add_run(match.group(2))  # Texte en gras
                        bold_run.bold = True
                        if match.group(3):  # Texte après le gras
                            paragraph.add_run(match.group(3).strip())
                    else:
                        doc.add_paragraph(line[2:], style='List Bullet')
                elif line.startswith('|'):  # Détection des lignes de tableau
                    if re.match(r'\|?\s*[-:]+\s*\|', line):
                        inside_table = True
                        continue  # Ignorer les lignes de séparation
                    else:
                        inside_table = True
                        table_data.append([cell.strip() for cell in line.split('|')[1:-1]])  # Enlever les bords vides et espaces
                elif re.match(r'^\*\*.+?\*\*\s*:', line):  # Texte en gras suivi de texte normal
                    match = re.match(r'^\*\*(.+?)\*\*\s*:(.*)', line)
                    if match:
                        paragraph = doc.add_paragraph()
                        bold_run = paragraph.add_run(match.group(1))  # Texte en gras
                        bold_run.bold = True
                        if match.group(2):  # Texte normal après le `:`
                            paragraph.add_run(f":{match.group(2)}")
                elif re.match(r'^\*\*.+?\*\*$', line):  # Texte entièrement en gras
                    paragraph = doc.add_paragraph()
                    bold_run = paragraph.add_run(line[2:-2])  # Texte sans `**`
                    bold_run.bold = True
                elif re.match(r'^\*\*.+?\*\*\s[\d.,]+\s?[€$%]$', line):  # Nombres avec symboles monétaires
                    match = re.match(r'^\*\*(.+?)\*\*\s([\d.,]+\s?[€$%])$', line)
                    if match:
                        paragraph = doc.add_paragraph()
                        bold_run = paragraph.add_run(match.group(1))  # Texte en gras
                        bold_run.bold = True
                        paragraph.add_run(f" {match.group(2)}")  # Montant avec symbole
                elif not inside_table:  # Paragraphe normal
                    doc.add_paragraph(line)

            # Traiter les données de table restantes
            if table_data:
                add_table_with_borders(doc, table_data)

        def add_table_with_borders(doc, table_data):
            """
            Ajoute un tableau au document Word avec bordures et gestion du texte en gras dans les cellules.
            """
            num_cols = len(table_data[0])
            table = doc.add_table(rows=len(table_data), cols=num_cols)
            table.style = 'Table Grid'  # Appliquer un style de tableau avec bordures

            for i, row in enumerate(table_data):
                for j, cell in enumerate(row):
                    cell_content = table.cell(i, j).paragraphs[0]
                    parts = re.split(r'(\*\*.+?\*\*)', cell)  # Diviser par texte en gras
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):  # Texte en gras
                            run = cell_content.add_run(part[2:-2])
                            run.bold = True
                        else:  # Texte normal
                            cell_content.add_run(part.strip())




        # Ajouter la création et le téléchargement du fichier Word
        
        doc = Document()
        markdown_to_word_via_text(markdown_content, doc)
        # doc.add_paragraph(content_result)
        
        # Vérifier et ajouter le contenu
       

        # Ajouter les sections du Business Plan
        """ for section_name, content in results.items():
            doc.add_heading(section_name, level=1)
            doc.add_paragraph(content)"""

        # Récupérer les données des tableaux depuis la session Streamlit
       # Récupérer les données des tableaux depuis la session Streamlit
        export_data_investissements = st.session_state.get('export_data_investissements', {})
        export_data_salaires = st.session_state.get('export_data_salaires_charges_sociales', {})
        export_data_amortissements = st.session_state.get('export_data_detail_amortissements', {})
        export_data_compte = st.session_state.get('export_data_compte_resultats_previsionnel', {})
        export_data_soldes = st.session_state.get('export_data_soldes_intermediaires_de_gestion', {})
        export_data_capacite = st.session_state.get('export_data_capacite_autofinancement', {})
        export_data_seuil = st.session_state.get('export_data_seuil_rentabilite_economique', {})
        export_data_bfr = st.session_state.get('export_data_besoin_fonds_roulement', {})

        # Ajouter une section pour les tableaux
        doc.add_heading('Résumé des Données Financières', level=1)

        # Fonction pour ajouter un tableau dans le document Word
        def ajouter_tableau(donnees, headers, titre):
            """
            Ajoute un tableau au document Word avec bordures et gestion du texte en gras dans les cellules.
            """
            doc.add_heading(titre, level=2)
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Light List Accent 1'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Ajouter les en-têtes
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                for paragraph in hdr_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Ajouter les données des tableaux
            for row in donnees:
                row_cells = table.add_row().cells
                for i, header in enumerate(headers):
                    cell_value = row.get(header, "")
                    cell_text = str(cell_value)  # Convertir en chaîne de caractères
                    row_cells[i].text = cell_text
                    row_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

            # Ajouter une note
            doc.add_paragraph()
            doc.add_paragraph("Les résultats sont calculés selon les données fournies.")


        # Ajouter les différents tableaux
        if export_data_investissements.get("table_data"):
            ajouter_tableau(export_data_investissements["table_data"], ["Investissements", "Taux (%)", "Durée (mois)", "Montant (€)"], "Investissements et Financements")
        if export_data_salaires.get("table_data"):
            ajouter_tableau(export_data_salaires["table_data"], ["Description", "Année 1", "Année 2", "Année 3"], "Salaires et Charges Sociales")
        if export_data_amortissements.get("amortissements"):
            ajouter_tableau(export_data_amortissements["amortissements"], ["Amortissement", "Année 1", "Année 2", "Année 3"], "Détail des Amortissements")
        if export_data_compte.get("table_data"):
            ajouter_tableau(export_data_compte["table_data"], ["Description", "Année 1", "Année 2", "Année 3"], "Compte de Résultats Prévisionnel")
        if export_data_soldes.get("table_data"):
            ajouter_tableau(export_data_soldes["table_data"], ["Description", "Année 1", "Année 2", "Année 3"], "Soldes Intermédiaires de Gestion")
        if export_data_capacite.get("table_data"):
            ajouter_tableau(export_data_capacite["table_data"], ["Description", "Année 1", "Année 2", "Année 3"], "Capacité d'Autofinancement")
        if export_data_seuil.get("table_data"):
            ajouter_tableau(export_data_seuil["table_data"], ["Description", "Année 1", "Année 2", "Année 3"], "Seuil de Rentabilité Économique")
        if export_data_bfr.get("table_data"):
            ajouter_tableau(export_data_bfr["table_data"], ["Analyse clients / fournisseurs", "Délai jours", "Année 1", "Année 2", "Année 3"], "Besoin en Fonds de Roulement")

        # **Nouvelles sections ajoutées :**

        # Ajouter la section Plan de Financement à Trois Ans
        export_data_plan_financement = st.session_state.get('export_data_plan_financement_trois_ans', {})
        if export_data_plan_financement.get("table_data"):
            ajouter_tableau(
                export_data_plan_financement["table_data"],
                ["Description", "Année 1", "Année 2", "Année 3"],
                "Plan de Financement à Trois Ans"
            )

        # Ajouter la section Budget Prévisionnel de Trésorerie Partie 1
        export_data_budget_part1 = st.session_state.get('export_data_budget_previsionnel_tresorerie_part1', {})
        if export_data_budget_part1.get("table_data"):
            ajouter_tableau(
                export_data_budget_part1["table_data"],
                ["Description", "Mois 1", "Mois 2", "Mois 3", "Mois 4", "Mois 5", "TOTAL"],
                "Budget Prévisionnel de Trésorerie - Partie 1"
            )

        # Ajouter la section Budget Prévisionnel de Trésorerie Partie 2
        export_data_budget_part2 = st.session_state.get('export_data_budget_previsionnel_tresorerie_part2', {})
        if export_data_budget_part2.get("table_data"):
            ajouter_tableau(
                export_data_budget_part2["table_data"],
                ["Description", "Mois 6", "Mois 7", "Mois 8", "Mois 9", "Mois 10", "Mois 11", "Mois 12", "TOTAL"],
                "Budget Prévisionnel de Trésorerie - Partie 2"
            )

        # Enregistrer le document dans un buffer
        word_buffer = BytesIO()
        doc.save(word_buffer)
        word_buffer.seek(0)

        # Télécharger les fichiers générés
        st.success("Le PDF et le document Word ont été générés avec succès.")
        with open(pdf_file_path, "rb") as f:
            st.download_button("Téléchargez le PDF", f, file_name="business_plan.pdf", mime="application/pdf")

        st.download_button("Téléchargez le document Word", word_buffer, file_name="business_plan.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# Mise à jour des noms d'onglets
# Mise à jour des noms d'onglets
tab_names = [
    "Informations Générales", "Besoins de Démarrage", "Financement",
    "Charges Fixes", "Chiffre d'Affaires", "Charges Variables",
    "Fonds de Roulement", "Salaires", "Rentabilité", "Trésorerie", "Récapitulatif",
    "Investissements et Financements", "Salaires et Charges Sociales", "Détail des Amortissements",
    "Compte de Résultats Prévisionnel", "Soldes Intermédiaires de Gestion",
    "Capacité d'Autofinancement", "Seuil de Rentabilité Économique",
    "Besoin en Fonds de Roulement", "Plan de Financement sur 3 Ans",
    "Budget Prévisionnel de Trésorerie","Génération du Business Plan", "Tableaux d'Analyse Financière"
]

# Mise à jour de la liste des fonctions correspondantes
sections = [
    page_informations_generales, page_besoins_demarrage, page_financement,
    page_charges_fixes, page_chiffre_affaires, page_charges_variables,
    page_fonds_roulement, page_salaires, page_rentabilite, page_tresorerie,
    page_recapitulatif, page_investissements_et_financements,
    page_salaires_charges_sociales, page_detail_amortissements,
    page_compte_resultats_previsionnel, page_soldes_intermediaires_de_gestion,
    page_capacite_autofinancement, page_seuil_rentabilite_economique,
    page_besoin_fonds_roulement, page_plan_financement_trois_ans,
    page_budget_previsionnel_tresorerie,page_generation_business_plan, page_douze_tableaux
]


# Création des onglets
tabs = st.tabs(tab_names)

# Parcours des onglets
for i, tab in enumerate(tabs):
    with tab:
        sections[i]()
